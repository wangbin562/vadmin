# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
html转vadmin
html转doc
1、暂不支持表格
"""
import xml.etree.cElementTree as ET

import docx
from docx import Document
from docx import shared
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.opc import constants
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from utils.word_html_api import CssStyleApi
from utils.word_html_api import DocApi
from utils.word_html_api import FontApi
from utils.word_html_api import UnitApi
from vadmin import widgets


class Vadmin2Word(object):
    def __init__(self, lst_panel):
        self.doc = Document()
        section = self.doc.sections[0]
        pt = UnitApi.px_2_pt(120)
        section.left_margin = shared.Pt(pt)
        section.right_margin = shared.Pt(pt)
        section.page_width = shared.Pt(UnitApi.px_2_pt(794))
        self.lst_panel = lst_panel
        self.table = None
        self.cell = None
        self.paragraph = None
        self.lst_paragraph = []
        self.content_width = 554

    def set_text_style(self, run, o_widget):
        """
        设置文字样式
        """
        # 设置字体
        if o_widget.font_family:
            run.font.name = FontApi.html_2_doc(o_widget.font_family)
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)  # 中文加这句设置才会生效
            except (BaseException,):
                pass

        # 设置文字大小
        if o_widget.text_size:
            pt = UnitApi.px_2_pt(o_widget.text_size)
            run.font.size = shared.Pt(pt)

        # 设置文字颜色
        if o_widget.text_color:
            color = o_widget.text_color
            r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            run.font.color.rgb = shared.RGBColor(r, g, b)

        # 加粗
        if o_widget.bold:
            run.font.bold = True

        # 斜体
        if o_widget.italics:
            run.font.italic = True

        # 删除线
        if o_widget.del_line:
            run.font.strike = True

        # 下划线
        if o_widget.under_line:
            run.font.underline = True

        if o_widget.background_color:
            DocApi.set_text_background_color(run, o_widget.background_color)

    def add_hyperlink(self, o_widget):
        """
        增加链接对象
        """
        src = ""
        if o_widget.url:
            src = o_widget.url.strip()
        elif o_widget.href:
            src = o_widget.href.strip()
        # if src.find("http") != 0:
        #     src = "%s%s" % (CON_BASE_URL, url)
        src = src.strip("(){}\r\n;")
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = self.paragraph.part
        r_id = part.relate_to(src, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = o_widget.text
        hyperlink.append(new_run)

        self.paragraph._p.append(hyperlink)
        r = self.paragraph.add_run()
        r._r.append(hyperlink)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
        self.set_text_style(r, o_widget)

        return hyperlink

    def add_text(self, o_widget):
        run = self.paragraph.add_run(o_widget.text)
        self.set_text_style(run, o_widget)

    def add_image(self, o_widget):
        pass

    def add_table(self, o_widget):
        row_num = len(o_widget.data)
        if o_widget.data:
            col_num = len(o_widget.data[0])
        else:
            col_num = 0

        # 表格在DIV下面，要先删除1个
        paragraph = self.lst_paragraph[-1]
        del self.lst_paragraph[-1]
        DocApi.delete_paragraph(paragraph)

        self.table = self.doc.add_table(row_num, col_num, style='TableGrid')
        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.table.autofit = False

        # 设置列宽
        table_width = 0.0
        for col_idx, width in o_widget.col_width.items():
            col = self.table.columns[col_idx]
            px = int(width.strip("px"))
            table_width += px
            pt = UnitApi.px_2_pt(px)
            col.width = shared.Pt(pt)

        # 设置行高
        for row_idx, height in o_widget.row_height.items():
            # px = int(height.strip("px"))
            pt = UnitApi.px_2_pt(height)
            row = self.table.rows[row_idx]
            DocApi.set_row_height(row, pt, "atLeast")

        DocApi.set_table_width(self.table, table_width)

        # 合并单元格
        for merged in o_widget.merged_cells:
            begin_row_col, end_row_col = merged.split(":")
            begin_row, begin_col = begin_row_col.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)

            end_row, end_col = end_row_col.split("-")
            end_row = int(end_row)
            end_col = int(end_col)

            cell = self.table.cell(begin_row, begin_col)
            b_cell = self.table.cell(end_row, end_col)
            cell.merge(b_cell)

        for row_idx, row_data in enumerate(o_widget.data):
            for col_idx, col_data in enumerate(row_data):
                if col_data is None:
                    continue

                self.cell = self.table.cell(row_idx, col_idx)
                self.paragraph = self.cell.paragraphs[-1]  # cell新建后默认有第一个paragraph
                DocApi.set_cell_vertical(self.cell, col_data.vertical)
                DocApi.set_cell_horizontal(self.cell, col_data.horizontal)
                DocApi.set_cell_background_color(self.cell, col_data.background_color)
                for i, o_sub in enumerate(col_data.get_widget()):
                    if i > 0:
                        # lst_widget = o_sub.get_widget()
                        if o_sub.height or o_sub.min_height:  # 如果没有子图例且高度为0的情况下，不新建paragraph
                            self.paragraph = self.cell.add_paragraph()

                    self.widget_2_run(o_sub)

        self.table = None

    def create_paragraph(self, o_widget):
        if self.table is None:
            self.paragraph = self.doc.add_paragraph()
            self.lst_paragraph.append(self.paragraph)
        # else:
        #     self.paragraph = self.cell.add_paragraph()

        if o_widget.horizontal == "center":
            self.paragraph.alignment = WD_TAB_ALIGNMENT.CENTER

        elif o_widget.horizontal == "right":
            self.paragraph.alignment = WD_TAB_ALIGNMENT.RIGHT

        if o_widget.margin_left:
            pt = UnitApi.px_2_pt(o_widget.margin_left)
            self.paragraph.paragraph_format.left_indent = shared.Pt(pt)

        if o_widget.width:
            try:
                width = int(o_widget.width)
                if width < self.content_width:
                    pt = UnitApi.px_2_pt(self.content_width - o_widget.width)
                    self.paragraph.paragraph_format.right_indent = shared.Pt(pt)
            except (BaseException,):
                pass

        # if parent:
        lst_widget = o_widget.get_widget()
        if lst_widget:
            if isinstance(lst_widget[0], widgets.ColSpacing):
                pt = UnitApi.px_2_pt(lst_widget[0].spacing)
                self.paragraph.paragraph_format.first_line_indent = shared.Pt(pt)

            if isinstance(lst_widget[-1], widgets.ColSpacing):
                pt = UnitApi.px_2_pt(lst_widget[-1].spacing)
                self.paragraph.paragraph_format.right_indent = shared.Pt(pt)

            if o_widget.min_height:
                max_size = 0.0
                for o_sub in lst_widget:
                    if isinstance(o_sub, widgets.Text):
                        if o_sub.text_size > max_size:
                            max_size = o_sub.text_size

                if o_widget.min_height > max_size:
                    pt = UnitApi.px_2_pt((o_widget.min_height - max_size) / 2.0)
                    self.paragraph.paragraph_format.space_after = shared.Pt(pt)
                    self.paragraph.paragraph_format.space_before = shared.Pt(pt)

    def widget_2_run(self, o_widget):
        if isinstance(o_widget, widgets.Text):
            if o_widget.url or o_widget.href:
                self.add_hyperlink(o_widget)
            else:
                self.add_text(o_widget)

        elif isinstance(o_widget, widgets.Image):
            self.add_image(o_widget)

        elif isinstance(o_widget, widgets.Panel):
            # lst_widget = o_widget.get_widget()
            if o_widget.height or o_widget.min_height:
                self.create_paragraph(o_widget)

            for o_sub in o_widget.get_widget():
                self.widget_2_run(o_sub)

        elif isinstance(o_widget, widgets.LayoutTable):
            self.add_table(o_widget)

        # elif isinstance(o_widget, widgets.Row):
        #     run = self.paragraph.add_run()
        #     run.add_break()

    def to_word(self, path):
        for o_widget in self.lst_panel:
            if isinstance(o_widget, widgets.LayoutTable):
                self.add_table(o_widget)
            else:
                self.create_paragraph(o_widget)
                for o_sub in o_widget.get_widget():
                    self.widget_2_run(o_sub)

        self.doc.save(path)


class Html(object):
    """
    1、新建页面默认8.5英寸 7772400 / 914400.0 doc.sections[0].page_width
    section = self.doc.sections[0]
    section.page_height
    10058400
    section.page_width
    7772400
    section.left_margin
    1143000
    section.right_margin
    1143000
    2、
    """

    def __init__(self, content):
        self.content = content
        # self.doc = None
        # self.dict_page_style = {}  # 页面样式，使用html关键字
        # self.dict_paragraph_style = {}  # 段落样式，使用html关键字
        # self.dict_element_style = {}  # 节点样式，使用html关键字
        # self.begin_paragraph = False
        # self.paragraph = None
        # self.table = None
        # self.cell = None
        self.lst_style = []
        self.lst_label = []
        self.lst_panel = []
        # self.parent = None
        self.table = None

    def _parse_style(self, element):
        style = element.get("style")
        dict_style = element.attrib or {}
        if style is None:
            return dict_style

        del dict_style["style"]
        lst_style = style.strip(";").split(";")
        for s_style in lst_style:
            s_style = s_style.strip()
            if s_style == "":
                continue

            #  'inherit !important'
            try:
                key, val = s_style.split(":")
            except (BaseException,):
                continue

            val = val.replace("!important", "").strip()
            if val == "":
                continue

            dict_style[key.strip().lower()] = val

        return dict_style

    def make_style(self):
        dict_style = {}
        for style in self.lst_style:
            if style:
                dict_style.update(style)

        return dict_style

    def create_panel(self, dict_style):
        # o_panel = widgets.Panel(width="100%", vertical="center", border_color="#FF0000")
        o_panel = widgets.Panel(width="100%", vertical="center")
        if "text-align" in dict_style:
            o_panel.horizontal = dict_style["text-align"]

        if "vertical-align" in dict_style:
            vertical = dict_style["vertical-align"]
            if vertical == "middle":
                vertical = "center"
            o_panel.vertical = vertical

        if "line-height" in dict_style:
            o_panel.min_height = CssStyleApi.format_unit(dict_style["line-height"], dict_style)

        if "min-height" in dict_style:
            o_panel.min_height = CssStyleApi.format_unit(dict_style["min-height"], dict_style)

        if "height" in dict_style:
            o_panel.height = CssStyleApi.format_unit(dict_style["height"], dict_style)

        if "width" in dict_style:
            o_panel.width = CssStyleApi.format_unit(dict_style["width"], dict_style)

        if "background" in dict_style:
            o_panel.background_color = CssStyleApi.format_color(dict_style["background"])

        if "background-color" in dict_style:
            o_panel.background_color = CssStyleApi.format_color(dict_style["background-color"])

        margin = CssStyleApi.format_margin_padding(dict_style)
        if margin != [0, 0, 0, 0]:
            o_panel.margin = margin

        padding = CssStyleApi.format_margin_padding(dict_style, "padding")
        if padding != [0, 0, 0, 0]:
            o_panel.padding = padding

        # if "position" in dict_style:
        #     if dict_style["position"] == "absolute":
        #         if "top" in dict_style["top"]:
        #             o_panel.y = self.format_unit(dict_style["top"], dict_style)
        #
        #         if "right" in dict_style["right"]:
        #             o_panel.y = -1 * CssStyleApi.format_unit(dict_style["top"], dict_style)
        if "text-indent" in dict_style:
            px = CssStyleApi.format_unit(dict_style["text-indent"], dict_style)
            o_panel.append(widgets.ColSpacing(px))
        return o_panel

    def create_text(self, text, dict_style, parent=None):
        o_text = widgets.Text(text=text, text_horizontal="left", inline=True)
        if "color" in dict_style:
            o_text.text_color = CssStyleApi.format_color(dict_style["color"])

        if "font-size" in dict_style:
            o_text.text_size = CssStyleApi.format_unit(dict_style["font-size"], dict_style)
        else:
            o_text.text_size = 16

        if "href" in dict_style:
            o_text.under_line = True
            if o_text.text_color is None:
                o_text.text_color = "#0563C1"
            o_text.href = dict_style["href"]

        if "text-decoration-line" in dict_style or "text-decoration" in dict_style:
            line = dict_style.get("text-decoration-line", None) or dict_style["text-decoration"]
            if line == "none":
                o_text.under_line = False
            elif line == "underline":  # 下划线
                o_text.under_line = True
            elif line == "line-through":  # 删除线
                o_text.del_line = False

        if "font-family" in dict_style:
            o_text.font_family = FontApi.html_2_doc(dict_style["font-family"])
        else:
            o_text.font_family = "宋体"

        if "font-weight" in dict_style:
            if dict_style["font-weight"] in ["normal", "lighter", "inherit"]:
                pass
            elif dict_style["font-weight"] in ["bold", "bolder"]:
                o_text.bold = True
            else:
                if int(dict_style["font-weight"]) > 400:
                    o_text.bold = True

        if "background" in dict_style:
            o_text.background_color = CssStyleApi.format_color(dict_style["background"])

        if "background-color" in dict_style:
            o_text.background_color = CssStyleApi.format_color(dict_style["background-color"])

        if "display" in dict_style:
            if dict_style["display"] == "display:inline-block" and "width" in dict_style:
                o_text.margin_left = CssStyleApi.format_color(dict_style["width"])

        if parent:
            if o_text.text_size and "line-height" in dict_style and "%" in dict_style["line-height"]:
                height = dict_style["line-height"]
                height = round(int(height.replace("%", "")) / 100.0 * o_text.text_size)
                parent.min_height = height

        return o_text

    def create_col_spacing(self, dict_style, parent):
        if "display" in dict_style:
            if dict_style["display"] == "inline-block" and "width" in dict_style:
                px = CssStyleApi.format_unit(dict_style["width"], dict_style)
                return widgets.ColSpacing(px)

    def create_image(self, dict_style, parent=None):
        o_image = widgets.Image()
        if "src" in dict_style:
            o_image.url = dict_style["src"]

        return o_image

    def create_table(self, root, dict_style):
        col_width = {}
        row_height = {}
        merged_cells = []
        data = []
        table = widgets.LayoutTable(col_width=col_width, row_height=row_height, merged_cells=merged_cells,
                                    row_border=True, col_border=True, data=data)
        for element in root.getchildren():
            if element.tag in ["colgroup"]:
                for i, o_sub in enumerate(element.getchildren()):
                    col_width[i] = o_sub.attrib["width"]

            elif element.tag in ["tbody"]:
                row = 0
                for row_idx, o_tr in enumerate(element.getchildren()):
                    col = 0
                    lst_td = []
                    for o_td in o_tr.getchildren():
                        if o_td.tag in ["td"]:
                            add_col = 0
                            if "colspan" in o_td.attrib or "rowspan" in o_td.attrib:
                                add_row = int(o_td.get("rowspane", 1)) - 1
                                add_col = int(o_td.get("colspan", 1)) - 1
                                merged = "%s-%s:%s-%s" % (row, col, row + add_row, col + add_col)
                                merged_cells.append(merged)
                                row += add_row

                            self._parse_element(o_td, lst_td)
                            if add_col:
                                col += add_col
                                lst_td.extend([None] * add_col)
                            col += 1
                    height = 0
                    for o_td in lst_td:
                        if o_td:
                            if o_td.height and o_td.height > height:
                                height = o_td.height

                    if height > 0:
                        row_height[row_idx] = height

                    data.append(lst_td)
                    row += 1

            # elif element.tag in ["tr", "tbody"]:
            #     lst_td = []
            #     for o_sub in element.getchildren():
            #         if o_sub.tag in ["td"]:
            #             self._parse_element(o_sub, lst_td)

        # self._parse_element(root, table)
        return table

    def add_text(self, parent, o_text):
        lst_widget = parent.get_widget()
        if lst_widget:
            o_text_b = lst_widget[-1]
            if isinstance(o_text_b, widgets.Text) and o_text == o_text_b:  # 属性相等，则合并
                lst_widget[-1].text += o_text.text
            else:
                parent.append(o_text)
        else:
            parent.append(o_text)

    def add_style(self, element, dict_style):
        if element.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            dict_style["font-size"] = element.tag
            # dict_style["font-weight"] = "bold"

        elif element.tag in ["strong"]:
            dict_style["font-weight"] = "bold"

        elif element.tag in ["mark"]:
            dict_style["font-weight"] = "bold"
            dict_style["background"] = "#FFFF00"

        elif element.tag in ["code"]:
            dict_style["background"] = "#EEF1F6"
            dict_style["padding"] = "15px 16px"
            # dict_style["padding"] = "15px 16px"

    def _parse_element(self, root, parent):
        for element in root.getchildren():
            o_widget = None

            dict_style = self._parse_style(element)
            self.lst_style.append(dict_style)
            if element.tag in ["span"]:
                pass

            elif element.tag in ["div"]:
                if dict_style == {"height": "-1px"}:
                    continue

                o_widget = self.create_panel(dict_style)
                o_widget.tag = element.tag

            elif element.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                self.add_style(element, dict_style)
                o_widget = self.create_panel(dict_style)
                o_widget.tag = element.tag

            elif element.tag in ["strong"]:
                self.add_style(element, dict_style)

            elif element.tag in ["mark"]:
                self.add_style(element, dict_style)

            elif element.tag in ["br"]:
                parent.append(widgets.Row(height=3))

            elif element.tag in ["code"]:
                self.add_style(element, dict_style)
                o_widget = self.create_panel(dict_style)
                o_widget.tag = element.tag

            elif element.tag in ["p", "div"]:
                o_widget = self.create_panel(dict_style)
                o_widget.tag = element.tag

            elif element.tag in ["img"]:
                o_widget = self.create_image(dict_style)

            elif element.tag in ["table"]:
                o_widget = self.create_table(element, dict_style)
                parent.append(o_widget)
                continue

            if element.text:
                text = element.text
                text = text.replace(" \xa0", "\xa0").replace("\t", "   ")
                if o_widget:
                    o_text = self.create_text(text, self.make_style(), o_widget)
                    o_widget.append(o_text)
                else:
                    o_text = self.create_text(text, self.make_style(), parent)
                    self.add_text(parent, o_text)
            else:
                if o_widget:
                    o_spacing = self.create_col_spacing(self.make_style(), o_widget)
                    if o_spacing:
                        o_widget.append(o_spacing)

                else:
                    o_spacing = self.create_col_spacing(self.make_style(), parent)
                    if o_spacing:
                        parent.append(o_spacing)

            if o_widget:
                self._parse_element(element, o_widget)
            else:
                self._parse_element(element, parent)

            if o_widget:
                parent.append(o_widget)

            del self.lst_style[-1]

            if element.tail:
                text = element.tail
                text = text.replace(" \xa0", "\xa0").replace("\t", "   ")
                o_text = self.create_text(text, self.make_style(), parent)
                self.add_text(parent, o_text)

    def parse_element(self, root):
        for (i, element) in enumerate(root):
            dict_style = self._parse_style(element)
            if element.tag in ["div"]:
                if dict_style == {"height": "-1px"}:  # 人工加的要去掉
                    continue

            self.lst_style.append(dict_style)

            self.add_style(element, dict_style)
            o_panel = self.create_panel(dict_style)
            o_panel.tag = element.tag

            if element.text:
                text = element.text
                o_text = self.create_text(text, self.make_style(), o_panel)
                o_panel.append(o_text)

            self._parse_element(element, o_panel)

            if element.tail:
                del self.lst_style[-1]
                text = element.tail
                o_text = self.create_text(text, self.make_style(), o_panel)
                self.add_text(o_panel, o_text)

            self.lst_panel.append(o_panel)

            self.lst_style = []

    def to_vadmin(self):
        html = etree.HTML(self.content)
        tree = ET.fromstring(etree.tostring(html))  # 使用etrr格式化html
        root = tree.find("body")  # 去掉html, 从body开始
        self.parse_element(root)
        return self.lst_panel

    def to_doc(self, path):
        if not self.lst_panel:
            self.to_vadmin()
        Vadmin2Word(self.lst_panel).to_word(path)
