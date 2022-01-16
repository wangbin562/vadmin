# !/usr/bin/python
# -*- coding=utf-8 -*-

"""
v1.3
1、不支持内嵌的图表
2、不支持*.doc文件，只支持*.docx文件
3、不支持分页显示
4、不支持两端对齐
"""

import codecs
import copy
import importlib
import math
import os
import re
import time

import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Twips
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

try:
    import xml.etree.cElementTree as ET
except (BaseException,):
    import xml.etree.ElementTree as ET

from xpinyin import Pinyin
from vadmin import widgets
from vadmin import step
from utils.word_html_api import UnitApi
from utils.word_html_api import FontApi
from utils.word_html_api import CssStyleApi
from utils.word_html_api import DocApi

BEGIN_KEY = "<%"
END_KEY = "%>"
MIN_HEIGHT = 16
RE_PATTERN = "([a-zA-Z0-9.])"
RE_LOWERCASE = "emsp;+|[0-9a-zA-Z().,`~!@#$%^&*()_+-='/?<>\\|\{}\[\]\" ]+"


# RATIO = 0.2


class BasePage(object):
    def __init__(self, output):
        self.output = output
        self.lst_widget = []
        self.o_parent = None
        self.dict_param = None

    def create(self, **kwargs):
        self.init_para(**kwargs)
        self.write()
        return self.lst_widget

    def set_param(self, o_parent=None, **kwargs):
        self.o_parent = o_parent
        self.dict_param = kwargs

    def _write(self, content, **kwargs):
        kwargs["text"] = content
        param = dict()
        if self.dict_param and kwargs:
            param = copy.copy(self.dict_param)
            param.update(kwargs)
        elif self.dict_param:
            param.update(self.dict_param)
        elif kwargs:
            param.update(kwargs)

        self.write_text(self.o_parent, **param)

    def write_text(self, o_parent=None, **kwargs):
        if "text_size" not in kwargs:
            kwargs["text_size"] = 16

        if "text_horizontal" not in kwargs:
            kwargs["text_horizontal"] = "left"

        o_text = widgets.Text(**kwargs)
        if o_parent is not None:
            o_parent.append(o_text)
        else:
            self.lst_widget.append(o_text)

    def write_panel(self, o_parent=None, **kwargs):
        if "min_height" not in kwargs:
            kwargs["min_height"] = MIN_HEIGHT

        o_panel = widgets.Panel(**kwargs)
        if o_parent is not None:
            o_parent.append(o_panel)
        else:
            self.lst_widget.append(o_panel)

        return o_panel

    def write_layout_table(self, o_parent=None, **kwargs):
        o_table = widgets.LayoutTable(**kwargs)
        if o_parent is not None:
            o_parent.append(o_table)
        else:
            self.lst_widget.append(o_table)

        return o_table

    def write_row_spacing(self):
        self.lst_widget.append(widgets.Row(-1))

    def write_widget(self, o_parent=None, **kwargs):
        o_widget = widgets.param_2_widget(**kwargs)
        if isinstance(o_widget, (widgets.Input, widgets.InputNumber)):
            # o_widget.border_config = [False, False, True, False]
            # self.o_parent = o_parent
            self._write(o_widget.value)

        if o_parent is not None:
            o_parent.append(o_widget)
        elif self.o_parent is not None:
            self.o_parent.append(o_widget)
        else:
            self.lst_widget.append(o_widget)


class Compiler(object):
    def __init__(self, path):
        self.lst_import = []
        self.lst_global = []
        self.lst_code = []
        self.path = path
        self.doc = docx.Document(path)
        # self.tree = ET.ElementTree()
        # self.root = ET.Element("body")
        self.lst_paragraph = []
        # self.dict_run = {}
        self.page_height = None
        self.page_width = None
        self.content_width = None
        self.script = None
        # self.runs = None  # 已解析过的run
        # self.prev_run = None  # 上一段
        # self.next_run = None  # 下一段
        # self.idx = None
        self.paragraph = None
        self.paragraph_text_height = None
        self.number_idx = None

    def compile_page(self, **kwargs):
        for key, val in kwargs.items():
            self.lst_global.append(key)

        self.parse()
        return self.generate_code()

    def parse_paragraph(self, paragraph, width=None, cell=None):
        self.paragraph = paragraph
        # run_len = len(paragraph.runs)
        runs = paragraph.runs
        # if run_len == 0:
        #     runs = []
        #     for run in paragraph._p.xpath("./w:pPr"):
        #         runs.append(Run(run, paragraph))
        # else:
        #     runs = paragraph.runs
        run_len = len(runs)

        if run_len == 0 and cell:
            return None

        o_panel = widgets.Panel(width=width, vertical="bottom")
        if paragraph.paragraph_format.line_spacing is None:
            min_height = UnitApi.twips_2_px(Twips(480))
        elif paragraph.paragraph_format.line_spacing in [1.0]:
            min_height = UnitApi.twips_2_px(Twips(240))
        else:
            if paragraph.paragraph_format.line_spacing_rule in \
                    [WD_LINE_SPACING.AT_LEAST, WD_LINE_SPACING.EXACTLY]:
                min_height = UnitApi.twips_2_px(paragraph.paragraph_format.line_spacing)
                # o_panel.vertical = "bottom"
                # o_panel.height = min_height
            else:  # WD_LINE_SPACING.AT_LEAST
                line_spacing = paragraph.paragraph_format.line_spacing
                min_height = UnitApi.twips_2_px(line_spacing * Twips(240))

        if paragraph.paragraph_format.space_after:
            min_height += UnitApi.twips_2_px(paragraph.paragraph_format.space_after)

        if paragraph.paragraph_format.space_before:
            min_height += UnitApi.twips_2_px(paragraph.paragraph_format.space_before)

        if width is None:
            width = self.content_width
        o_panel.width = width

        # else:
        #     o_panel = widgets.Panel(width=self.content_width)
        # o_panel = widgets.Panel(width=571)
        if run_len == 0:
            lst_sz = paragraph._p.xpath('./w:pPr/w:rPr/w:sz')
            if lst_sz:
                sz = lst_sz[0].val
                min_height += UnitApi.twips_2_px(sz)

        if min_height != MIN_HEIGHT:
            o_panel.min_height = min_height

        self.paragraph_text_height = min_height

        # 对齐
        if paragraph.alignment == WD_TAB_ALIGNMENT.CENTER:
            o_panel.horizontal = "center"
        elif paragraph.alignment == WD_TAB_ALIGNMENT.RIGHT:
            o_panel.horizontal = "right"
        # elif paragraph.alignment == WD_TAB_ALIGNMENT.LEFT:
        # else:
        #     o_panel.horizontal = "left"

        if cell:
            if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                # o_panel.vertical = "center"
                pass
            elif cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
                o_panel.vertical = "bottom"
            else:
                o_panel.vertical = "top"

        if paragraph.paragraph_format.first_line_indent:  # 第一行左缩进
            spacing = UnitApi.twips_2_px(paragraph.paragraph_format.first_line_indent)
            o_panel.append(widgets.ColSpacing(spacing))

        if paragraph.paragraph_format.left_indent:  # 块全部左缩进
            indent_width = int(UnitApi.twips_2_px(paragraph.paragraph_format.left_indent))
            # o_spacing = widgets.ColSpacing(indent_width)
            # o_spacing.tag = "left_indent"
            # o_panel.append(o_spacing)
            o_panel.margin_left = indent_width
            o_widget = o_panel
            # col_width = {0: indent_width, 1: width - indent_width}
            # o_panel.width = col_width[1]
            # data = [[None, o_panel]]
            #
            # o_table = widgets.LayoutTable(data=data, focus_color="transparent", width=width,
            #                               col_width=col_width, auto_wrap=True)
            # o_widget = o_table

        elif paragraph.paragraph_format.right_indent:
            indent_width = int(UnitApi.twips_2_px(paragraph.paragraph_format.right_indent))
            col_width = {0: width - indent_width, 1: indent_width}
            o_panel.width = col_width[0]
            o_widget = o_panel
            # data = [[o_panel, None]]
            #
            # o_table = widgets.LayoutTable(data=data, focus_color="transparent", width=width,
            #                               col_width=col_width, auto_wrap=True, border_color="#FF0000")
            # o_widget = o_table
        else:
            o_widget = o_panel

        self.parse_number(paragraph, o_panel)
        for i, run in enumerate(runs):
            if run.text:
                hyperlink, hyperlink_run = self.get_hyperlink(run, paragraph)
                if hyperlink:
                    if hyperlink_run:
                        self.parse_run(hyperlink_run, o_panel, hyperlink=hyperlink)
                        self.parse_run(run, o_panel)

                    else:
                        self.parse_run(run, o_panel, hyperlink=hyperlink)
                else:
                    self.parse_run(run, o_panel)
            else:
                self.parse_image(run, o_panel)

        self.format_paragraph(o_panel)
        return o_widget

    def format_paragraph(self, o_panel):
        width = 0.0
        lst_width = copy.copy(o_panel.get_widget())
        i = 0
        for o_widget in lst_width:
            if isinstance(o_widget, widgets.Text) and o_widget.text:
                text = o_widget.text.replace("&ensp;", " ").replace("&thinsp;", "")
                lst_text = re.findall(RE_LOWERCASE, text, re.S)
                if lst_text:
                    text = "".join(lst_text)
                    tmp_width = width + len(text) * o_widget.text_size / 2.0
                    width = 0.0
                    if tmp_width > o_panel.width:
                        o_panel.insert_widget(i, widgets.Row(height=0))
                        i += 1
                else:
                    count = text.count(" ")
                    text_len = len(text)
                    width = width + (text_len - count) * o_widget.text_size + count * o_widget.text_size / 2.0

                    if o_widget.padding_left:
                        width += o_widget.padding_left

                    if o_widget.padding_right:
                        width += o_widget.padding_right

            elif isinstance(o_widget, widgets.Image):
                width += o_widget.width
                if o_widget.padding_left:
                    width += o_widget.padding_left

                if o_widget.padding_right:
                    width += o_widget.padding_right

            elif isinstance(o_widget, widgets.ColSpacing):
                # if o_widget.tag == "left_indent":
                width += o_widget.spacing

            i += 1

        return o_panel

    def parse_script(self, run, script, o_parent):
        dict_style = {"script": script}
        font = run.font
        if font.size:
            dict_style["text_size"] = UnitApi.twips_2_px(font.size)
        # else:
        #     dict_style["text_size"] = 16  # 默认小四

        if font.color.rgb:
            dict_style["text_color"] = "#%s" % str(font.color.rgb)

        if font.highlight_color:
            dict_style["background_color"] = StyleApi.color_idx_2_rgb(font.highlight_color)

        if font.underline:
            dict_style["under_line"] = font.underline

        if font.strike:
            dict_style["del_line"] = font.strike

        if run.style.font.bold:
            dict_style["bold"] = run.style.font.bold

        if font.italic:
            dict_style["italics"] = font.italic

        o_parent.append(dict_style)
        self.script = None

    def get_image_id(self, run):
        namespace = {
            'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
            'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

        if hasattr(run.element, "drawing_lst"):
            for drawing in run.element.drawing_lst:
                for node in drawing.findall('.//a:blip', namespace):
                    for k, v in node.items():
                        if k.find("}embed") >= 0:
                            return v

    def get_image_size(self, rId):
        for shape in self.doc.inline_shapes:
            if rId == shape._inline.graphic.graphicData.pic.blipFill.blip.embed:
                if shape.width:
                    return UnitApi.twips_2_px(shape.width), UnitApi.twips_2_px(shape.height)
                return None, None

        return None, None

    def get_row_height(self, row):
        if row.height:
            return UnitApi.twips_2_px(row.height)

        cell = row.cells[0]
        tc = cell._tc
        # size = tc.xpath('./w:trPr/w:trHeight')
        lst_sz = tc.xpath('./w:p/w:pPr/w:rPr/w:sz')
        if lst_sz:
            sz = lst_sz[0]
            return UnitApi.twips_2_px(sz.val)

    def get_font_family(self, run):
        r = run._r
        lst_font = r.xpath('./w:rPr/w:rFonts')
        if lst_font:
            font = lst_font[0]
            return font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or font.get(qn("w:cs"))

    def get_hyperlink(self, run, paragraph):
        """
        链接有两种格式
        1、在paragraph的run中存在，但没有链接内容，通过style_id中的name_val判断是否为链接文本
        2、在paragraph中不存在，在paragraph中查找是否有链接结点，判断是否在run之前
        """
        if not run.style:
            return None, None

        for style in self.doc.styles.element.style_lst:
            if run.style.style_id == style.styleId:
                if style.name_val == "Hyperlink":
                    p = paragraph._p
                    lst_instr = p.xpath('./w:r/w:instrText')
                    if lst_instr:
                        instr = lst_instr[0]
                        o_re = re.search('HYPERLINK "(.*)" ', instr.text, re.I)
                        if o_re:
                            hyperlink = o_re.groups()[0]
                        else:
                            hyperlink = instr.text

                        return hyperlink, None
                break

        lst_hyperlink = paragraph._p.xpath("./w:hyperlink")
        if lst_hyperlink:
            for hyperlink in lst_hyperlink:
                if hyperlink.getnext() == run.element:
                    rId = hyperlink.get(qn("r:id"))
                    hyperlink_run = None
                    for node in lst_hyperlink[0].getchildren():
                        # hyperlink_text += node.text
                        hyperlink_run = Run(node, paragraph)
                        break

                    for rel in self.doc.part.rels:
                        if rel == rId and self.doc.part.rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                            return self.doc.part.rels[rel]._target, hyperlink_run

        return None, None

    def get_table(self, paragraph, cell):
        for table in cell.tables:
            if paragraph._p.getprevious() == table._element:
                return table

    def parse_image(self, run, o_parent):
        rId = self.get_image_id(run)
        if not rId:
            return
        image_part = self.doc.part.related_parts[rId]
        name = image_part.partname

        try:
            from django.conf import settings
            path = os.path.abspath(os.path.join(settings.MEDIA_ROOT, "images", name.strip("/\\")))
            url = os.path.join(settings.MEDIA_URL, "images", name.strip("/\\"))
        except (BaseException,):
            path = os.path.join(os.getcwd(), "images", name.strip("/\\"))
            url = path

        tmp_path = os.path.split(path)[0]
        if not os.path.exists(tmp_path):
            os.makedirs(tmp_path)

        image_data = image_part._blob
        with open(path, 'wb') as fp:
            fp.write(image_data)

        width, height = self.get_image_size(rId)
        if width:
            o_parent.append(widgets.Image(url=url, width=width, height=height, padding_right=3))
        else:
            o_parent.append(widgets.Image(url=url, padding_right=3))

    def parse_number(self, paragraph, o_parent):
        """解析编号"""
        lst_num = paragraph._p.xpath('./w:pPr/w:numPr/w:numId')
        if lst_num:
            num_id = lst_num[0].val
            ilvl = paragraph._p.xpath('./w:pPr/w:numPr/w:ilvl')[0].val
            text = ""
            if self.number_idx is None:
                self.number_idx = 1
            else:
                self.number_idx += 1

            if num_id in [1, 2]:
                dict_number = {'0': '零', '1': '一', '2': '二', '3': '三', '4': '四',
                               '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
                s_number = str(self.number_idx)
                lst_number = []
                for num in s_number:
                    lst_number.append(dict_number[num])
                text = "%s、" % "".join(lst_number)

            elif num_id == 3:
                text = "%s. " % self.number_idx
            else:
                return

            o_text = widgets.Text(text=text, inline=True, margin_left=43)

            # 字体
            lst_sz = paragraph._p.xpath('./w:pPr/w:rPr/w:sz')

            if lst_sz:
                o_text.text_size = UnitApi.twips_2_px(lst_sz[0].val)
            else:
                o_text.text_size = 14  # 默认五号

            lst_font = paragraph._p.xpath('./w:pPr/w:rPr/w:rFonts')
            if lst_font:
                font = lst_font[0]
                font_family = font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or font.get(qn("w:cs"))
                o_text.font_family = FontApi.doc_2_html(font_family, o_text.text)

            # if font.color.rgb:
            #     o_text.text_color = "#%s" % str(font.color.rgb)

            # if font.highlight_color:
            #     o_text.background_color = StyleApi.color_idx_2_rgb(font.highlight_color)

            # if font.underline:
            #     o_text.under_line = font.underline

            # if font.strike:
            #     o_text.del_line = font.strike

            # if run.style.font.bold:
            #     o_text.bold = run.style.font.bold
            #
            # if font.italic:
            #     o_text.italics = font.italic

            o_parent.append(o_text)

        else:
            self.number_idx = None

    def parse_run(self, run, o_parent, text=None, hyperlink=None):
        if text is None:
            text = run.text

        begin_idx = text.find(BEGIN_KEY)
        end_idx = text.find(END_KEY)
        if self.script is not None:
            if end_idx < 0:
                self.script += text
            else:
                self.script += text[:end_idx]
                self.parse_script(run, script=self.script, o_parent=o_parent)
                text = text[end_idx + 2:]
                if text:
                    self.parse_run(run, text=text, o_parent=o_parent)
            return

        if (begin_idx >= 0) and (end_idx >= 0):  # 在同一段
            if begin_idx > 0:
                self.parse_run(run, text=text[:begin_idx], o_parent=o_parent)

            self.parse_script(run, script=text[begin_idx + 2: end_idx], o_parent=o_parent)
            if (end_idx + 2) < len(text):
                self.parse_run(run, text=text[end_idx + 2:], o_parent=o_parent)
            return

        elif begin_idx >= 0:
            self.script = text[begin_idx + 2:]
            text = text[:begin_idx]
            if text:
                self.parse_run(run, text=text, o_parent=o_parent)

            return

        if hyperlink:
            try:
                import shutil
                from django.conf import settings
                path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, "hyperlink"))
                if not os.path.exists(path):
                    os.makedirs(path)
                file_path = os.path.join(path, hyperlink)
                shutil.copy(os.path.join(os.path.split(self.path)[0], hyperlink), file_path)
                hyperlink = "%shyperlink/%s" % (settings.MEDIA_URL, hyperlink.strip("/\\"))
            except (BaseException,):
                pass

            o_text = widgets.Text(text=text, inline=True, under_line=True,
                                  text_color="#0563C1", step=step.DownloadFile(url=hyperlink))
        else:
            o_text = widgets.Text(text=text, inline=True)

        # 字体
        font = run.font
        if font.size:
            o_text.text_size = UnitApi.twips_2_px(font.size)
        else:
            o_text.text_size = 14  # 默认五号

        # 在非表格内，word字体上下都会空了一点（测量值），如下特殊处理
        if self.paragraph.paragraph_format.line_spacing is None:
            o_text.height = o_text.text_size + 21
        else:
            o_text.height = o_text.text_size + 8

        o_text.font_family = FontApi.doc_2_html(self.get_font_family(run), o_text.text)

        if font.color.rgb:
            o_text.text_color = "#%s" % str(font.color.rgb)
        else:
            o_text.text_color = "#000000"

        if font.highlight_color:
            o_text.background_color = CssStyleApi.color_idx_2_rgb(font.highlight_color)

        if font.underline:
            o_text.under_line = font.underline

        if font.strike:
            o_text.del_line = font.strike

        if run.style.font.bold:
            o_text.bold = run.style.font.bold

        if font.italic:
            o_text.italics = font.italic

        if o_text.text:
            o_text.text = "&thinsp;".join(re.findall("emsp;+|[-=（）()0-9a-zA-Z().]+|[^-=（）()0-9a-zA-Z().]+",
                                                     o_text.text, re.S))
            if re.findall(RE_LOWERCASE, o_text.text, re.S):
                o_text.text = o_text.text.replace("\xa0", "&ensp;").replace(" ", "&ensp;") \
                    .replace("\u3000", "&ensp;")
            else:
                o_text.text = o_text.text.replace("\xa0", "&ensp;&ensp;").replace(" ", "&ensp;") \
                    .replace("\u3000", "&ensp;&ensp;")

            lst_widget = o_parent.get_widget()
            if lst_widget and isinstance(lst_widget[-1], widgets.Text):
                o_widget = lst_widget[-1]
                s = o_widget.text[-1]
                o_re_1 = re.search(RE_PATTERN, s, re.S)
                o_re_2 = re.search(RE_PATTERN, o_text.text[0], re.S)
                if o_re_1 is None and o_re_2:
                    o_widget.padding_right = 8
                    o_parent.append(o_text)
                else:
                    if o_widget == o_text:
                        text_1 = o_text.text.replace("&ensp;", "")
                        text_2 = o_widget.text.replace("&ensp;", "")
                        if text_1 == "":  # 等于空要合并
                            o_widget.text += o_text.text

                        # 同为汉字或同为字母要合并
                        elif bool(re.findall(RE_LOWERCASE, text_1, re.S)) == \
                                bool(re.findall(RE_LOWERCASE, text_2, re.S)):
                            o_widget.text += o_text.text
                        else:
                            o_parent.append(o_text)
                    else:
                        o_parent.append(o_text)
            else:
                o_parent.append(o_text)

    def get_merged_cell(self, table, cell, row_idx, col_idx):
        row_num = len(table.rows)
        col_num = len(table.columns)
        end_row = row_idx
        end_col = col_idx
        for r in range(row_idx, row_num):
            for c in range(col_idx, col_num):
                if table.cell(r, c)._tc == cell._tc:
                    end_row = r
                    end_col = c

        return end_row, end_col

    def get_merged_info(self, merged_cells, row_idx, col_idx):
        """获取合并的数据"""
        b_merged = False
        b_range = False
        row_span = 0
        col_span = 0
        row_col = '%s-%s' % (row_idx, col_idx)
        for s_range in merged_cells:
            s_range = str(s_range)
            begin_row_col, end_row_col = s_range.split(":")
            # begin_col, begin_row = re.findall("[A-Z]+|[0-9]+", begin_col_row, re.I)
            # end_col, end_row = re.findall("[A-Z]+|[0-9]+", end_col_row, re.I)
            begin_row, begin_col = begin_row_col.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)
            end_row, end_col = end_row_col.split("-")
            end_row = int(end_row)
            end_col = int(end_col)

            if begin_row_col == row_col:  # 开始合并
                b_merged = True
                row_span = end_row - begin_row + 1
                col_span = end_col - begin_col + 1
                break

            elif (row_idx >= begin_row) and (row_idx <= end_row) and \
                    (col_idx >= begin_col) and (col_idx <= end_col):
                b_range = True
                break

        return b_merged, b_range, row_span, col_span

    def get_merged_cells(self, table):
        """获取合并单元格"""
        row_num = len(table.rows)
        col_num = len(table.columns)

        merged_cells = []
        for row_idx in range(row_num):
            for col_idx in range(col_num):
                b_merged, b_range, row_span, col_span = self.get_merged_info(merged_cells, row_idx, col_idx)
                if b_merged or b_range:
                    continue

                cell = table.cell(row_idx, col_idx)
                end_row, end_col = self.get_merged_cell(table, cell, row_idx, col_idx)
                if (row_idx != end_row) or (col_idx != end_col):
                    # merged_cells.append("%s%s:%s%s" % (get_column_letter(col_idx + 1), row_idx + 1,
                    #                                    get_column_letter(c + 1), r + 1))
                    merged_cells.append("%s-%s:%s-%s" % (row_idx, col_idx, end_row, end_col))

        return merged_cells

    def parse_table(self, table):
        data = list()
        min_row_height = dict()
        merged_cells = self.get_merged_cells(table)
        for row_idx, row in enumerate(table.rows):
            min_row_height[row_idx] = self.get_row_height(row) or 20
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                b_merged, b_range, row_span, col_span = self.get_merged_info(merged_cells, row_idx, col_idx)
                if b_range:
                    row_data.append(None)
                    continue

                # 获取表格文字对齐（水平、垂直）
                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                    vertical = "center"
                elif cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
                    vertical = "bottom"
                else:
                    vertical = "top"

                horizontal = DocApi.get_cell_horizontal(cell)

                width = 0.0
                if b_merged and col_span:
                    for i in range(col_idx, col_idx + col_span):
                        width += UnitApi.twips_2_px(table.columns[i].width)
                else:
                    width = UnitApi.twips_2_px(table.columns[col_idx].width)
                width = int(width)

                background_color = DocApi.get_cell_background_color(cell)
                o_panel_cell = widgets.Panel(vertical=vertical, horizontal=horizontal,
                                             background_color=background_color, height=min_row_height[row_idx])
                paragraphs_len = len(cell.paragraphs)
                for paragraph in cell.paragraphs:
                    o_table_sub = self.get_table(paragraph, cell)
                    if o_table_sub:
                        o_panel = self.parse_table(o_table_sub)
                        if o_panel:
                            #     o_panel.padding_left = 3
                            #     o_panel.padding_right = 3
                            #     # o_panel_cell.append(widgets.ColSpacing(3))
                            o_panel_cell.append(o_panel)
                            # o_panel_cell.append(widgets.Row(3))

                        o_panel = self.parse_paragraph(paragraph, width, cell=cell)
                        if o_panel:
                            o_panel.padding_left = 3
                            o_panel_cell.append(o_panel)

                    else:
                        o_panel = None
                        if not paragraph.runs:
                            """
                            表格中有空行的情况，docx不解析到runs中
                            要解析xml获取，获取占位的高度
                            """
                            lst_sz = paragraph._p.xpath("./w:pPr/w:rPr/w:sz")
                            if lst_sz:
                                size = UnitApi.twips_2_px(lst_sz[0].val)
                                if paragraph.paragraph_format.line_spacing is None:
                                    height = size + 21
                                else:
                                    height = size + 8
                                o_panel = widgets.Panel(min_height=height)
                        else:
                            o_panel = self.parse_paragraph(paragraph, width, cell=cell)

                        if o_panel:
                            # o_panel.padding_left = 3
                            # o_panel.border_color = "#FF0000"
                            # o_panel.vertical = vertical
                            # o_panel.horizontal = horizontal
                            if paragraphs_len == 1:
                                # if vertical != "top":  # 有中，下对齐的情况下，要有固定高度
                                o_panel.height = min_row_height[row_idx]
                                o_panel.background_color = background_color
                                o_panel_cell = o_panel
                            else:
                                o_panel_cell.append(o_panel)
                                o_panel_cell.append(widgets.RowSpacing())

                        elif paragraphs_len == 1:
                            o_panel_cell.height = min_row_height[row_idx]

                row_data.append(o_panel_cell)
            data.append(row_data)

        col_width = dict()
        for col_idx, col in enumerate(table.columns):
            col_width[col_idx] = UnitApi.twips_2_px(col.width)

        o_panel = widgets.Panel(width=math.ceil(sum(col_width.values())))
        o_table = widgets.LayoutTable(data=data, col_border=True, row_border=True, focus_color="transparent",
                                      col_width=col_width, auto_wrap=True, border_color="#000000",
                                      min_row_height=min_row_height, merged_cells=merged_cells)
        o_panel.add_widget(o_table)
        return o_panel

    def parse(self):
        """
        解释文件
        """
        section = self.doc.sections[0]
        # section.page_height = Mm(297)
        # section.page_width = Mm(210)
        # section.left_margin = Mm(25.4)
        # section.right_margin = Mm(25.4)
        # section.top_margin = Mm(25.4)
        # section.bottom_margin = Mm(25.4)
        # section.header_distance = Mm(12.7)
        # section.footer_distance = Mm(12.7)
        self.page_width = math.ceil(UnitApi.twips_2_px(section.page_width))
        self.content_width = math.ceil(
            UnitApi.twips_2_px(section.page_width - section.left_margin - section.right_margin))
        # self.content_width += 17  # 不知道什么原因，差有17个px距离
        height = UnitApi.twips_2_px(section.top_margin)
        o_panel = widgets.Panel(height=height, width=self.content_width)
        self.lst_paragraph.append(o_panel)

        for child in self.doc._body._body:
            if isinstance(child, CT_P):
                # paragraph._p.xpath("./w:pPr/w:rPr/w:szCs")[0].get(qn("w:val")) 获取高度
                paragraph = Paragraph(child, self.doc)
                if paragraph._p.xpath('./w:pPr/w:adjustRightInd'):  # 头部空行，经验值
                    continue

                o_panel = self.parse_paragraph(paragraph)
                if o_panel is not None:
                    self.lst_paragraph.append(o_panel)

            elif isinstance(child, CT_Tbl):
                table = Table(child, self.doc)
                o_panel = self.parse_table(table)
                self.lst_paragraph.append(o_panel)
            else:
                # print(child)
                pass

        height = UnitApi.twips_2_px(section.bottom_margin)
        o_panel = widgets.Panel(height=height, width=self.content_width)
        self.lst_paragraph.append(o_panel)

    def generate_run_code(self, run, parent=None):
        if isinstance(run, widgets.Panel):
            if parent:
                parent_name = "o_panel_%s" % parent.replace(".", "_").replace("[", "_").replace("]", "")
                self.lst_code.append('        %s = self.write_panel(o_parent=%s, **%s)' %
                                     (parent_name, parent, run.render_self()))
            else:
                parent_name = "o_panel"
                self.lst_code.append('        o_panel = self.write_panel(**%s)' % run.render_self())

            for o_sub in run.lst_widget:
                self.generate_run_code(o_sub, parent=parent_name)

        elif isinstance(run, widgets.LayoutTable):
            if parent:
                table_name = "o_table_%s" % parent.replace(".", "_").replace("[", "_").replace("]", "")
            else:
                table_name = "o_table"

            self.lst_code.append('        %s = self.write_layout_table(o_parent=%s, **%s)' %
                                 (table_name, parent, run.render_self()))

            for row, lst_row in enumerate(run.data):
                self.lst_code.append('')
                if parent:
                    row_name = "row_%s" % parent.replace(".", "_").replace("[", "_").replace("]", "")
                else:
                    row_name = "row"

                self.lst_code.append('        %s = []' % row_name)
                for col, cell in enumerate(lst_row):
                    if (cell is None) and (col == 0):
                        self.lst_code.append('        self.write_text(o_parent=%s, **{"widget_type": "text", '
                                             '"text": "&nbsp;", "inline": 1})' % row_name)
                    elif cell is None:
                        self.lst_code.append('        %s.append(None)' % row_name)
                    else:
                        self.generate_run_code(cell, parent=row_name)

                self.lst_code.append('        %s.data.append(%s)' % (table_name, row_name))

        elif isinstance(run, widgets.Text):
            if parent:
                self.lst_code.append('        self.write_text(o_parent=%s, **%s)' % (parent, run.render()))
            else:
                self.lst_code.append('        self.write_text(**%s)' % run.render())

        elif isinstance(run, widgets.Widget):
            if parent:
                self.lst_code.append('        self.write_widget(o_parent=%s, **%s)' % (parent, run.render()))
            else:
                self.lst_code.append('        self.write_widget(**%s)' % run.render())

        elif isinstance(run, dict):  # 脚本
            script = run["script"]
            del run["script"]
            if parent:
                parent_name = parent.replace(".", "_").replace("[", "_").replace("]", "")
                self.lst_code.append('        self.set_param(o_parent=%s, **%s)' % (parent_name, run))
            else:
                self.lst_code.append('        self.set_param(**%s)' % run)
            self.lst_code.append('        %s' % script.strip().replace("\n", "\n        "))

        elif isinstance(run, str):
            self.lst_code.append('        %s' % run)
        else:
            self.lst_code.append('        pass')

    def generate_code(self):
        """
        生成代码
        """
        self.lst_code.append('# !/usr/bin/python')
        self.lst_code.append('# -*- coding: utf-8 -*-')
        self.lst_code.append('"""')
        self.lst_code.append('code generated by word template compiler')
        self.lst_code.append('"""')
        # lst_code.append('import copy')
        self.lst_code.append('from utils.word_2_html import BasePage')
        self.lst_code.append('')

        if self.lst_import:
            for s_import in self.lst_import:
                self.lst_code.append(s_import)

        if self.lst_global:
            for s_global in self.lst_global:
                self.lst_code.append("%s = None" % s_global)

        self.lst_code.append('''

class Page(BasePage):
    def __init__(self, output=None):
        super().__init__(output)
        global _write_widget
        global _w
        _write_widget = self.write_widget
        _w = self._write
''')

        self.lst_code.append('    def init_para(self, **kwargs):')
        self.lst_code.append('        self.page_width = %s' % self.page_width)
        self.lst_code.append('        self.content_width = %s' % self.content_width)
        for s_global in self.lst_global:
            self.lst_code.append('        global %s' % s_global)
            self.lst_code.append('        %s = kwargs["%s"]' % (s_global, s_global))
        self.lst_code.append('')

        self.lst_code.append('    def write(self):')
        for i, paragraph in enumerate(self.lst_paragraph):
            self.lst_code.append('        self.write_%s()' % (i + 1))
            self.lst_code.append('        self.write_row_spacing()')
        self.lst_code.append('')

        for i, paragraph in enumerate(self.lst_paragraph):
            self.lst_code.append('    def write_%s(self):' % (i + 1))
            self.generate_run_code(paragraph)
            self.lst_code.append('')

        self.lst_code.append('')
        return "\n".join(self.lst_code)


def generate_page(template_file, is_compile=False, **kwargs):
    path, name = os.path.split(template_file)
    module_name = "word_2_html_%s" % os.path.splitext(name)[0]
    p = Pinyin()
    module_name_py = p.get_pinyin(module_name, '').lower()

    if not os.path.exists("templates_py"):
        os.makedirs("templates_py")
        open(os.path.join("templates_py", "__init__.py"), "w").write("# !/usr/bin/python")

    file_path = os.path.join("templates_py", "%s.py" % module_name_py)
    template_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(template_file)))
    template_py_time = None
    if os.path.exists(file_path):
        template_py_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(file_path)))

    if (not os.path.exists(file_path)) or is_compile or (template_py_time < template_time):
        compiler = Compiler(template_file)
        fh = codecs.open(file_path, "w", encoding='utf-8')
        result = compiler.compile_page(**kwargs)
        fh.write(result)
        fh.close()
        module = importlib.import_module("templates_py.%s" % module_name_py)
        importlib.reload(module)
    else:
        module = importlib.import_module("templates_py.%s" % module_name_py)
    o_page = module.Page()
    return o_page


def generate(template_file, is_compile=False, **kwargs):
    """
    template_file：模板文件(只支持*.docx)
    is_compile:强制编译（模板转换成代码）
    """
    o_page = generate_page(template_file, is_compile, **kwargs)
    lst_widget = o_page.create(**kwargs)

    lst_output = []
    for o_widget in lst_widget:
        lst_output.append(o_widget.render_html())
    return '<div style="width:%s;margin:0 auto;">%s</div>' % (o_page.page_width, "\r".join(lst_output))


if __name__ == "__main__":
    final_price = 123
    begin_year = 2019
    begin_month = 9
    begin_day = 10
    project_name = "AAAAAAAAAA"
    provider_name = "AAAAAAAAAA"
    output = generate(r"/Volumes/share/wangbin/code_sync/project_management/templates/acceptance.docx",
                      is_compile=True, final_price=final_price, begin_year=begin_year, begin_month=begin_month,
                      begin_day=begin_day, project_name=project_name, provider_name=provider_name, aa=2, bb=3)

    fh = open("2.html", "w", encoding='utf-8')
    fh.write(output)
    fh.close()
