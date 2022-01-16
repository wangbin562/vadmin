# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
word操作api, vadmin数据转换word使用，主要是text、image、lite_table组件
"""
import re

import requests
import unicodedata
from PIL import ImageFont

try:
    from docx import shared
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.run import Run
except (BaseException,):
    pass
from utils import conver

try:
    import xml.etree.cElementTree as ET
except (BaseException,):
    import xml.etree.ElementTree as ET

DEFAULT_TABLE_STYLE = "TABLE STYLE"
CON_DEFAULT_FONT = "微软雅黑"
CON_DEFAULT_FONT_EN = "SansSerif"
CON_DEFAULT_FONT_SIZE = 16

DICT_COLOR_NAME_2_HEX = {
    "black": "#000000",
    "navy": "#000080",
    "darkblue": "#00008b",
    "mediumblue": "#0000cd",
    "blue": "#0000ff",
    "darkgreen": "#006400",
    "green": "#008000",
    "teal": "#008080",
    "darkcyan": "#008b8b",
    "deepskyblue": "#00bfff",
    "darkturquoise": "#00ced1",
    "mediumspringgreen": "#00fa9a",
    "lime": "#00ff00",
    "springgreen": "#00ff7f",
    "aqua": "#00ffff",
    "cyan": "#00ffff",
    "midnightblue": "#191970",
    "dodgerblue": "#1e90ff",
    "lightseagreen": "#20b2aa",
    "forestgreen": "#228b22",
    "seagreen": "#2e8b57",
    "darkslategray": "#2f4f4f",
    "limegreen": "#32cd32",
    "mediumseagreen": "#3cb371",
    "turquoise": "#40e0d0",
    "royalblue": "#4169e1",
    "steelblue": "#4682b4",
    "darkslateblue": "#483d8b",
    "mediumturquoise": "#48d1cc",
    "indigo": "#4b0082",
    "darkolivegreen": "#556b2f",
    "cadetblue": "#5f9ea0",
    "cornflowerblue": "#6495ed",
    "mediumaquamarine": "#66cdaa",
    "dimgray": "#696969",
    "slateblue": "#6a5acd",
    "olivedrab": "#6b8e23",
    "slategray": "#708090",
    "lightslategray": "#778899",
    "mediumslateblue": "#7b68ee",
    "lawngreen": "#7cfc00",
    "chartreuse": "#7fff00",
    "aquamarine": "#7fffd4",
    "maroon": "#800000",
    "purple": "#800080",
    "olive": "#808000",
    "gray": "#808080",
    "skyblue": "#87ceeb",
    "lightskyblue": "#87cefa",
    "blueviolet": "#8a2be2",
    "darkred": "#8b0000",
    "darkmagenta": "#8b008b",
    "saddlebrown": "#8b4513",
    "darkseagreen": "#8fbc8f",
    "lightgreen": "#90ee90",
    "mediumpurple": "#9370db",
    "darkviolet": "#9400d3",
    "palegreen": "#98fb98",
    "darkorchid": "#9932cc",
    "yellowgreen": "#9acd32",
    "sienna": "#a0522d",
    "brown": "#a52a2a",
    "darkgray": "#a9a9a9",
    "lightblue": "#add8e6",
    "greenyellow": "#adff2f",
    "paleturquoise": "#afeeee",
    "lightsteelblue": "#b0c4de",
    "powderblue": "#b0e0e6",
    "firebrick": "#b22222",
    "darkgoldenrod": "#b8860b",
    "mediumorchid": "#ba55d3",
    "rosybrown": "#bc8f8f",
    "darkkhaki": "#bdb76b",
    "silver": "#c0c0c0",
    "mediumvioletred": "#c71585",
    "indianred": "#cd5c5c",
    "peru": "#cd853f",
    "chocolate": "#d2691e",
    "tan": "#d2b48c",
    "lightgray": "#d3d3d3",
    "thistle": "#d8bfd8",
    "orchid": "#da70d6",
    "goldenrod": "#daa520",
    "palevioletred": "#db7093",
    "crimson": "#dc143c",
    "gainsboro": "#dcdcdc",
    "plum": "#dda0dd",
    "burlywood": "#deb887",
    "lightcyan": "#e0ffff",
    "lavender": "#e6e6fa",
    "darksalmon": "#e9967a",
    "violet": "#ee82ee",
    "palegoldenrod": "#eee8aa",
    "lightcoral": "#f08080",
    "khaki": "#f0e68c",
    "aliceblue": "#f0f8ff",
    "honeydew": "#f0fff0",
    "azure": "#f0ffff",
    "sandybrown": "#f4a460",
    "wheat": "#f5deb3",
    "beige": "#f5f5dc",
    "whitesmoke": "#f5f5f5",
    "mintcream": "#f5fffa",
    "ghostwhite": "#f8f8ff",
    "salmon": "#fa8072",
    "antiquewhite": "#faebd7",
    "linen": "#faf0e6",
    "lightgoldenrodyellow": "#fafad2",
    "oldlace": "#fdf5e6",
    "red": "#ff0000",
    "fuchsia": "#ff00ff",
    "magenta": "#ff00ff",
    "deeppink": "#ff1493",
    "orangered": "#ff4500",
    "tomato": "#ff6347",
    "hotpink": "#ff69b4",
    "coral": "#ff7f50",
    "darkorange": "#ff8c00",
    "lightsalmon": "#ffa07a",
    "orange": "#ffa500",
    "lightpink": "#ffb6c1",
    "pink": "#ffc0cb",
    "gold": "#ffd700",
    "peachpuff": "#ffdab9",
    "navajowhite": "#ffdead",
    "moccasin": "#ffe4b5",
    "bisque": "#ffe4c4",
    "mistyrose": "#ffe4e1",
    "blanchedalmond": "#ffebcd",
    "papayawhip": "#ffefd5",
    "lavenderblush": "#fff0f5",
    "seashell": "#fff5ee",
    "cornsilk": "#fff8dc",
    "lemonchiffon": "#fffacd",
    "floralwhite": "#fffaf0",
    "snow": "#fffafa",
    "yellow": "#ffff00",
    "lightyellow": "#ffffe0",
    "ivory": "#fffff0",
    "white": "#ffffff",
}

CSS_2_DOC_EN = {
    u"verdana": u"Verdana",
    u"arial": u"Arial",
    u"helvetica": u"Helvetica",
    u"sans-serif": u"SansSerif",
    u"consolas": u"Consolas",
    u"simsun": u"SimSun-ExtB",
    u"宋体": u"SimSun-ExtB",
    u"微软雅黑": u"Microsoft YaHei UI",
    u"simkai": u"楷体",
    u"楷体": u"楷体",
    u"楷体_gb2312": u"楷体",
    u"黑体": u"黑体",
    u"simhei": u"SimHei",
    u"隶书": u"隶书",
    u"simli": u"隶书",
    u"andale mono": u"SansSerif",
    u"arial black": u"Arial Black",
    u"avant garde": u"Arial Black",
    u"comic sans ms": u"Comic Sans MS",
    u"impact": u"Impact",
    u"chicago": u"Impact",
    u"times new roman": u"Times New Roman",
    u"courier new": u"Courier New",
    u"wingdings": u"Wingdings",
    u"华文细黑": u"华文细黑",
}

CSS_2_DOC_CHS = {
    u"verdana": u"微软雅黑",
    u"arial": u"微软雅黑",
    u"helvetica": u"微软雅黑",
    u"sans-serif": u"微软雅黑",
    u"consolas": u"微软雅黑",
    u"simsun": u"宋体",
    u"宋体": u"宋体",
    u"微软雅黑": u"微软雅黑",
    u"simkai": u"楷体",
    u"楷体": u"楷体",
    u"楷体_gb2312": u"楷体",
    u"黑体": u"黑体",
    u"隶书": u"隶书",
    u"simhei": u"隶书",
    u"andale mono": u"微软雅黑",
    u"arial black": u"微软雅黑",
    u"avant garde": u"微软雅黑",
    u"comic sans ms": u"微软雅黑",
    u"impact": u"微软雅黑",
    u"chicago": u"微软雅黑",
    u"times new roman": u"微软雅黑",
    u"courier new": u"微软雅黑",
    u"wingdings": u"Wingdings",
    u"华文细黑": u"华文细黑",
}

default_size_height = {
    56: 78,
    55: 77,
    54: 75,
    53: 74,
    52: 73,
    51: 71,
    50: 70,
    49: 69,
    48: 67,
    47: 66,
    46: 65,
    45: 63,
    44: 62,
    43: 61,
    42: 59,
    41: 57,
    40: 56,
    39: 54,
    38: 53,
    37: 52,
    36: 50,
    35: 49,
    34: 48,
    33: 46,
    32: 45,
    31: 44,
    30: 42,
    29: 41,
    28: 40,
    27: 38,
    26: 37,
    25: 36,
    24: 33,
    23: 32,
    22: 30,
    21: 29,
    20: 28,
    19: 26,
    18: 25,
    17: 24,
    16: 22,
    15: 21,
    14: 20,
    13: 18,
    12: 17,
}


def delete_paragraph(paragraph):
    """
    删除段落
    Args:
        paragraph:

    Returns:

    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class ParagraphApi(object):
    @staticmethod
    def create(parent):
        paragrap = parent.add_paragraph()
        paragrap.paragraph_format.line_spacing = 1.0
        return paragrap

    @staticmethod
    def delete(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def set_horizontal(paragraph, horizontal):
        if horizontal == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif horizontal == "right":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    @staticmethod
    def set_space(paragraph, space_after=0, space_before=0):
        paragraph.paragraph_format.space_after = conver.px_2_twips(space_after)
        paragraph.paragraph_format.space_before = conver.px_2_twips(space_before)

    @staticmethod
    def set_indent(paragraph, left_indent=0, right_indent=0):
        paragraph.paragraph_format.left_indent = conver.px_2_twips(left_indent)
        paragraph.paragraph_format.right_indent = conver.px_2_twips(right_indent)


class TextApi(object):

    @staticmethod
    def insert(paragraph, run, text):
        new_run_element = paragraph._element._new_r()
        run._element.addnext(new_run_element)
        new_run = Run(new_run_element, run._parent)
        new_run.text = text
        return new_run

    @staticmethod
    def split(text):
        """
        将整体文字分割成段落（根据中英文，在word上中英文显示的宽度不一致）
        文字分割(中英文分割)
        isalnum() 方法检测字符串是否由字母和数字组成。
        isalpha() 方法检测字符串是否只由字母组成
        isdigit() 方法检测字符串是否只由数字组成。
        isnumeric() 方法检测字符串是否只由数字组成。这种方法是只针对unicode对象。
        isspace() 方法检测字符串是否只由空格组成。
        istitle() 方法检测字符串中所有的单词拼写首字母是否为大写，且其他字母为小写。
        isupper() 方法检测字符串中所有的字母是否都为大写。
        """
        lst_split = []
        lst_en = []
        lst_chs = []
        for s in text:
            ascii_val = ord(s)
            if (ascii_val > 31) and (ascii_val < 128):
                lst_en.append(s)
                if lst_chs:
                    lst_split.append(("".join(lst_chs), False))
                    lst_chs = []
            else:
                lst_chs.append(s)
                if lst_en:
                    lst_split.append(("".join(lst_en), True))
                    lst_en = []

        if lst_chs:
            lst_split.append(("".join(lst_chs), False))

        if lst_en:
            lst_split.append(("".join(lst_en), True))

        return lst_split

    @staticmethod
    def get_default_family(name, is_en=False):
        """
        字体转换（根据中英文分别转换）
        """
        name = name.replace("'", "").replace('"', "").lower()
        for single_name in name.split(","):
            single_name = single_name.strip()
            if single_name == "":
                continue
            if is_en:
                if single_name in CSS_2_DOC_EN:
                    return CSS_2_DOC_EN[single_name]
            else:
                if single_name in CSS_2_DOC_CHS:
                    return CSS_2_DOC_CHS[single_name]

        if is_en:
            return CON_DEFAULT_FONT_EN

        return CON_DEFAULT_FONT

    @staticmethod
    def set_style(run, dict_text):
        """
        设置文字样式
        """
        # 获取字体
        font = dict_text.get("font", {})
        if "family" in font:
            font_name = font["family"]
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)  # 中文加这句设置才会生效

        # 文字大小
        if "size" in font:
            run.font.size = conver.px_2_twips(font["size"])

        # 文字颜色
        color = font.get("color", None)
        if color is not None:
            r, g, b = conver.color_2_rgb(color)
            run.font.color.rgb = shared.RGBColor(r, g, b)

        # 文字划线
        decoration = font.get("decoration", None)
        if decoration is not None:
            # 下划线
            if decoration == "underline":
                run.font.underline = True

            # 删除线
            elif decoration == "line-through":
                run.font.strike = True

            # 上划线
            elif decoration == "overline":
                pass

        # 粗体
        weight = font.get("weight", None)
        if weight == "bold":
            run.font.bold = True

        # 斜体
        style = font.get("style", None)
        if style in ["italic", "oblique"]:
            run.font.italic = True

        # 文字背景颜色
        bg = dict_text.get("bg", {})
        bg_color = bg.get("color", None)
        if bg_color is not None:
            run.font.highlight_color = conver.color_2_background_color(color)


class ImageApi(object):
    @staticmethod
    def download(href, path):
        """
        下载图片
        :return:
        """
        # name = os.path.split(href)[1]
        # path = os.path.abspath(os.path.join(CON_BASE_ADDRESS, name))
        # if os.path.exists(path):
        #     return path

        ir = requests.get(href)
        if ir.status_code == 200:
            open(path, 'wb').write(ir.content)
            return path

        return None

    @staticmethod
    def add(run, path, width, height, float=None, paragraph=None):
        if float is not None:
            if "top" in float and "left" in float:
                from utils.template import add_float_picture
                add_float_picture.add_float_picture(paragraph, path, width=width,
                                                    height=height, pos_x=float["left"], pos_y=float["top"])
            else:
                run.add_picture(path, width=width, height=height)
                node = run._element.xpath("w:drawing/wp:inline")[0]
                extent_extent = node.find(qn("wp:extent"))
                extent_docPr = node.find(qn("wp:docPr"))
                extent_cNvGraphicFramePr = node.find(qn("wp:cNvGraphicFramePr"))
                extent_graphic = node.find(qn("a:graphic"))

                new_node = OxmlElement("wp:anchor")
                new_node.set('distT', "0")
                new_node.set('distB', "0")
                new_node.set('distL', "114300")
                new_node.set('distR', "114300")
                new_node.set('simplePos', "0")
                new_node.set('relativeHeight', "251658240")
                new_node.set('behindDoc', "1")
                new_node.set('locked', "0")
                new_node.set('layoutInCell', "1")
                new_node.set('allowOverlap', "1")

                node_simplePos = OxmlElement("wp:simplePos")
                node_simplePos.set('x', "0")
                node_simplePos.set('y', "0")

                node_positionH = OxmlElement("wp:positionH")
                node_positionH.set('relativeFrom', 'character')
                node_posOffset = OxmlElement("wp:align")
                node_posOffset.text = "center"
                node_positionH.append(node_posOffset)

                node_positionV = OxmlElement("wp:positionV")
                node_positionV.set('relativeFrom', 'line')
                node_posOffset = OxmlElement("wp:align")
                node_posOffset.text = "center"
                node_positionV.append(node_posOffset)

                new_node.append(node_simplePos)
                new_node.append(node_positionH)
                new_node.append(node_positionV)
                new_node.append(extent_extent)
                new_node.append(extent_docPr)
                new_node.append(extent_cNvGraphicFramePr)
                new_node.append(extent_graphic)

                node.getparent().replace(node, new_node)

                pass

        else:
            run.add_picture(path, width=width, height=height)


class TableApi(object):

    @staticmethod
    def move(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    @staticmethod
    def delete():
        pass

    @staticmethod
    def add_style(doc, style_name=DEFAULT_TABLE_STYLE, dict_border=None):
        """
        <w:style w:type="table" w:styleId="a1">
    <w:name w:val="Table Grid"/>
  # <w:basedOn w:val="TableNormal"/>
  # <w:uiPriority w:val="59"/>
  # <w:rsid w:val="00FC693F"/>
  # <w:pPr> # 段落使用
  #  <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
  # </w:pPr>
  <w:tblPr>
    <w:tblInd w:w="0" w:type="dxa"/> # 指定要在表格的前边缘（从左到右的表格的左边缘）之前添加的缩进 dxa-指定该值的单位为点的二十分之一（1/1440英寸）
    <w:tblBorders>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>
    <w:tblCellMar>
      <w:top w:w="0" w:type="dxa"/>
      <w:left w:w="108" w:type="dxa"/>
      <w:bottom w:w="0" w:type="dxa"/>
      <w:right w:w="108" w:type="dxa"/>
    </w:tblCellMar>
  </w:tblPr>
</w:style>
        """
        if doc.styles._element.get_by_name(style_name):
            return style_name

        style = OxmlElement('w:style')
        style.set(qn('w:type'), "table")
        style.set(qn('w:styleId'), "a%s" % (len(doc.styles) + 1))
        name = OxmlElement('w:name')
        name.set(qn('w:val'), style_name)
        style.append(name)

        tblPr = OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        if dict_border is None:
            dict_border = {
                "top": {"w:val": "single", "w:color": "auto"},
                "left": {"w:val": "single", "w:color": "auto"},
                "bottom": {"w:val": "single", "w:color": "auto"},
                "right": {"w:val": "single", "w:color": "auto"},
                "insideH": {"w:val": "single", "w:color": "auto"},
                "insideV": {"w:val": "single", "w:color": "auto"},
            }
        for k, v in dict_border.items():
            element = OxmlElement('w:%s' % k)
            element.set(qn('w:sz'), "4")
            element.set(qn('w:space'), "0")
            for k1, v1 in v.items():
                element.set(qn(k1), v1)
            tblBorders.append(element)
        tblPr.append(tblBorders)

        tblCellMar = OxmlElement('w:tblCellMar')
        for k, v in [["top", "0"], ["left", "108"], ["bottom", "0"], ["right", "108"]]:
            element = OxmlElement('w:%s' % k)
            element.set(qn('w:val'), v)
            element.set(qn('w:type'), "dxa")
            tblCellMar.append(element)
        tblPr.append(tblCellMar)
        style.append(tblPr)
        doc.styles._element.append(style)
        return style_name

    @staticmethod
    def set_border(doc, dict_table):
        dict_border = {
            "top": {"w:val": "none", "w:color": "auto"},
            "left": {"w:val": "none", "w:color": "auto"},
            "bottom": {"w:val": "none", "w:color": "auto"},
            "right": {"w:val": "none", "w:color": "auto"},
            "insideH": {"w:val": "none", "w:color": "auto"},
            "insideV": {"w:val": "none", "w:color": "auto"},
        }

        lst_color = []
        if dict_table.get("border", None):
            border = dict_table.get("border")
            color = border.get("color")
            if color:
                lst_color.append(color)
                dict_border["top"]["w:color"] = color[1:]
                dict_border["top"]["w:val"] = "single"
                dict_border["left"]["w:color"] = color[1:]
                dict_border["left"]["w:val"] = "single"
                dict_border["bottom"]["w:color"] = color[1:]
                dict_border["bottom"]["w:val"] = "single"
                dict_border["right"]["w:color"] = color[1:]
                dict_border["right"]["w:val"] = "single"

        if dict_table.get("col_border", None):
            border = dict_table.get("col_border")
            color = border.get("color")
            if color:
                lst_color.append(color)
                dict_border["insideV"]["w:color"] = color[1:]
                dict_border["insideV"]["w:val"] = "single"

        if dict_table.get("row_border", None):
            border = dict_table.get("row_border")
            color = border.get("color")
            if color:
                lst_color.append(color)
                dict_border["insideH"]["w:color"] = color[1:]
                dict_border["insideH"]["w:val"] = "single"

        style_name = "TABLE_STYLE_NOT_BORDER"
        if lst_color:
            style_name = "TABLE_STYLE-%s" % "-".join(lst_color)

        TableApi.add_style(doc, style_name, dict_border)
        return style_name

    @staticmethod
    def set_row_height(table, dict_table, rule="exact"):
        """
        设置行高
        """
        row_height = dict_table.get("row_height", {})
        if row_height:
            for row_idx, row_data in enumerate(dict_table.get("data", [])):
                row = table.rows[row_idx]
                height_pt = conver.px_2_pt(row_height)
                # height_twips = shared.Pt(height_pt)
                tr = row._tr
                tr_pr = tr.get_or_add_trPr()
                tr_height = OxmlElement('w:trHeight')
                tr_height.set(qn('w:val'), "%s" % int(height_pt * 20))  # val - 指定行的高度，在二十分之一的点。
                tr_height.set(qn('w:hRule'), "atLeast")
                tr_pr.append(tr_height)

    @staticmethod
    def set_col_width(table, dict_table, content_width):
        """
        设置列宽，由于office和wps要求不一样，所有cell和col都设置
        """
        for col_idx, width in dict_table.get("col_width", {}).items():
            col_idx = int(col_idx)
            width_twips = conver.px_2_twips(conver.calc_2_px(width, content_width))
            # 设置列宽
            for row in table.rows:
                cell = row.cells[col_idx]
                cell.width = width_twips

            col = table.columns[col_idx]
            col.width = width_twips

    @staticmethod
    def set_left_indent(table, left_indent):
        """设置表格X轴位置"""
        # left_indent = conver.px_2_pt(left_indent - self.page["left_margin"])
        # element = OxmlElement('w:tblInd')
        # element.set(qn('w:w'), "%s" % int(round(left_indent * 20.0)))
        # element.set(qn('w:type'), "dxa")
        # table._tblPr.append(element)
        pass

    @staticmethod
    def set_padding(table, dict_table):
        """
        设置表格padding
        单位pt
        """
        #     < w:tblCellMar >
        #     < w:top        w:w = "57"        w:type = "dxa" / >
        #     < w:left        w:w = "85"        w:type = "dxa" / >
        #     < w:bottom        w:w = "113"        w:type = "dxa" / >
        #     < w:right        w:w = "142"        w:type = "dxa" / >
        # < / w:tblCellMar >

        top = dict_table.get("space_top", 0)
        bottom = dict_table.get("space_bottom", 0)
        left = dict_table.get("space_left", 0)
        right = dict_table.get("space_right", 0)
        if top is None:
            top_20_pt = 0
        else:
            top_20_pt = int(round(top * 20))
        element = OxmlElement('w:tblCellMar')
        element_sub = OxmlElement('w:top')
        element_sub.set(qn('w:w'), "%s" % top_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if bottom is None:
            bottom_20_pt = 0
        else:
            bottom_20_pt = int(round(bottom * 20))
        element_sub = OxmlElement('w:bottom')
        element_sub.set(qn('w:w'), "%s" % bottom_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if left is None:
            left_20_pt = 0
        else:
            left_20_pt = int(round(left * 20))
        element_sub = OxmlElement('w:left')
        element_sub.set(qn('w:w'), "%s" % left_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if right is None:
            right_20_pt = 0
        else:
            right_20_pt = int(round(right * 20))
        element_sub = OxmlElement('w:right')
        element_sub.set(qn('w:w'), "%s" % right_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)
        table._tblPr.append(element)

    @staticmethod
    def set_vertical(table, dict_table):
        """
        设置垂直对齐
        :param table:
        :param dict_table:
        :return:
        """
        vertical = dict_table.get("vertical", None)
        row_vertical = dict_table.get("row_vertical", {})
        data = dict_table.get("data", [])
        if vertical in ["middle", "bottom"]:  # middle和bottom，默认是居上，不用调用此函数设置
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    TableApi.set_cell_vertical(cell, vertical)

        for idx, v in row_vertical.items():
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_data in enumerate(row_data):
                    if row_idx != idx:
                        continue

                    cell = table.cell(row_idx, col_idx)
                    TableApi.set_cell_vertical(cell, v)

    @staticmethod
    def set_horizontal(table, dict_table):
        """
        设置水平对齐
        :param table:
        :param dict_table:
        :return:
        """
        horizontal = dict_table.get("horizontal", None)
        col_horizontal = dict_table.get("col_horizontal", {})
        data = dict_table.get("data", [])
        if horizontal in ["center", "right"]:
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    TableApi.set_cell_horizontal(cell, horizontal)

        for idx, h in col_horizontal.items():
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx != idx:
                        continue

                    cell = table.cell(row_idx, col_idx)
                    TableApi.set_cell_horizontal(cell, h)

    @staticmethod
    def set_cell_vertical(cell, vertical):
        if vertical == "middle":
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif vertical == "bottom":
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

    @staticmethod
    def set_cell_horizontal(cell, horizontal):
        if horizontal in [None, "left"]:
            return

        lst_p = cell._tc.xpath("./w:p")
        if lst_p:
            p = lst_p[0]
            element = OxmlElement('w:pPr')
            p.append(element)
            sub = OxmlElement('w:jc')
            sub.set(qn('w:val'), "%s" % horizontal)
            element.append(sub)

    @staticmethod
    def merge(table, dict_table):
        """
        合并单元格
        :param table:
        """
        for merged_cell in dict_table.get("merged_cells", []):
            begin, end = merged_cell.split(":")
            begin_row, begin_col = begin.split("-")
            end_row, end_col = end.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)
            end_row = int(end_row)
            end_col = int(end_col)
            cell = table.cell(begin_row, begin_col)
            merge_cell = table.cell(end_row, end_col)
            cell.merge(merge_cell)

    @staticmethod
    def get_merged_cells(table):
        """
        获取合并单元格
        :param table:word table对象
        """
        row_num = len(table.rows)
        col_num = len(table.columns)

        merged_cells = []
        for row_idx in range(row_num):
            for col_idx in range(col_num):
                b_merged, b_range, row_span, col_span = \
                    TableApi.get_merged_info(merged_cells, row_idx, col_idx)

                if b_merged or b_range:
                    continue

                cell = table.cell(row_idx, col_idx)
                end_row, end_col = TableApi.get_merged_cell(table, cell, row_idx, col_idx)
                if (row_idx != end_row) or (col_idx != end_col):
                    # merged_cells.append("%s%s:%s%s" % (get_column_letter(col_idx + 1), row_idx + 1,
                    #                                    get_column_letter(c + 1), r + 1))
                    merged_cells.append("%s-%s:%s-%s" % (row_idx, col_idx, end_row, end_col))

        return merged_cells

    @staticmethod
    def get_merged_cell(table, cell, row_idx, col_idx):
        row_num = len(table.rows)
        col_num = len(table.columns)

        end_row = row_idx
        end_col = col_idx

        for r in range(row_idx + 1, row_num):
            if table.cell(r, col_idx)._tc == cell._tc:
                end_row = r
            else:
                break

        for c in range(col_idx + 1, col_num):
            if table.cell(row_idx, c)._tc == cell._tc:
                end_col = c
            else:
                break

        return end_row, end_col

    @staticmethod
    def get_merged_info(merged_cells, row_idx, col_idx):
        """获取合并的数据"""
        b_merged = False
        b_range = False
        row_span = 0
        col_span = 0
        row_col = '%s-%s' % (row_idx, col_idx)

        for s_range in merged_cells:
            s_range = str(s_range)
            begin_row_col, end_row_col = s_range.split(":")
            # begin_col, begin_row = re.findall("[A-Z]+|[0-9]+", begin_col_row, re.I)
            # end_col, end_row = re.findall("[A-Z]+|[0-9]+", end_col_row, re.I)
            begin_row, begin_col = begin_row_col.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)
            end_row, end_col = end_row_col.split("-")
            end_row = int(end_row)
            end_col = int(end_col)

            if begin_row_col == row_col:  # 开始合并
                b_merged = True
                row_span = end_row - begin_row + 1
                col_span = end_col - begin_col + 1
                break

            elif (row_idx >= begin_row) and (row_idx <= end_row) and \
                    (col_idx >= begin_col) and (col_idx <= end_col):
                b_range = True
                break

        return b_merged, b_range, row_span, col_span


# def get_half_width(font_size):
#     """获取半角字符的显示宽度"""
#     return font_size / 2.0 * 0.74
#

def get_half_char(width, font_size, is_pdf=False):
    """根据宽度获取可以显示的半角字符个数"""
    w = get_text_width('.', font_size, is_pdf=is_pdf)
    num = int(width / w)
    return num


def get_chinese_char(text):
    """获取中文字符，含全角符号"""
    pass


def replace_chinese_char(text):
    """替换中文字符，含全角符号"""
    # https://www.jb51.net/article/86577.htm
    return re.sub(u'[\u4e00-\u9fa5，。？]+', "", text)


def get_paragraph_space_pos(paragraph, width=0):
    font_size = paragraph.get("font", {}).get("size", 16)
    page_width = paragraph.width
    for widget in paragraph.children:
        text = widget.text.strip(" ")
        text_width = get_text_width(widget.text, widget.get("font", {}).get("size") or font_size)
        width += text_width
        if (text != "") and (width > page_width):
            width -= page_width

    return width


def get_text_width(text, font_size, font_family="utils/fonts/Arial.ttf", is_pdf=False):
    """
    ea ; A     ; Ambiguous  不确定
    ea ; F     ; Fullwidth  全宽
    ea ; H     ; Halfwidth  半宽
    ea ; N     ; Neutral   中性
    ea ; Na    ; Narrow    窄
    ea ; W     ; Wide     宽
    :param text:
    :param font_size:
    :param font_family:
    :param is_pdf:
    :return:
    """
    lst_char_w = []
    lst_char = []
    lst_char_space = []
    for char in text:
        if char == " ":
            lst_char_space.append(char)
        elif unicodedata.east_asian_width(char) in ('F', 'W', 'A'):
            lst_char_w.append(char)
        else:
            lst_char.append(char)

    if is_pdf:
        return len(lst_char) * font_size / 2.0 + len(lst_char_w) * font_size + len(lst_char_space) * font_size / 2.0

    font = ImageFont.truetype(font_family, round(font_size), 0)
    width, height = font.getsize("".join(lst_char))
    return width + len(lst_char_w) * font_size + len(lst_char_space) * font_size / 2.0


def get_paragraph_line_height(line_spacing_rule, line_spacing, font_size, font_family="utils/fonts/Arial.ttf"):
    # // ONE_POINT_FIVE = 1，1.5倍行距
    # // AT_LEAST = 3，最小行距
    # // DOUBLE = 2，双倍行距
    # // EXACTLY = 4，固定值
    # // MULTIPLE = 5，多倍行距
    # // SINGL = 0，单倍行距
    font = ImageFont.truetype(font_family, font_size, 0)
    width, font_height = font.getsize(" ")
    font_height += 2
    if line_spacing_rule == 1:
        height = font_height * 1.5
    elif line_spacing_rule == 2:
        height = font_height * 2
    elif line_spacing_rule == 3:
        height = font_height
    elif line_spacing_rule == 4:
        height = line_spacing
    elif line_spacing_rule == 5:
        height = font_height * line_spacing
    else:
        height = font_height * 1.5

    return height
