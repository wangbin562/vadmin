# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
"""
from PIL import Image

from utils import conver
from utils import word_api
from vadmin import common

MIN_HEIGHT = 16
DEFAULT_TABLE_STYLE = "TABLE STYLE"
DEFAULT_PARAGRAPH_STYLE = "PARAGRAPH STYLE"


class Compiler(object):
    def __init__(self, node, template_path=None, **kwargs):
        import docx
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        if isinstance(node, (tuple, list)):
            self.lst_node = node
        else:
            self.lst_node = [node, ]

        if template_path:
            try:
                self.doc = docx.Document(template_path)  # 使用模板样式
                for child in self.doc._body._body:
                    if isinstance(child, CT_P):
                        word_api.ParagraphApi.delete(child)
                    elif isinstance(child, CT_Tbl):
                        word_api.ParagraphApi.delete(child)

                # for paragraph in self.doc.paragraphs:
                #     self.delete_paragraph(paragraph._element)
            except (BaseException,):
                self.doc = docx.Document()
        else:
            self.doc = docx.Document()

        self.paragraph = None
        self.parent = self.doc
        self.parent_node = {}
        section = self.doc.sections[0]
        self.width = conver.twips_2_px(section.page_width)
        section.left_margin = kwargs.get("left_margin", section.left_margin)
        section.right_margin = kwargs.get("left_margin", section.right_margin)
        section.top_margin = kwargs.get("top_margin", section.top_margin)
        section.bottom_margin = kwargs.get("bottom_margin", section.bottom_margin)
        self.content_width = conver.twips_2_px(section.page_width - section.left_margin -
                                               section.right_margin)

        # section.page_width = shared.Pt(conver.px_2_pt(page["page_width"] or 794))
        # section.page_height = shared.Pt(conver.px_2_pt(page["page_height"] or 1123))

    def write(self, lst_node):
        for i, node in enumerate(lst_node):
            if not node:
                continue

            widget_type = node.get("type", None)
            if widget_type == "text":
                # paragraph = self.write_paragraph(self.doc, node)
                self.write_text(node)

            elif widget_type == "image":
                self.write_image(node)

            elif widget_type == "row":
                self.paragraph = word_api.ParagraphApi.create(self.parent)
                word_api.ParagraphApi.set_horizontal(self.paragraph, self.parent_node.get("horizontal", None))
                # if node.get("height", None):
                self.paragraph.paragraph_format.line_spacing = conver.twips_2_px(node.get("height", 0))

            elif widget_type == "panel":
                self.paragraph = word_api.ParagraphApi.create(self.parent)
                word_api.ParagraphApi.set_horizontal(self.paragraph, node.get("horizontal", None))
                self.parent_node = node
                self.write(node.get("children", []))

            elif widget_type == "lite_table":
                self.write_table(node)

    def write_text(self, dict_widget):
        if self.paragraph is None:
            self.paragraph = word_api.ParagraphApi.create(self.parent)

        run = self.paragraph.add_run()
        text = dict_widget.get("text", None) or ""

        run.text = text.replace("\r\n", "\n")
        word_api.TextApi.set_style(run, dict_widget)
        new_run = run

        # for run_text, is_en in word_api.TextApi.split(text):
        #     # 插入text
        #     new_run = word_api.TextApi.insert(self.paragraph, run, run_text)
        #     # run._element.addprevious(new_run_element)
        #     # self.run = self.paragraph.add_run(run_text)
        #     word_api.TextApi.set_style(new_run, dict_widget, is_en)

        if "image" in dict_widget.get("bg", {}):
            bg = dict_widget.get("bg", {})
            self.write_image({"type": "image", "href": bg.get("image"),
                              "float": {}, "width": dict_widget.get("width", None),
                              "height": dict_widget.get("height", None)}, new_run)

    def write_image(self, dict_image, run=None):
        if run is None:
            run = self.paragraph.add_run()

        href = dict_image.get("href", None) or dict_image.get("active_href", None)
        if href:
            path = common.get_file_path(href)
            width = dict_image.get("width", None)
            height = dict_image.get("height", None)

            if width is not None:
                width = conver.px_2_twips(width)

            if height is not None:
                height = conver.px_2_twips(height)

            if (width is None) or (height is None):
                im = Image.open(path)
                if width is None:
                    if im.size[0] > self.content_width:
                        width = conver.px_2_twips(self.content_width)
                    else:
                        width = conver.px_2_twips(im.size[0])

            word_api.ImageApi.add(run=run, path=path, width=width, height=height,
                                  float=dict_image.get("float", None))

    def write_table(self, dict_table):
        data = dict_table.get("data", [])
        if data:
            row_num = len(data)
            col_num = len(data[0])
        else:
            row_num = 0
            col_num = 0

        table = self.parent.add_table(row_num, col_num)
        # if self.parent.paragraphs:
        #     self.paragraph = self.parent.paragraphs[-1]
        # self.paragraph._p.addnext(table._element)

        style_name = word_api.TableApi.set_border(self.doc, dict_table)
        if not style_name:
            style_name = word_api.TableApi.add_style(self.doc)
        table.style = style_name

        # 设置行高
        word_api.TableApi.set_row_height(table, dict_table)

        # 设置列宽
        word_api.TableApi.set_col_width(table, dict_table, self.content_width)

        # word_api.TableApi.set_left_indent(table, dict_table)
        word_api.TableApi.set_horizontal(table, dict_table)
        word_api.TableApi.set_vertical(table, dict_table)
        # 合并单元格
        word_api.TableApi.merge(table, dict_table)
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        word_api.TableApi.set_padding(table, dict_table)

        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                if not isinstance(cell_data, list):
                    cell_data = [cell_data]

                self.parent = cell
                self.paragraph = cell.paragraphs[-1]
                self.paragraph.paragraph_format.line_spacing = 1.0

                # 规避第一个子组件是表格，有一个空行的情况
                if cell.paragraphs and cell_data and cell_data[0] and \
                        cell_data[0].get("type", None) in ["lite_table"]:
                    word_api.ParagraphApi.delete(cell.paragraphs[-1])

                self.write(cell_data)

    def to_doc(self, path):
        self.doc.save(path)


def generate(path, lst_node, template_path=None, **kwargs):
    """
    path:生成的文件路径
    lst_node:节点数据
    template_file：模板文件(只支持*.docx)
    """
    obj = Compiler(lst_node, template_path, **kwargs)
    obj.write(obj.lst_node)
    obj.to_doc(path)


if __name__ == "__main__":
    pass
