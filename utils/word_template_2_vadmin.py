# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
"""
import json
import os
import re

import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import RGBColor
from docx.shared import Twips
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from utils import conver
from utils import word_api
from vadmin import const
from vadmin import step
from vadmin import widgets

MIN_HEIGHT = 16


class Compiler(object):
    def __init__(self, path, param=None, show_width=None, is_margin=None, only_content=False):
        self.path = path
        self.doc = docx.Document(path)
        if param is None:
            self.param = {}
        else:
            self.param = param
        self.lst_paragraph = []
        self.width = None
        self.height = None
        self.show_width = show_width  # 显示的宽度
        self.is_margin = is_margin  # 是否有边框
        self.only_content = only_content  # 是否只构造内容区域
        self.ratio = 1
        self.col_width_ratio = 1

        self.page_width = None
        self.page_height = None
        self.content_width = None
        self.content_height = None
        self.top_margin = 0
        self.right_margin = 0
        self.bottom_margin = 0
        self.left_margin = 0

        self.margin = []
        self.lst_widget = []
        self.prev_text_run = None
        self.dict_abstract = {}
        self.paragraph = None
        self.paragraph_number = {}

    def parse(self):
        """
        解释文件
        """
        section = self.doc.sections[0]
        self.width = conver.twips_2_px(section.page_width)
        if (self.show_width is not None) and (self.show_width < self.width):
            self.ratio = self.show_width / self.width  # 显示宽度小于实际宽度，所有内容要同比缩小
            if self.is_margin is None:
                self.is_margin = False

        elif self.is_margin is None:
            self.is_margin = True
            # self.width = self.show_width

        self.height = conver.twips_2_px(section.page_height)

        # 文档网络类型（只支持无网络和只指定行网络）
        self.line_type = section._sectPr.find(qn("w:docGrid")).get(qn("w:type"))
        line_pitch = Twips(int(section._sectPr.find(qn("w:docGrid")).get(qn("w:linePitch"))))
        self.line_pitch = conver.twips_2_px(line_pitch)

        self.page_width = conver.twips_2_px(section.page_width)
        self.page_height = self.height
        self.content_width = conver.twips_2_px(section.page_width - section.left_margin -
                                               section.right_margin)
        self.content_height = conver.twips_2_px(section.page_height - section.top_margin -
                                                section.bottom_margin)
        try:
            self.parse_paragraph_number()
        except (BaseException,):
            pass

        self.top_margin = conver.twips_2_px(section.top_margin)
        self.right_margin = conver.twips_2_px(section.right_margin)
        self.bottom_margin = conver.twips_2_px(section.bottom_margin)
        self.left_margin = conver.twips_2_px(section.left_margin)

        line_height = round(30 * self.ratio)
        if self.only_content:
            margin_top = 0
            margin_right = 0
            margin_bottom = 0
            margin_left = 0

        elif self.is_margin:
            self.col_width_ratio = self.ratio
            margin_top = round(self.top_margin * self.ratio)
            margin_right = round(self.right_margin * self.ratio)
            margin_bottom = round(self.bottom_margin * self.ratio)
            margin_left = round(self.left_margin * self.ratio)
            self.lst_widget.append(widgets.Row(margin_top - line_height))

            self.content_width = self.content_width * self.ratio

            left = widgets.Panel(width=line_height, height=line_height,
                                 border={"width": [0, 2, 2, 0], "color": "#AAAAAA"},
                                 margin_left=margin_left - line_height)
            right = widgets.Panel(width=line_height, height=line_height,
                                  border={"width": [0, 0, 2, 2], "color": "#AAAAAA"}, )

            self.lst_widget.append(left)
            self.lst_widget.append(widgets.Col(round(self.content_width)))
            self.lst_widget.append(right)
            self.lst_widget.append(widgets.Row())

        else:
            if self.show_width is None:
                self.show_width = self.content_width
            self.col_width_ratio = (self.show_width - 32) / self.content_width

            margin_top = 0
            margin_right = 8
            margin_bottom = 0
            margin_left = 8
            if self.show_width < self.content_width:
                self.content_width = self.show_width

        self.margin = [margin_top, margin_right, margin_bottom, margin_left]
        o_content = widgets.Panel(width="100%%-%s" % (margin_left + margin_right),
                                  margin=[0, margin_right, 0, margin_left])

        for child in self.doc._body._body:
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, self.doc)
                o_panel = self.parse_paragraph(paragraph, "100%")
                o_content.append(o_panel)
                o_content.append(widgets.Row())

            elif isinstance(child, CT_Tbl):
                table = Table(child, self.doc)
                o_table = self.parse_table(table)
                o_content.append(o_table)
                o_content.append(widgets.Row())
            else:
                # print(type(child))
                pass

        self.lst_widget.append(o_content)
        self.lst_widget.append(widgets.Row())
        if self.only_content:
            pass
        elif self.is_margin:
            left = widgets.Panel(width=line_height, height=line_height,
                                 border={"width": [2, 2, 0, 0], "color": "#AAAAAA"},
                                 margin_left=self.margin[3] - line_height)
            right = widgets.Panel(width=line_height, height=line_height,
                                  border={"width": [2, 0, 0, 2], "color": "#AAAAAA"})
            self.lst_widget.append(left)
            self.lst_widget.append(widgets.Col(round(self.content_width * self.ratio)))
            self.lst_widget.append(right)

        self.lst_widget.append(widgets.Row(margin_bottom))

        return self.lst_widget

    def parse_paragraph_number(self):
        """解析段落符号"""
        numbering = self.doc.part.numbering_part.numbering_definitions._numbering
        dict_abstract_id = {}
        lst_num = numbering.xpath("./w:num")
        for num in lst_num:
            num_id = num.get(qn("w:numId"))
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            dict_abstract_id[abstract_id] = num_id

        self.dict_abstract = {}
        lst_abstract = numbering.xpath("./w:abstractNum")
        for abstract in lst_abstract:
            abstract_id = abstract.get(qn("w:abstractNumId"))

            lst_symbol = []
            lst_lvl = abstract.findall(qn("w:lvl"))
            for lvl in lst_lvl:
                type = lvl.find(qn("w:numFmt")).get(qn("w:val"))
                level = int(lvl.get(qn("w:ilvl")))
                left_indent = 0
                try:
                    align = lvl.find(qn("w:lvlJc")).get(qn("w:val"))
                    if align == "left":
                        ind = lvl.find(qn("w:pPr")).find(qn("w:ind"))
                        left = ind.get(qn("w:left"))
                        hanging = ind.get(qn("w:hanging"))
                        left_indent = conver.twips_2_px(Twips(int(left) - int(hanging)))
                except (BaseException,):
                    left_indent = 0

                format = lvl.find(qn("w:lvlText")).get(qn("w:val"))
                format = format.replace(str(level + 1), "s")
                symbol = {"id": "%s-%s" % (type, format), "type": type,
                          "format": format, "level": level, "left_indent": left_indent}
                lst_symbol.append(symbol)

            num_id = dict_abstract_id[abstract_id]
            self.dict_abstract[num_id] = lst_symbol

    def parse_paragraph(self, paragraph, width):
        self.paragraph = paragraph
        o_panel = widgets.Panel(width=width)
        runs = paragraph.runs
        # print(paragraph.text)

        line_spacing_rule, line_spacing = self.get_line_spacing(paragraph)
        space_before = self.get_space_before(paragraph)
        space_after = self.get_space_after(paragraph)
        o_panel.padding = [round(space_before * self.ratio), 0, round(space_after * self.ratio), 0]

        # 对齐
        horizontal = None
        if paragraph.alignment == WD_TAB_ALIGNMENT.CENTER:
            horizontal = "center"
        elif paragraph.alignment == WD_TAB_ALIGNMENT.RIGHT:
            horizontal = "right"
        # else:
        #     o_panel.horizontal = "left"

        left_indent = self.get_left_indent(paragraph)
        right_indent = self.get_right_indent(paragraph)
        first_line_indent = self.get_first_line_indent(paragraph)
        if first_line_indent < 0:
            left_indent = left_indent + first_line_indent
            first_line_indent = 0

        if left_indent:
            o_panel.padding[3] = round(left_indent * self.ratio)

        elif first_line_indent:
            o_panel.append(widgets.Col(round(first_line_indent * self.ratio)))

        if right_indent:
            o_panel.padding[1] = round(right_indent * self.ratio)

        o_panel.horizontal = horizontal

        font_size = 14
        if paragraph.style.font.size:
            font_size = conver.twips_2_px(paragraph.style.font.size)
        else:
            lst_sz = paragraph._p.xpath('./w:pPr/w:rPr/w:sz')
            if lst_sz:
                font_size = conver.twips_2_px(lst_sz[0].val)

        o_panel.min_height = round(conver.font_size_2_height(font_size) * self.ratio)

        try:
            snap_to_grid = paragraph._p.find(qn("w:pPr")).find(qn("w:snapToGrid")).val
        except (BaseException,):
            snap_to_grid = True

        # o_paragraph.snap_to_grid = snap_to_grid
        family = self.get_font_family(paragraph=paragraph)

        num_id, level = self.get_number(paragraph)
        if num_id:
            self.paragraph_number.setdefault(level, 0)
            paragraph_symbol = self.dict_abstract[str(num_id)][level]
            text = conver.paragraph_number_2_text(paragraph_symbol, self.paragraph_number[level])
            self.paragraph_number[level] += 1
            if paragraph_symbol["left_indent"]:
                o_panel.append(widgets.Col(round(paragraph_symbol["left_indent"] * self.ratio)))

            height = word_api.get_paragraph_line_height(line_spacing_rule, line_spacing, font_size * self.ratio)
            o_widget = widgets.Text(text=text, font={"size": round(font_size * self.ratio)},
                                    height=round(height * self.ratio), inline=True)
            o_panel.append(o_widget)

        else:
            if self.paragraph_number:
                lst_key = self.paragraph_number.keys()
                del self.paragraph_number[max(lst_key)]

        lst_script = re.findall(r"( *{{.*?}} *)", paragraph.text)
        if not lst_script:
            for run in runs:
                if run.text:
                    o_widget = self.parse_run(run, run.text, font_size, family)
                    o_widget.line_height = round(word_api.get_paragraph_line_height(line_spacing_rule, line_spacing,
                                                                                    o_widget.font["size"]) * self.ratio)
                else:
                    o_widget = self.parse_image(run)

                if not o_widget:
                    continue

                o_panel.append(o_widget)

            return o_panel

        dict_run = {}
        idx = 0
        for run in runs:
            dict_run[idx] = run
            idx += len(run.text)

        is_script = False
        is_script_end = False
        script = ""
        text = ""
        run = runs[0]
        text_len = len(paragraph.text)
        for i in range(text_len + 1):
            if (not is_script) and (i in dict_run) and text:
                text = text.replace("\t", "  ")
                o_widget = self.parse_run(run, text, font_size, family)
                o_widget.line_height = round(word_api.get_paragraph_line_height(line_spacing_rule, line_spacing,
                                                                                o_widget.font["size"]) * self.ratio)
                o_panel.append(o_widget)
                text = ""

            if i == text_len:
                if text:
                    o_widget = self.parse_run(run, text, font_size, family)
                    o_widget.line_height = round(word_api.get_paragraph_line_height(line_spacing_rule, line_spacing,
                                                                                    o_widget.font["size"]) * self.ratio)
                    o_panel.append(o_widget)

                if is_script_end:
                    if run.font.size:
                        widget_width = int(len(script) / 2.0) * conver.twips_2_px(run.font.size)
                    else:
                        widget_width = int(len(script) / 2.0) * font_size

                    name = re.search(r"{{(.*?)}}", script).groups()[0]
                    o_widget = self.make_widget(name, widget_width)
                    o_panel.append(o_widget)
                break

            char = paragraph.text[i]
            if i in dict_run:
                run = dict_run[i]
                # if is_script:

            if char == "{" and (text_len > i + 1) and paragraph.text[i + 1] == "{":
                is_script = True
                script += char

            elif char == "}" and (text_len > i + 1) and paragraph.text[i + 1] == "}":
                script += char
            elif char == "}":
                script += char
                is_script_end = True

            elif char == " " and run.underline:
                script += char

            elif is_script:
                if is_script_end:
                    if text:
                        o_widget = self.parse_run(run, text, font_size, family)
                        o_widget.line_height = word_api.get_paragraph_line_height(line_spacing_rule, line_spacing,
                                                                                  o_widget.font["size"] * self.ratio)
                        o_panel.append(o_widget)
                        text = ""

                    if run.font.size:
                        widget_width = int(len(script) / 2.0) * conver.twips_2_px(run.font.size)
                    else:
                        widget_width = int(len(script) / 2.0) * font_size

                    name = re.search(r"{{(.*?)}}", script).groups()[0]
                    o_widget = self.make_widget(name, widget_width)
                    o_panel.append(o_widget)
                    is_script = False
                    is_script_end = False
                    script = ""
                    text += char
                else:
                    script += char

            else:
                script = ""
                text += char

        return o_panel

    def parse_run(self, run, text=None, font_size=None, family=None):
        text = text or run.text
        text = text.replace("   ", "    ")
        font = {}
        bg = {}
        hyperlink, hyperlink_run = self.get_hyperlink(run, self.paragraph)
        if hyperlink:
            try:
                import shutil
                from django.conf import settings
                path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, "hyperlink"))
                if not os.path.exists(path):
                    os.makedirs(path)
                file_path = os.path.join(path, hyperlink)
                shutil.copy(os.path.join(os.path.split(self.path)[0], hyperlink), file_path)
                hyperlink = "%shyperlink/%s" % (settings.MEDIA_URL, hyperlink.strip("/\\"))
            except (BaseException,):
                pass

            o_text = widgets.Text(text=text, inline=True)
            font["color"] = "#0563C1"
            font["decoration"] = "underline"
        else:
            o_text = widgets.Text(text=text, inline=True)

        rId, width, height = self.get_bg_image(run)
        if rId:
            image_part = self.doc.part.related_parts[rId]
            name = image_part.partname

            try:
                from django.conf import settings
                path = os.path.abspath(os.path.join(settings.MEDIA_ROOT, "images", name.strip("/\\")))
                url = os.path.join(settings.MEDIA_URL, "images", name.strip("/\\"))
            except (BaseException,):
                path = os.path.join(os.getcwd(), "images", name.strip("/\\"))
                url = path

            tmp_path = os.path.split(path)[0]
            if not os.path.exists(tmp_path):
                os.makedirs(tmp_path)

            with open(path, 'wb') as fp:
                fp.write(image_part._blob)

            o_text.set("bg", {"image": url})
            o_text.set("height", height)

        if run.font.hidden:
            o_text.hide = 1

        # 字体
        if run.font.size:
            font["size"] = round(conver.twips_2_px(run.font.size) * self.ratio)
        else:
            font["size"] = round((font_size or 14) * self.ratio)

        # 在非表格内，word字体上下都会空了一点（测量值），如下特殊处理
        # if self.paragraph.paragraph_format.line_spacing is None:
        #     o_text.height = o_text.text_size + 21
        # else:
        #     o_text.height = o_text.text_size + 8

        if family:
            font["family"] = family
        else:
            font["family"] = self.get_font_family(run)

        if run.font.color.rgb:
            if run.font.color.rgb != RGBColor(0, 0, 0):  # docx bug，实际word没有颜色
                font["color"] = "#%s" % str(run.font.color.rgb)
        # else:
        #     font["color"] = "#000000"

        try:
            if run.font and run.font.highlight_color:
                bg["color"] = conver.word_color_idx_2_rgb(run.font.highlight_color)
        except (BaseException,):
            pass

        if run.font.underline:
            font["decoration"] = "underline"
            # o_text.under_line = font.underline

        if run.font.strike:
            font["decoration"] = "line-through"

        if run.bold or run.style.font.bold:
            font["weight"] = "bold"

        if run.font.italic:
            font["style"] = "italic"

        # if o_text.text:
        #     o_text.text = "&thinsp;".join(re.findall("emsp;+|[-=（）()0-9a-zA-Z().]+|[^-=（）()0-9a-zA-Z().]+",
        #                                              o_text.text, re.S))
        #     if re.findall(RE_LOWERCASE, o_text.text, re.S):
        #         o_text.text = o_text.text.replace("\xa0", "&ensp;").replace(" ", "&ensp;") \
        #             .replace("\u3000", "&ensp;")
        #     else:
        #         o_text.text = o_text.text.replace("\xa0", "&ensp;&ensp;").replace(" ", "&ensp;") \
        #             .replace("\u3000", "&ensp;&ensp;")

        if font:
            o_text.font = font
        if bg:
            o_text.bg = bg
        return o_text

    def parse_image(self, run):
        rId = self.get_image_id(run)
        if not rId:
            return
        image_part = self.doc.part.related_parts[rId]
        # name = image_part.partname

        # try:
        #     from django.conf import settings
        #     path = os.path.abspath(os.path.join(settings.MEDIA_ROOT, "images", name.strip("/\\")))
        #     url = os.path.join(settings.MEDIA_URL, "images", name.strip("/\\"))
        # except (BaseException,):
        #     path = os.path.join(os.getcwd(), "images", name.strip("/\\"))
        #     url = path

        # tmp_path = os.path.split(path)[0]
        # if not os.path.exists(tmp_path):
        #     os.makedirs(tmp_path)
        #
        # image_data = image_part._blob
        # with open(path, 'wb') as fp:
        #     fp.write(image_data)

        float = {}
        lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:positionH/wp:posOffset')
        if lst_node:
            node = lst_node[0]
            float["left"] = conver.twips_2_px(int(node.text))

        lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:positionV/wp:posOffset')
        if lst_node:
            node = lst_node[0]
            float["top"] = conver.twips_2_px(int(node.text))

        width, height = self.get_image_size(rId)
        if width is None and height is None:
            lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:extent')
            if lst_node:
                node = lst_node[0]
                width = conver.twips_2_px(int(node.get("cx")))
                height = conver.twips_2_px(int(node.get("cy")))

        o_image = widgets.Image(padding_right=3, data=image_part._blob, float=float)
        if width:
            o_image.width = round(width * self.ratio)
            o_image.height = round(height * self.ratio)

        return o_image

    def parse_table(self, table):
        data = list()
        min_row_height = dict()
        cell_horizontal = dict()
        cell_vertical = dict()
        merged_cells = self.get_table_merged_cells(table)
        for row_idx, row in enumerate(table.rows):
            row_height = self.get_table_row_height(row)
            if row_height:
                min_row_height[row_idx] = row_height

            row_data = []
            for col_idx, cell in enumerate(row.cells):
                b_merged, b_range, row_span, col_span = \
                    self.get_table_merged_info(merged_cells, row_idx, col_idx)

                if b_range:
                    row_data.append(None)
                    continue

                # 获取表格文字对齐（水平、垂直）
                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                    vertical = "center"
                elif cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
                    vertical = "bottom"
                else:
                    vertical = "top"

                if vertical != "top":
                    cell_vertical["%s-%s" % (row_idx, col_idx)] = vertical

                horizontal = self.get_table_cell_horizontal(cell)
                if horizontal != "left":
                    cell_horizontal["%s-%s" % (row_idx, col_idx)] = horizontal

                # background_color = self.get_table_cell_background_color(cell)
                cell_data = []
                for paragraph in cell.paragraphs:
                    o_table_sub = self.get_table(paragraph, cell)
                    if o_table_sub:
                        pass

                    else:
                        o_panel = self.parse_paragraph(paragraph, "100%")
                        cell_data.append(o_panel)

                row_data.append(cell_data)

            data.append(row_data)

        col_width = dict()
        width = 0
        for col_idx, col in enumerate(table.columns):
            width += col.width

        width = conver.twips_2_px(width)
        if width > self.content_width:  # 表格不能大于内容宽度
            ratio = self.content_width / width
            for col_idx, col in enumerate(table.columns):
                col_width[col_idx] = round(conver.twips_2_px(col.width) * ratio)
        else:
            for col_idx, col in enumerate(table.columns):
                col_width[col_idx] = round(conver.twips_2_px(col.width) * self.col_width_ratio)

        space_left = round(6 * self.ratio)
        space_right = round(6 * self.ratio)

        left_indent = round(self.get_table_left_indent(table) * self.ratio)
        o_table = widgets.LiteTable(data=data, row_height=min_row_height,
                                    # width="auto", max_width=self.width - self.padding[1] - self.padding[3],
                                    col_width=col_width,
                                    width=sum(col_width.values()),
                                    space_left=space_left, space_right=space_right,
                                    merged_cells=merged_cells,
                                    left_indent=left_indent,
                                    cell_horizontal=cell_horizontal, cell_vertical=cell_vertical,
                                    row_border={"color": "#000000"}, col_border={"color": "#000000"},
                                    auto_wrap=True)
        return o_table

    def make_widget(self, name, widget_width):
        o_widget = None
        name = name.strip()
        if name.lower().find("widget(") == 0:
            name = name.replace('”', '"').replace("，", ",")
            script = "widgets.W%s" % name[1:]
            o_widget = eval(script)

        elif name.lower().find("widgets.") == 0:
            name = name.replace('”', '"').replace("，", ",")
            o_widget = eval(name)

        if o_widget:
            if o_widget.get_attr_value("type", None):
                o_widget.type = o_widget.type.strip()
                if o_widget.type == "input" and o_widget.get_attr_value("input_type", None) == "textarea":
                    o_widget.border = {"color": "#FF0000"}
                    o_widget.clearable = False
                else:
                    o_widget.border = {"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0}
            else:
                o_widget.type = "input"
                o_widget.clearable = False
                o_widget.border = {"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0}

            if o_widget.get_attr_value("name", None):
                o_widget.name = o_widget.name.strip()
                o_widget.value = self.param.get(o_widget.name, None)

            if o_widget.get_attr_value("width", None) and isinstance(o_widget.width, str):
                if "%" in o_widget.width:
                    pass
                else:
                    o_widget.width = round(int(o_widget.width.strip()) * self.ratio)

            if o_widget.get_attr_value("height", None) and isinstance(o_widget.height, str):
                o_widget.height = round(int(o_widget.height.strip()) * self.ratio)

            if o_widget.get_attr_value("width", None) is None:
                if o_widget.type == "date_picker":
                    o_widget.width = round(140 * self.ratio)
                elif o_widget.type == "time_picker":
                    o_widget.width = round(120 * self.ratio)
                elif o_widget.type == "datetime_picker":
                    o_widget.width = round(200 * self.ratio)

        elif self.param.get(name, None):
            value = self.param[name]
            if isinstance(value, widgets.Widget):
                o_widget = value
                if not o_widget.get_attr_value("name", None):
                    o_widget.name = name

            elif isinstance(value, dict):
                o_widget = value
                if not o_widget.get("name", None):
                    o_widget["name"] = name
            else:
                o_widget = widgets.Input(name=name, value=value, width=round(widget_width * self.ratio),
                                         clearable=False,
                                         border={"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0},
                                         active={"border_color": "#000000"}, focus={"border_color": "#000000"})
        else:
            o_widget = widgets.Input(name=name, width=round(widget_width * self.ratio), clearable=False,
                                     border={"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0},
                                     active={"border_color": "#000000"}, focus={"border_color": "#000000"})

        return o_widget

    def get_font_family(self, run=None, paragraph=None):
        if run:
            r = run._r
            lst_font = r.xpath('./w:rPr/w:rFonts')
        else:
            r = paragraph._p
            lst_font = r.xpath('./w:pPr/w:rPr/w:rFonts')

        if lst_font:
            font = lst_font[0]
            return font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or \
                   font.get(qn("w:cs"))

    def get_table(self, paragraph, cell):
        for table in cell.tables:
            if paragraph._p.getprevious() == table._element:
                return table

    def get_table_merged_cells(self, table):
        """获取合并单元格"""
        row_num = len(table.rows)
        col_num = len(table.columns)

        merged_cells = []
        for row_idx in range(row_num):
            for col_idx in range(col_num):
                b_merged, b_range, row_span, col_span = \
                    self.get_table_merged_info(merged_cells, row_idx, col_idx)

                if b_merged or b_range:
                    continue

                cell = table.cell(row_idx, col_idx)
                end_row, end_col = self.get_table_merged_cell(table, cell, row_idx, col_idx)
                if (row_idx != end_row) or (col_idx != end_col):
                    # merged_cells.append("%s%s:%s%s" % (get_column_letter(col_idx + 1), row_idx + 1,
                    #                                    get_column_letter(c + 1), r + 1))
                    merged_cells.append("%s-%s:%s-%s" % (row_idx, col_idx, end_row, end_col))

        return merged_cells

    def get_table_merged_cell(self, table, cell, row_idx, col_idx):
        row_num = len(table.rows)
        col_num = len(table.columns)

        end_row = row_idx
        end_col = col_idx

        for r in range(row_idx + 1, row_num):
            if table.cell(r, col_idx)._tc == cell._tc:
                end_row = r
            else:
                break

        for c in range(col_idx + 1, col_num):
            if table.cell(row_idx, c)._tc == cell._tc:
                end_col = c
            else:
                break

        return end_row, end_col

    def get_table_merged_info(self, merged_cells, row_idx, col_idx):
        """获取合并的数据"""
        b_merged = False
        b_range = False
        row_span = 0
        col_span = 0
        row_col = '%s-%s' % (row_idx, col_idx)

        for s_range in merged_cells:
            s_range = str(s_range)
            begin_row_col, end_row_col = s_range.split(":")
            # begin_col, begin_row = re.findall("[A-Z]+|[0-9]+", begin_col_row, re.I)
            # end_col, end_row = re.findall("[A-Z]+|[0-9]+", end_col_row, re.I)
            begin_row, begin_col = begin_row_col.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)
            end_row, end_col = end_row_col.split("-")
            end_row = int(end_row)
            end_col = int(end_col)

            if begin_row_col == row_col:  # 开始合并
                b_merged = True
                row_span = end_row - begin_row + 1
                col_span = end_col - begin_col + 1
                break

            elif (row_idx >= begin_row) and (row_idx <= end_row) and \
                    (col_idx >= begin_col) and (col_idx <= end_col):
                b_range = True
                break

        return b_merged, b_range, row_span, col_span

    def get_table_row_height(self, row):
        if row.height:
            return conver.twips_2_px(row.height)

        cell = row.cells[0]
        tc = cell._tc
        # size = tc.xpath('./w:trPr/w:trHeight')
        lst_sz = tc.xpath('./w:p/w:pPr/w:rPr/w:sz')
        if lst_sz:
            sz = lst_sz[0]
            return conver.twips_2_px(sz.val)

    def get_table_cell_horizontal(self, cell):
        horizontal = "left"
        lst_horizontal = cell._tc.xpath("./w:p/w:pPr/w:jc")
        if lst_horizontal:
            horizontal = lst_horizontal[0].val
            if horizontal == WD_TABLE_ALIGNMENT.LEFT:
                horizontal = "left"
            elif horizontal == WD_TABLE_ALIGNMENT.CENTER:
                horizontal = "center"
            elif horizontal == WD_TABLE_ALIGNMENT.RIGHT:
                horizontal = "right"
            else:
                horizontal = "left"
        return horizontal

    def get_table_cell_background_color(self, cell):
        lst_shd = cell._tc.xpath('./w:tcPr/w:shd')
        if lst_shd:
            element = lst_shd[0]
            fill = element.get(qn("w:fill"))
            return "#%s" % fill

    def get_table_left_indent(self, table):
        left_indent = 0
        lst_path = table._tbl.xpath('./w:tblPr/w:tblInd')
        if lst_path:
            left_indent = conver.in_2_px(
                int(lst_path[0].get(qn("w:w"))) * 1 / 1440)  # dxa-指定该值的单位为点的二十分之一（1/1440英寸）
        else:
            lst_path = table._tbl.xpath('./w:tblPr/w:tblpPr')
            if lst_path:
                left_indent = conver.in_2_px(int(lst_path[0].get(qn("w:tblpX"))) / 1400) * -1
        return left_indent

    def get_image_id(self, run):
        namespace = {
            'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
            'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

        if hasattr(run.element, "drawing_lst"):
            for drawing in run.element.drawing_lst:
                for node in drawing.findall('.//a:blip', namespace):
                    for k, v in node.items():
                        if k.find("}embed") >= 0:
                            return v

    def get_image_size(self, rId):
        for shape in self.doc.inline_shapes:
            if rId == shape._inline.graphic.graphicData.pic.blipFill.blip.embed:
                if shape.width:
                    return conver.twips_2_px(shape.width), conver.twips_2_px(shape.height)
                return None, None

        return None, None

    def get_number(self, paragraph):
        """解析编号"""
        lst_num = paragraph._p.xpath('./w:pPr/w:numPr/w:numId')
        num_id, ilvl = None, None
        if lst_num:
            num_id = lst_num[0].val
            ilvl = paragraph._p.xpath('./w:pPr/w:numPr/w:ilvl')[0].val

        return num_id, ilvl

    def get_hyperlink(self, run, paragraph):
        """
        链接有两种格式
        1、在paragraph的run中存在，但没有链接内容，通过style_id中的name_val判断是否为链接文本
        2、在paragraph中不存在，在paragraph中查找是否有链接结点，判断是否在run之前
        """
        if not run.style:
            return None, None

        for style in self.doc.styles.element.style_lst:
            if run.style.style_id == style.styleId:
                if style.name_val == "Hyperlink":
                    p = paragraph._p
                    lst_instr = p.xpath('./w:r/w:instrText')
                    if lst_instr:
                        instr = lst_instr[0]
                        o_re = re.search('HYPERLINK "(.*)" ', instr.text, re.I)
                        if o_re:
                            hyperlink = o_re.groups()[0]
                        else:
                            hyperlink = instr.text

                        return hyperlink, None
                break

        lst_hyperlink = paragraph._p.xpath("./w:hyperlink")
        if lst_hyperlink:
            for hyperlink in lst_hyperlink:
                if hyperlink.getnext() == run.element:
                    rId = hyperlink.get(qn("r:id"))
                    hyperlink_run = None
                    for node in lst_hyperlink[0].getchildren():
                        # hyperlink_text += node.text
                        hyperlink_run = Run(node, paragraph)
                        break

                    for rel in self.doc.part.rels:
                        if rel == rId and \
                                self.doc.part.rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                            return self.doc.part.rels[rel]._target, hyperlink_run

        return None, None

    def get_bg_image(self, run):
        rId, width, height = None, None, None
        lst_element = run._r.xpath('./w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip')
        if lst_element:
            element = lst_element[0]
            rId = element.get(qn("r:embed"))

        lst_element = run._r.xpath('./w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic/pic:spPr/a:xfrm/a:ext')
        if lst_element:
            element = lst_element[0]
            width = element.get("cx")
            height = element.get("cy")
            if width is not None:
                width = conver.twips_2_px(int(width))

            if height is not None:
                height = conver.twips_2_px(int(height))

        return rId, width, height

    def get_left_indent(self, paragraph):
        left_indent = 0
        if paragraph.paragraph_format.left_indent:  # 块全部左缩进
            left_indent = conver.twips_2_px(paragraph.paragraph_format.left_indent)

        return left_indent

    def get_right_indent(self, paragraph):
        right_indent = 0
        if paragraph.paragraph_format.right_indent:
            right_indent = conver.twips_2_px(paragraph.paragraph_format.right_indent)
        return right_indent

    def get_first_line_indent(self, paragraph):
        first_line_indent = 0
        if paragraph.paragraph_format.first_line_indent:
            first_line_indent = conver.twips_2_px(paragraph.paragraph_format.first_line_indent)
        return first_line_indent

    def get_space_before(self, paragraph):
        space_before = 0
        if paragraph.paragraph_format.space_before:
            space_before = conver.twips_2_px(paragraph.paragraph_format.space_before)
        else:
            p = paragraph._p
            lst_xpath = p.xpath('./w:pPr/w:spacing')
            if lst_xpath:
                before = lst_xpath[0].get(qn("w:beforeLines"))
                if before:
                    space_before = before

        return space_before

    def get_space_after(self, paragraph):
        space_after = 0
        if paragraph.paragraph_format.space_after:
            space_after = int(conver.twips_2_px(paragraph.paragraph_format.space_after))
        else:
            p = paragraph._p
            lst_xpath = p.xpath('./w:pPr/w:spacing')
            if lst_xpath:
                after = lst_xpath[0].get(qn("w:afterLines"))
                if after:
                    space_after = after

        return space_after

    def get_line_spacing(self, paragraph):
        line_spacing = 0
        line_spacing_rule = None
        if paragraph.paragraph_format.line_spacing:
            line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
            if line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                line_spacing = conver.twips_2_px(paragraph.paragraph_format.line_spacing)
            else:
                line_spacing = paragraph.paragraph_format.line_spacing
            line_spacing_rule = int(line_spacing_rule)

        return line_spacing_rule, line_spacing


def generate(template_path, param=None, show_width=None,
             action_bar=False, save_url=None, export_url=None,
             has_export_word=True,
             has_export_pdf=False,
             has_print=False,
             has_close=True,
             custom_button=None,
             ):
    """
    template_file：模板文件(只支持*.docx)
    param:参数 {"模板关键字":"显示值", "模板关键字":"widget对象"}
    """
    obj = Compiler(template_path, param, show_width)

    lst_widget = obj.parse()
    if show_width is None or show_width > obj.width:
        width = obj.width
    else:
        width = show_width

    o_panel = widgets.Panel(vertical="top", horizontal="left", width=width,
                            bg_color="#FFFFFF", height="100%", scroll={"y": "auto"})
    if action_bar:
        o_panel_sub = widgets.Panel(width="100%", horizontal="right")
        if save_url:
            o_step = step.RunScript(save_url)
            o_panel_sub.append(widgets.Button(prefix="v-save", text="保存", font_size=14, width=70,
                                              margin=2, step=o_step))

        o_panel_sub.append(widgets.Input(name="v-template-file", value=template_path, hide=True, width=0))
        if export_url is None:
            o_step = step.RunScript("utils.word_template_2_vadmin.export_word")
        else:
            o_step = step.RunScript(export_url)

        if has_export_word:
            o_panel_sub.append(widgets.Button(prefix="el-icon-download", text="导出Word", font_size=14,
                                              margin=2, step=o_step))

        if has_export_pdf:
            o_step = step.RunScript("utils.word_template_2_vadmin.export_pdf")
            o_panel_sub.append(widgets.Button(prefix="el-icon-download", text="导出PDF", font_size=14,
                                              margin=2, step=o_step))

        if custom_button:
            o_panel_sub.append(custom_button)

        if has_print:
            o_panel_sub.append(widgets.Button(prefix="v-print", text="打印", font_size=14, width=70,
                                              margin=2, step=step.Print(name="panel-word")))

        if has_close:
            o_panel_sub.append(widgets.Button(prefix="el-icon-close", text="关闭", font_size=14, width=70,
                                              margin=2, step=step.LayerClose()))
        o_panel.append(o_panel_sub)

    o_panel_word = widgets.Panel(name="panel-word", width="100%")
    o_panel_word.append(lst_widget)
    return o_panel.append(o_panel_word)


def export_word(request):
    from django.conf import settings
    from utils import word_template_2_word
    from vadmin import common
    widget_data = json.loads(request.POST.get(const.SUBMIT_WIDGET, "{}"))
    template_path = widget_data['v-template-file']
    path, file_name = os.path.split(template_path)
    file_path = common.get_export_path(file_name)
    word_template_2_word.generate(template_path, file_path, widget_data)
    file_url = file_path.replace(settings.BASE_DIR, "")
    return step.DownloadFile(file_url)


def export_pdf(request, template_path=None, param=None):
    from utils import vadmin_2_pdf
    from vadmin import common
    if template_path is None:
        param = json.loads(request.POST.get(const.SUBMIT_WIDGET, "{}"))
        template_path = param['v-template-file']

    path, file_name = os.path.split(template_path)
    pdf_path = common.get_export_path(file_name, suffix=".pdf")

    obj = Compiler(template_path, param, only_content=True)

    lst_widget = obj.parse()
    vadmin_2_pdf.generate(pdf_path, lst_widget,
                          page_width=obj.page_width, page_height=obj.page_height,
                          top_margin=obj.top_margin, bottom_margin=obj.bottom_margin,
                          left_margin=obj.left_margin, right_margin=obj.right_margin)
    file_url = common.get_download_url(pdf_path)
    return step.DownloadFile(file_url)


# def export_pdf(request, template_path=None, param=None):
#
#     """
#     Chrome驱动文件下载
#     # https://selenium-python.readthedocs.io/installation.html
#     # https://sites.google.com/a/chromium.org/chromedriver/downloads
#     mac os
#     chromedriver文件放在“/usr/local/bin”目录下
#
#     linux版本安装
#     https://www.cnblogs.com/xiaomifeng0510/p/12072081.html
#     """
#     # import pdfkit
#     from django.conf import settings
#     from utils import word_template_2_word
#     from utils import word_convert
#     from vadmin import common
#     if template_path is None:
#         param = json.loads(request.POST.get(const.SUBMIT_WIDGET, "{}"))
#         template_path = param['v-template-file']
#
#     path, file_name = os.path.split(template_path)
#     word_path = common.get_export_path(file_name)
#     word_template_2_word.generate(template_path, word_path, param)
#
#     pdf_path = common.get_export_path(file_name, suffix=".pdf")
#     word_convert.word2pdf(word_path, pdf_path)
#     file_url = pdf_path.replace(settings.BASE_DIR, "")
#     return step.DownloadFile(file_url)
#     # obj = Compiler(template_path, param, show_width)
#     # lst_widget = obj.parse()
#     #
#     # if show_width is None or show_width > obj.width:
#     #     width = obj.width
#     # else:
#     #     width = show_width
#     #
#     # o_panel = widgets.Panel(width="100%", readonly=True)
#     # o_panel_word = widgets.Panel(width=width)
#     # o_panel_word.append(lst_widget)
#     # o_panel.append(o_panel_word)
#     # html = o_panel.render_html(readonly=True)
#
#     # from selenium import webdriver
#     # url = request.META['HTTP_REFERER'] + "#" + const.URL_SHOW_WORD_TEMPLATE % template_path
#     # url = common.make_url(url, param=param)
#     # browser = webdriver.Chrome()
#     # browser.get(url)
#     # html = browser.page_source
#     # browser.close()
#
#     # options = {
#     #     'quiet': ''
#     # }
#     #
#     # options = {
#     #     # 'page-size': 'Letter',
#     #     'margin-top': '0.75in',
#     #     'margin-right': '0.75in',
#     #     'margin-bottom': '0.75in',
#     #     'margin-left': '0.75in',
#     #     # 'encoding': "UTF-8",
#     #     # 'custom-header': [
#     #     #     ('Accept-Encoding', 'gzip')
#     #     # ],
#     #     # 'cookie': [
#     #     #     ('cookie-name1', 'cookie-value1'),
#     #     #     ('cookie-name2', 'cookie-value2'),
#     #     # ],
#     #     # 'no-outline': None
#     # }
#
#     # path, file_name = os.path.split(template_path)
#     # file_path = common.get_export_path(file_name, suffix=".pdf")
#     # html_path = common.get_export_path(file_name, suffix=".html")
#     # open(html_path, "w").write('<head><meta charset="UTF-8"></head>' + html)
#     #
#     # if "mac" in platform():
#     #     path = "/usr/local/bin/wkhtmltopdf"
#     # elif "Windows" in platform():
#     #     path = ""
#     # else:
#     #     path = "/usr/local/bin/wkhtmltopdf"
#     #
#     # config = pdfkit.configuration(wkhtmltopdf=path)
#     # pdfkit.from_url(html_path, file_path, configuration=config)  # 从URL生成


if __name__ == "__main__":
    generate(r"/Users/wangbin/Documents/docx/word/报请指定居所监视居住意见书.docx")
    pass
