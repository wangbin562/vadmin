# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
html转doc
doc转html
"""
import sys
import os
import logging
import traceback
import docx
import requests
import copy
import math
import re
import collections
from docx import Document
from docx import shared
from docx.opc import constants
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.numbering import CT_Num
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn

# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
# from lxml import etree
from colorsys import rgb_to_hsv
from PIL import Image

try:
    import xml.etree.cElementTree as ET
except (BaseException,):
    import xml.etree.ElementTree as ET
    # ElementTree.tostring(root)
try:
    from django.conf import settings

    DEBUG = settings.DEBUG
except (BaseException,):
    DEBUG = True

from logging.handlers import RotatingFileHandler

if DEBUG:
    console = logging.StreamHandler()
    formatter = logging.Formatter(
        '%(asctime)s [%(threadName)s:%(thread)d] [%(name)s:%(lineno)d] [%(levelname)s]- %(message)s')
    console.setFormatter(formatter)
    logger = logging.getLogger('')
    logger.addHandler(console)
    logger.setLevel(logging.INFO)
else:
    file_path = sys.modules['__main__'].__file__
    file_name = os.path.split(file_path)[1]
    module_name = os.path.splitext(file_name)[0]
    o_handler = RotatingFileHandler('%s.log' % module_name, maxBytes=10 * 1024 * 1024, backupCount=5)
    formatter = logging.Formatter(
        '%(asctime)s [%(threadName)s:%(thread)d] [%(name)s:%(lineno)d] [%(levelname)s]- %(message)s')
    o_handler.setFormatter(formatter)
    logger = logging.getLogger('')
    logger.addHandler(o_handler)
    logger.setLevel(logging.INFO)

############################### 如下部分要根据为使用配置部分 #############################################
CON_BASE_ADDRESS = "C:\wangbin\code_sync\code_sync\zhwhhuiyu\server_v2"
CON_BASE_URL = "http://zh.whhuiyu.com/"

CON_DEFAULT_FONT = u"微软雅黑"
CON_DEFAULT_FONT_EN = "SansSerif"

CON_DEFAULT_FONT_SIZE = "16px"

# 如下是页面样式（页面css定义的样式，各富文本编辑器不一样，也可以增加，例如：增加a标签 A_LABEL_STYLE = {}
P_LABEL_STYLE = {"margin": "5px 0;"}  # p标签页面外的css样式
PRE_LABEL_STYLE = {"margin": ".5em 0;", "padding": ".4em .6em;", "background-color": "#f8f8f8"}  # pre标签页面外的css样式
TD_LABEL_STYLE = {"padding": "5px 10px;"}
TH_LABEL_STYLE = {"background-color": "#F7F7F7"}
H1_LABEL_STYLE = H2_LABEL_STYLE = H3_LABEL_STYLE = H4_LABEL_STYLE = H5_LABEL_STYLE = H6_LABEL_STYLE = {
    "font-weight": "bold"}

######################################## end #############################################################

MAX_WIDTH_IN = 8.5  # 最大宽度(英寸），不用修改docx包固定
MAX_CONTENT_WIDTH_IN = 6.0  # 最大内容宽度(英寸），不用修改docx包固定
CON_DPI = 96.0  # dpi

DICT_COLOR_NAME_2_HEX = {
    "black": "#000000",
    "navy": "#000080",
    "darkblue": "#00008b",
    "mediumblue": "#0000cd",
    "blue": "#0000ff",
    "darkgreen": "#006400",
    "green": "#008000",
    "teal": "#008080",
    "darkcyan": "#008b8b",
    "deepskyblue": "#00bfff",
    "darkturquoise": "#00ced1",
    "mediumspringgreen": "#00fa9a",
    "lime": "#00ff00",
    "springgreen": "#00ff7f",
    "aqua": "#00ffff",
    "cyan": "#00ffff",
    "midnightblue": "#191970",
    "dodgerblue": "#1e90ff",
    "lightseagreen": "#20b2aa",
    "forestgreen": "#228b22",
    "seagreen": "#2e8b57",
    "darkslategray": "#2f4f4f",
    "limegreen": "#32cd32",
    "mediumseagreen": "#3cb371",
    "turquoise": "#40e0d0",
    "royalblue": "#4169e1",
    "steelblue": "#4682b4",
    "darkslateblue": "#483d8b",
    "mediumturquoise": "#48d1cc",
    "indigo": "#4b0082",
    "darkolivegreen": "#556b2f",
    "cadetblue": "#5f9ea0",
    "cornflowerblue": "#6495ed",
    "mediumaquamarine": "#66cdaa",
    "dimgray": "#696969",
    "slateblue": "#6a5acd",
    "olivedrab": "#6b8e23",
    "slategray": "#708090",
    "lightslategray": "#778899",
    "mediumslateblue": "#7b68ee",
    "lawngreen": "#7cfc00",
    "chartreuse": "#7fff00",
    "aquamarine": "#7fffd4",
    "maroon": "#800000",
    "purple": "#800080",
    "olive": "#808000",
    "gray": "#808080",
    "skyblue": "#87ceeb",
    "lightskyblue": "#87cefa",
    "blueviolet": "#8a2be2",
    "darkred": "#8b0000",
    "darkmagenta": "#8b008b",
    "saddlebrown": "#8b4513",
    "darkseagreen": "#8fbc8f",
    "lightgreen": "#90ee90",
    "mediumpurple": "#9370db",
    "darkviolet": "#9400d3",
    "palegreen": "#98fb98",
    "darkorchid": "#9932cc",
    "yellowgreen": "#9acd32",
    "sienna": "#a0522d",
    "brown": "#a52a2a",
    "darkgray": "#a9a9a9",
    "lightblue": "#add8e6",
    "greenyellow": "#adff2f",
    "paleturquoise": "#afeeee",
    "lightsteelblue": "#b0c4de",
    "powderblue": "#b0e0e6",
    "firebrick": "#b22222",
    "darkgoldenrod": "#b8860b",
    "mediumorchid": "#ba55d3",
    "rosybrown": "#bc8f8f",
    "darkkhaki": "#bdb76b",
    "silver": "#c0c0c0",
    "mediumvioletred": "#c71585",
    "indianred": "#cd5c5c",
    "peru": "#cd853f",
    "chocolate": "#d2691e",
    "tan": "#d2b48c",
    "lightgray": "#d3d3d3",
    "thistle": "#d8bfd8",
    "orchid": "#da70d6",
    "goldenrod": "#daa520",
    "palevioletred": "#db7093",
    "crimson": "#dc143c",
    "gainsboro": "#dcdcdc",
    "plum": "#dda0dd",
    "burlywood": "#deb887",
    "lightcyan": "#e0ffff",
    "lavender": "#e6e6fa",
    "darksalmon": "#e9967a",
    "violet": "#ee82ee",
    "palegoldenrod": "#eee8aa",
    "lightcoral": "#f08080",
    "khaki": "#f0e68c",
    "aliceblue": "#f0f8ff",
    "honeydew": "#f0fff0",
    "azure": "#f0ffff",
    "sandybrown": "#f4a460",
    "wheat": "#f5deb3",
    "beige": "#f5f5dc",
    "whitesmoke": "#f5f5f5",
    "mintcream": "#f5fffa",
    "ghostwhite": "#f8f8ff",
    "salmon": "#fa8072",
    "antiquewhite": "#faebd7",
    "linen": "#faf0e6",
    "lightgoldenrodyellow": "#fafad2",
    "oldlace": "#fdf5e6",
    "red": "#ff0000",
    "fuchsia": "#ff00ff",
    "magenta": "#ff00ff",
    "deeppink": "#ff1493",
    "orangered": "#ff4500",
    "tomato": "#ff6347",
    "hotpink": "#ff69b4",
    "coral": "#ff7f50",
    "darkorange": "#ff8c00",
    "lightsalmon": "#ffa07a",
    "orange": "#ffa500",
    "lightpink": "#ffb6c1",
    "pink": "#ffc0cb",
    "gold": "#ffd700",
    "peachpuff": "#ffdab9",
    "navajowhite": "#ffdead",
    "moccasin": "#ffe4b5",
    "bisque": "#ffe4c4",
    "mistyrose": "#ffe4e1",
    "blanchedalmond": "#ffebcd",
    "papayawhip": "#ffefd5",
    "lavenderblush": "#fff0f5",
    "seashell": "#fff5ee",
    "cornsilk": "#fff8dc",
    "lemonchiffon": "#fffacd",
    "floralwhite": "#fffaf0",
    "snow": "#fffafa",
    "yellow": "#ffff00",
    "lightyellow": "#ffffe0",
    "ivory": "#fffff0",
    "white": "#ffffff",
}


def get_page_style(label):
    """
    获取标签页面样式
    """
    dict_global = globals()
    key = "%s_LABEL_STYLE" % label.upper()
    return dict_global.get(key, {})


def download_image(url):
    """
    下载图片
    :return:
    """
    name = os.path.split(url)[1]
    path = os.path.abspath(os.path.join(CON_BASE_ADDRESS, name))
    if os.path.exists(path):
        return path

    ir = requests.get(url)
    if ir.status_code == 200:
        open(path, 'wb').write(ir.content)
        return path

    return None


def delete_paragraph(paragraph):
    """
    删除段落
    Args:
        paragraph:

    Returns:

    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class UnitChangeApi(object):
    """
    pt (point，磅) 是一个物理长度单位，指的是72分之一英寸
    px (pixel，像素)是一个虚拟长度单位，是计算机系统的数字化图像长度单位，如果px要换算成物理长度，
    需要指定精度DPI(Dots Per Inch，每英寸像素数)，在扫描打印时一般都有DPI可选。
    Windows系统默认是96dpi，Apple系统默认是72dpi。
    em是一个相对长度单位，最初是指字母M的宽度，故名em。现指的是字符宽度的倍数，
    用法类似百分比，如：0.8em, 1.2em,2em等。通常1em=16px。

    """

    @staticmethod
    def pt_2_px(pt):
        return pt * 4 / 3.0

    @staticmethod
    def px_2_pt(px):
        return px * 3 / 4.0

    @staticmethod
    def cm_2_pt(cm):
        return cm * 72.0 / 2.54

    @staticmethod
    def pt_2_cm(pt):
        return pt * 2.54 / 76

    @staticmethod
    def px_2_mm(px):
        """
        转换成毫米 1lin = 2.54cm = 25.4 mm
        """
        return px * 1 / CON_DPI * 25.4

    @staticmethod
    def px_2_in(px):
        """
        像素转英寸
        """
        return px * 1 / CON_DPI

    @staticmethod
    def in_2_px(inc):
        """
        英求转像素
        Args:
            px:

        Returns:

        """
        return inc * CON_DPI

    @staticmethod
    def in_2_pt(inc):
        """
        英寸转pt
        """
        return inc * 72.0


class FontApi(object):
    """
    字体接口
    """
    # arial, helvetica, sans - serif
    CSS_2_DOC_EN = {
        u"verdana": u"Verdana",
        u"arial": u"Arial",
        u"helvetica": u"Helvetica",
        u"sans-serif": u"SansSerif",
        u"consolas": u"Consolas",
        u"simsun": u"SimSun-ExtB",
        u"宋体": u"SimSun-ExtB",
        u"微软雅黑": u"Microsoft YaHei UI",
        u"simkai": u"楷体",
        u"楷体": u"楷体",
        u"楷体_gb2312": u"楷体",
        u"黑体": u"黑体",
        u"simhei": u"SimHei",
        u"隶书": u"隶书",
        u"simli": u"隶书",
        u"andale mono": u"SansSerif",
        u"arial black": u"Arial Black",
        u"avant garde": u"Arial Black",
        u"comic sans ms": u"Comic Sans MS",
        u"impact": u"Impact",
        u"chicago": u"Impact",
        u"times new roman": u"Times New Roman",
        u"courier new": u"Courier New",
        u"wingdings": u"Wingdings",
        u"华文细黑": u"华文细黑",
    }

    CSS_2_DOC_CHS = {
        u"verdana": u"微软雅黑",
        u"arial": u"微软雅黑",
        u"helvetica": u"微软雅黑",
        u"sans-serif": u"微软雅黑",
        u"consolas": u"微软雅黑",
        u"simsun": u"宋体",
        u"宋体": u"宋体",
        u"微软雅黑": u"微软雅黑",
        u"simkai": u"楷体",
        u"楷体": u"楷体",
        u"楷体_gb2312": u"楷体",
        u"黑体": u"黑体",
        u"隶书": u"隶书",
        u"simhei": u"隶书",
        u"andale mono": u"微软雅黑",
        u"arial black": u"微软雅黑",
        u"avant garde": u"微软雅黑",
        u"comic sans ms": u"微软雅黑",
        u"impact": u"微软雅黑",
        u"chicago": u"微软雅黑",
        u"times new roman": u"微软雅黑",
        u"courier new": u"微软雅黑",
        u"wingdings": u"Wingdings",
        u"华文细黑": u"华文细黑",
    }

    @staticmethod
    def split(text):
        """
        文字分割(中英文分割)
        isalnum() 方法检测字符串是否由字母和数字组成。
        isalpha() 方法检测字符串是否只由字母组成
        isdigit() 方法检测字符串是否只由数字组成。
        isnumeric() 方法检测字符串是否只由数字组成。这种方法是只针对unicode对象。
        isspace() 方法检测字符串是否只由空格组成。
        istitle() 方法检测字符串中所有的单词拼写首字母是否为大写，且其他字母为小写。
        isupper() 方法检测字符串中所有的字母是否都为大写。
        """
        lst_split = []
        lst_en = []
        lst_chs = []
        for s in text:
            ascii_val = ord(s)
            if (ascii_val > 31) and (ascii_val < 128):
                lst_en.append(s)
                if lst_chs:
                    lst_split.append(("".join(lst_chs), False))
                    lst_chs = []
            else:
                lst_chs.append(s)
                if lst_en:
                    lst_split.append(("".join(lst_en), True))
                    lst_en = []

        if lst_chs:
            lst_split.append(("".join(lst_chs), False))

        if lst_en:
            lst_split.append(("".join(lst_en), True))

        return lst_split

    @staticmethod
    def change(name, is_en=False):
        """
        字体转换（根据中英文分别转换）
        """
        name = name.replace("'", "").replace('"', "").lower()
        for single_name in name.split(","):
            single_name = single_name.strip()
            if single_name == "":
                continue
            if is_en:
                if single_name in FontApi.CSS_2_DOC_EN:
                    return FontApi.CSS_2_DOC_EN[single_name]
            else:
                if single_name in FontApi.CSS_2_DOC_CHS:
                    return FontApi.CSS_2_DOC_CHS[single_name]

        if is_en:
            return CON_DEFAULT_FONT_EN

        return CON_DEFAULT_FONT


class StyleManageApi(object):

    @staticmethod
    def get_style(key, lst_style, dict_label_style, default_style=None, not_key=None):
        """
        先获取当前样式（标签中的sytle关键字），再获取css中标签样式，再获取上级样式
        Args:
            key:
            lst_style:
            dict_label_style:
            default_style:
        Returns:

        """
        val = default_style
        if len(lst_style) > 0:
            dict_style = lst_style[-1]
        else:
            dict_style = {}

        if key in dict_style:
            val = dict_style[key]
        elif key in dict_label_style:
            val = dict_label_style[key]
        else:
            for i in range(len(lst_style) - 2, -1, -1):
                dict_style = lst_style[i]
                if key in dict_style:
                    val = dict_style[key]

        if (not_key is not None) and (not_key in val):
            if val in dict_label_style.get(key, ""):
                val = StyleManageApi.get_style(key, lst_style[1:], {}, default_style, not_key)
            else:
                val = StyleManageApi.get_style(key, lst_style[1:], dict_label_style, default_style, not_key)

        return val

    @staticmethod
    def size_2_pt(size, lst_style, dict_label_style=None, key=None):
        """
        html单位格式统一转换成pt
        """
        size = str(size).lower().strip().strip(";")

        if "pt" in size:
            pt = size.strip("pt").strip()
            return float(pt)

        elif "px" in size:
            px = size.strip("px").strip()
            return UnitChangeApi.px_2_pt(float(px))

        elif "cm" in size:
            cm = size.strip("cm").strip()
            return UnitChangeApi.cm_2_pt(float(cm))

        elif "in" is size:
            inc = size.strip("in").strip()
            return UnitChangeApi.in_2_pt(float(inc))

        elif "mm" in size:
            mm = size.strip("mm").strip()
            return UnitChangeApi.cm_2_pt(float(mm) * 10)

        elif "rem" in size:
            rem = size.strip("rem").strip()
            return float(rem) * StyleManageApi.size_2_pt(CON_DEFAULT_FONT_SIZE, None, None)

        elif "em" in size:
            em = size.strip("em").strip()
            font_size = StyleManageApi.get_style("font-size", lst_style, dict_label_style, CON_DEFAULT_FONT_SIZE,
                                                 not_key="em")
            pt = StyleManageApi.size_2_pt(font_size, lst_style, dict_label_style)
            return float(em) * pt

        elif "pc" in size:
            pc = size.strip("pc").strip()
            return UnitChangeApi.px_2_pt(float(pc) * 12)

        elif "%" in size:
            percentage = size.strip("%").strip()
            font_size = StyleManageApi.get_style("font-size", lst_style, dict_label_style, CON_DEFAULT_FONT_SIZE,
                                                 not_key="%")
            pt = StyleManageApi.size_2_pt(font_size, lst_style, dict_label_style)
            return float(percentage) / 100.0 * pt

        elif "auto" in size:
            return 0.0

        elif "normal" in size:
            return UnitChangeApi.px_2_pt(float(CON_DEFAULT_FONT_SIZE.strip("px")))

        elif "inherit" in size:
            style = StyleManageApi.get_style(key, lst_style[0:-1], dict_label_style, CON_DEFAULT_FONT_SIZE,
                                             not_key="em")
            pt = StyleManageApi.size_2_pt(style, lst_style[0:-1], dict_label_style, key)
            return pt
        else:
            return float(size)  # 默认是pt

    @staticmethod
    def parse_size(size, lst_style, dict_label_style, key=None):
        """
        单位格式化
        Args:
            size:
            lst_style:
            dict_label_style:
        Returns:

        """
        try:
            pt = StyleManageApi.size_2_pt(size, lst_style, dict_label_style, key)
            return shared.Pt(pt)

        except (BaseException,):
            logger.error(traceback.format_exc())
            return shared.Pt(UnitChangeApi.px_2_pt(float(CON_DEFAULT_FONT_SIZE.strip("px"))))

    @staticmethod
    def line_spacing_change(size, lst_style, dict_label_style):
        """
        行距格式化
        """
        # size = str(size).lower().strip()
        try:
            return float(size)
        except (BaseException,):
            pt = StyleManageApi.size_2_pt(size, lst_style, dict_label_style)
            font_size = StyleManageApi.get_style("font-size", lst_style, dict_label_style, CON_DEFAULT_FONT_SIZE)
            font_size_pt = StyleManageApi.size_2_pt(font_size, lst_style, dict_label_style)
            return float(pt) / font_size_pt

    @staticmethod
    def style_color_2_rgb(color):
        """
        颜色格式化，css样式颜色转化成rgb
        """
        if color in DICT_COLOR_NAME_2_HEX:
            color = DICT_COLOR_NAME_2_HEX[color]
            return int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)

        elif "rgb" in color:
            rgb = color.strip().replace('rgb', '').replace("(", "").replace(")", "").strip().split(",")
            return int(rgb[0]), int(rgb[1]), int(rgb[2])

        elif "#" in color:
            try:
                return int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            except (BaseException,):
                pass

        elif type(color) in [type([]), type(())]:
            return color

        return 0xFF, 0xFF, 0xFF

    @staticmethod
    def style_color_2_hex(color):
        """
        颜色格式化，css样式颜色转化成十六进制
        """
        r, g, b = StyleManageApi.style_color_2_rgb(color)
        return "%02X%02X%02X" % (r, g, b)

    @staticmethod
    def background_color_change(color):
        """
        背景颜色格式化
        """
        # word文字背景颜色只固定几种格式，html可以自定义颜色，需要计算相近的颜色
        dict_background_color = {
            (255, 255, 255): WD_COLOR_INDEX.WHITE,  # 白色
            (0, 0, 0): WD_COLOR_INDEX.BLACK,  # 黑色
            (0, 0, 255): WD_COLOR_INDEX.BLUE,  # 蓝
            (0, 255, 0): WD_COLOR_INDEX.BRIGHT_GREEN,  # 鲜绿
            (0, 0, 128): WD_COLOR_INDEX.DARK_BLUE,  # 深蓝
            (128, 0, 0): WD_COLOR_INDEX.DARK_RED,  # 深红
            (128, 128, 0): WD_COLOR_INDEX.DARK_YELLOW,  # 深黄
            (192, 192, 192): WD_COLOR_INDEX.GRAY_25,  # 灰色， 25%
            (128, 128, 128): WD_COLOR_INDEX.GRAY_50,  # 灰色， 50%
            (0, 128, 0): WD_COLOR_INDEX.GREEN,  # 绿色
            (255, 0, 255): WD_COLOR_INDEX.PINK,  # 粉红
            (255, 0, 0): WD_COLOR_INDEX.RED,  # 红色
            (0, 128, 128): WD_COLOR_INDEX.TEAL,  # 青色
            (0, 255, 255): WD_COLOR_INDEX.TURQUOISE,  # 青绿
            (128, 0, 128): WD_COLOR_INDEX.VIOLET,  # 紫罗兰
            (255, 255, 0): WD_COLOR_INDEX.YELLOW,  # 黄色
        }

        def to_hsv(c):
            """ converts color tuples to floats and then to hsv """
            return rgb_to_hsv(*[x / 255.0 for x in c])  # rgb_to_hsv wants floats!

        def color_dist(c1, c2):
            """ returns the squared euklidian distance between two color vectors in hsv space """
            return sum((a - b) ** 2 for a, b in zip(to_hsv(c1), to_hsv(c2)))

        def min_color_diff(color_to_match, colors):
            """ returns the `(distance, color_name)` with the minimal distance to `colors`"""
            return min(  # overal best is the best match to any color:
                (color_dist(color_to_match, c), i)  # (distance to `test` color, color name)
                for c, i in colors.items())

        r, g, b = StyleManageApi.style_color_2_rgb(color)
        if (r, g, b) in dict_background_color:
            return dict_background_color[(r, g, b)]

        del dict_background_color[(255, 255, 255)]  # 删除白色
        value, idx = min_color_diff((r, g, b), dict_background_color)
        return idx

    @staticmethod
    def parse_padding(lst_style, dict_label_style):
        """
        css padding
        """

        def _parse_padding(lst_style, dict_label_style):
            dict_style = lst_style[-1]
            top = bottom = left = right = None
            if "padding" in dict_style:
                padding = dict_style["padding"]
                lst_padding = padding.strip().strip(";").split()
                num = len(lst_padding)
                if num == 4:
                    top = StyleManageApi.size_2_pt(lst_padding[0], lst_style, dict_label_style)
                    right = StyleManageApi.size_2_pt(lst_padding[1], lst_style, dict_label_style)
                    bottom = StyleManageApi.size_2_pt(lst_padding[2], lst_style, dict_label_style)
                    left = StyleManageApi.size_2_pt(lst_padding[3], lst_style, dict_label_style)

                elif num == 3:
                    top = StyleManageApi.size_2_pt(lst_padding[0], lst_style, dict_label_style)
                    left = right = StyleManageApi.size_2_pt(lst_padding[1], lst_style, dict_label_style)
                    bottom = StyleManageApi.size_2_pt(lst_padding[2], lst_style, dict_label_style)

                elif num == 2:
                    top = bottom = StyleManageApi.size_2_pt(lst_padding[0], lst_style, dict_label_style)
                    left = right = StyleManageApi.size_2_pt(lst_padding[1], lst_style, dict_label_style)
                elif num == 1:
                    top = bottom = left = right = StyleManageApi.size_2_pt(lst_padding[0], lst_style, dict_label_style)

            if "padding-top" in dict_style:
                top = StyleManageApi.size_2_pt(dict_style["padding-top"], lst_style, dict_label_style)

            if "padding-bottom" in dict_style:
                bottom = StyleManageApi.size_2_pt(dict_style["padding-bottom"], lst_style, dict_label_style)

            if "padding-left" in dict_style:
                left = StyleManageApi.size_2_pt(dict_style["padding-left"], lst_style, dict_label_style)

            if "padding-right" in dict_style:
                right = StyleManageApi.size_2_pt(dict_style["padding-right"], lst_style, dict_label_style)

            return top, bottom, left, right

        top, bottom, left, right = _parse_padding(lst_style, dict_label_style)

        if (top is None) or (bottom is None) or (left is None) is (right is None):
            lable_top, lable_bottom, lable_left, lable_right = _parse_padding([dict_label_style, ], dict_label_style)
            if top is None:
                top = lable_top

            if bottom is None:
                bottom = lable_bottom

            if left is None:
                left = lable_left

            if right is None:
                right = lable_right

        return top, bottom, left, right

    @staticmethod
    def parse_margin(lst_style, dict_label_style):
        """
        css margin
        """

        def _parse_margin(lst_style, dict_label_style):
            dict_style = lst_style[-1]
            top = bottom = left = right = None
            if "margin" in dict_style:
                margin = dict_style["margin"]
                lst_margin = margin.strip().strip(";").split()
                num = len(lst_margin)
                if num == 4:
                    top = StyleManageApi.size_2_pt(lst_margin[0], lst_style, dict_label_style)
                    right = StyleManageApi.size_2_pt(lst_margin[1], lst_style, dict_label_style)
                    bottom = StyleManageApi.size_2_pt(lst_margin[2], lst_style, dict_label_style)
                    left = StyleManageApi.size_2_pt(lst_margin[3], lst_style, dict_label_style)

                elif num == 3:
                    top = StyleManageApi.size_2_pt(lst_margin[0], lst_style, dict_label_style)
                    left = right = StyleManageApi.size_2_pt(lst_margin[1], lst_style, dict_label_style)
                    bottom = StyleManageApi.size_2_pt(lst_margin[2], lst_style, dict_label_style)

                elif num == 2:
                    top = bottom = StyleManageApi.size_2_pt(lst_margin[0], lst_style, dict_label_style)
                    left = right = StyleManageApi.size_2_pt(lst_margin[1], lst_style, dict_label_style)

                elif num == 1:
                    top = bottom = left = right = StyleManageApi.size_2_pt(lst_margin[0], lst_style, dict_label_style)

            if "margin-top" in dict_style:
                top = StyleManageApi.size_2_pt(dict_style["margin-top"], lst_style, dict_label_style)

            if "margin-bottom" in dict_style:
                bottom = StyleManageApi.size_2_pt(dict_style["margin-bottom"], lst_style, dict_label_style)

            if "margin-left" in dict_style:
                left = StyleManageApi.size_2_pt(dict_style["margin-left"], lst_style, dict_label_style)

            if "margin-right" in dict_style:
                right = StyleManageApi.size_2_pt(dict_style["margin-right"], lst_style, dict_label_style)

            return top, bottom, left, right

        top, bottom, left, right = _parse_margin(lst_style, dict_label_style)

        if (top is None) or (bottom is None) or (left is None) is (right is None):
            lable_top, lable_bottom, lable_left, lable_right = _parse_margin([dict_label_style, ], dict_label_style)
            if top is None:
                top = lable_top

            if bottom is None:
                bottom = lable_bottom

            if left is None:
                left = lable_left

            if right is None:
                right = lable_right

        return top, bottom, left, right


class DocApi(object):
    """
    生成word api（一般是特殊情况）
    """

    @staticmethod
    def set_list_number_restat(doc, paragraph):
        """
        设置列表重新开始编号（暂不支持多级列表）
        """
        next_num_id = doc.part.numbering_part.element._next_numId
        num = CT_Num.new(next_num_id, 7)
        num.add_lvlOverride(ilvl=0).add_startOverride(1)
        doc.part.numbering_part.element._insert_num(num)

        p = paragraph._p
        pPr = p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), "%s" % next_num_id)
        numPr.append(numId)
        pPr.append(numPr)

    @staticmethod
    def set_row_height(row, height, rule="exact"):
        """
        设置行高
        height:高度,单位:pt
        Returns:

        """
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tr_height = OxmlElement('w:trHeight')
        # pt = UnitChangeApi.px_2_pt(height)
        tr_height.set(qn('w:val'), "%s" % int(height * 20))  # val - 指定行的高度，在二十分之一的点。
        tr_height.set(qn('w:hRule'), rule)
        tr_pr.append(tr_height)

    @staticmethod
    def set_col_width(table, col_idx, width):
        """
        设置行高，由于office和wps要求不一样，所有cell和col都设置
        width:宽度,单位：像素
        Args:
            col_idx: 列索引
            width: 单位是PT
        """
        # inch = UnitChangeApi.px_2_in(width)
        # in_width = shared.Inches(inch)
        for row in table.rows:
            cell = row.cells[col_idx]
            cell.width = shared.Pt(width)

        col = table.columns[col_idx]
        col.width = shared.Pt(width)

    @staticmethod
    def set_table_background_color(table, color):
        """
        设置表格背景色
        """
        # element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "%s%s%s" % (hex(r)[2:], hex(g)[2:], hex(b)[2:])))
        # element = parse_xml(r'<w:shd {} w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), "%s%s%s" % (color_int_2_hex(r, g, b))))
        # < w:shd        w:val = "clear"        w:color = "auto"        w:fill = "FF0000" / >
        element = OxmlElement('w:shd')
        element.set(qn('w:w'), "clear")
        element.set(qn('w:color'), "auto")
        element.set(qn('w:fill'), StyleManageApi.style_color_2_hex(color))
        table._tblPr.append(element)  # 修改单元格颜色

    @staticmethod
    def set_table_padding(table, top, bottom, left, right):
        """
        设置表格padding
        单位pt
        """
        #     < w:tblCellMar >
        #     < w:top        w:w = "57"        w:type = "dxa" / >
        #     < w:left        w:w = "85"        w:type = "dxa" / >
        #     < w:bottom        w:w = "113"        w:type = "dxa" / >
        #     < w:right        w:w = "142"        w:type = "dxa" / >
        # < / w:tblCellMar >
        if top is None:
            top_20_pt = 0
        else:
            top_20_pt = int(round(top * 20))
        element = OxmlElement('w:tblCellMar')
        element_sub = OxmlElement('w:top')
        element_sub.set(qn('w:w'), "%s" % top_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if bottom is None:
            bottom_20_pt = 0
        else:
            bottom_20_pt = int(round(bottom * 20))
        element_sub = OxmlElement('w:bottom')
        element_sub.set(qn('w:w'), "%s" % bottom_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if left is None:
            left_20_pt = 0
        else:
            left_20_pt = int(round(left * 20))
        element_sub = OxmlElement('w:left')
        element_sub.set(qn('w:w'), "%s" % left_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if right is None:
            right_20_pt = 0
        else:
            right_20_pt = int(round(right * 20))
        element_sub = OxmlElement('w:right')
        element_sub.set(qn('w:w'), "%s" % right_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)
        table._tblPr.append(element)

    @staticmethod
    def set_cell_padding(cell, top, bottom, left, right):
        """
        设置单元格padding
        单位pt
        """
        # < w:tcMar >
        # < w:top        w:w = "85"        w:type = "dxa" / >
        # < w:left        w:w = "113"        w:type = "dxa" / >
        # < w:bottom        w:w = "170"        w:type = "dxa" / >
        # < w:right        w:w = "198"        w:type = "dxa" / >
        # < / w:tcMar >
        if top is None:
            top_20_pt = 0
        else:
            top_20_pt = int(round(top * 20))
        element = OxmlElement('w:tcMar')
        element_sub = OxmlElement('w:top')
        element_sub.set(qn('w:w'), "%s" % top_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if bottom is None:
            bottom_20_pt = 0
        else:
            bottom_20_pt = int(round(bottom * 20))
        element_sub = OxmlElement('w:bottom')
        element_sub.set(qn('w:w'), "%s" % bottom_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if left is None:
            left_20_pt = 0
        else:
            left_20_pt = int(round(left * 20))
        element_sub = OxmlElement('w:left')
        element_sub.set(qn('w:w'), "%s" % left_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)

        if right is None:
            right_20_pt = 0
        else:
            right_20_pt = int(round(right * 20))
        element_sub = OxmlElement('w:right')
        element_sub.set(qn('w:w'), "%s" % right_20_pt)
        element_sub.set(qn('w:type'), "dxa")
        element.append(element_sub)
        cell._tc.append(element)

    @staticmethod
    def set_table_left_indent(table, size):
        """
        设置表格整体移动
        # < w:tblInd
        # w:w = "210"
        # w:type = "dxa" / >
        size:单位pt
        """
        if size is not None:
            element = OxmlElement('w:tblInd')
            element.set(qn('w:w'), "%s" % int(round(size * 20.0)))
            element.set(qn('w:type'), "dxa")
            table._tblPr.append(element)

    @staticmethod
    def set_cell_background_color(cell, color):
        """
        设置单元格背景色
        """
        element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), StyleManageApi.style_color_2_hex(color)))
        cell._tc.get_or_add_tcPr().append(element)  # 修改单元格颜色

    @staticmethod
    def set_paragraph_background_color(paragraph, color):
        """
        设置段落背景颜色
        """
        element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), StyleManageApi.style_color_2_hex(color)))
        paragraph._p.get_or_add_pPr().append(element)

    @staticmethod
    def set_text_background_color(run, color):
        """
        设置文字背景颜色
        """
        element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), StyleManageApi.style_color_2_hex(color)))
        run._r.get_or_add_rPr().append(element)

    @staticmethod
    def set_cell_vertical_alingment(cell, alingment):
        """
        设置单元格内容垂直对齐方式
        Args:
            cell:单元格对象
            alingment:center和bottom，默认是居上，不用调用此函数设置
        """
        tc_pr = cell._tc.get_or_add_tcPr()
        element = OxmlElement('w:vAlign')
        element.set(qn('w:val'), "%s" % alingment)
        tc_pr.append(element)

    @staticmethod
    def set_text_alingment(paragraph, alingment):
        """
        设置内容水平对齐方式
        Args:
            paragraph:段落对象
            alingment:center和right和left，默认是居左
        """
        if alingment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alingment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def add_hyperlink(paragraph, text, url, dict_style):
        """
        增加链接对象
        """
        src = url.strip()
        if src.find("http") != 0:
            src = "%s%s" % (CON_BASE_URL, url)
        src = src.strip("(){}\r\n;")
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(src, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        r = paragraph.add_run()
        r._r.append(hyperlink)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
        DocApi.set_text_style(r, [dict_style, ], get_page_style("a"))

        return hyperlink

    @staticmethod
    def set_paragraph_style(paragraph, lst_style, dict_label_style):
        """
        设置段落样式 p标签 pre标签， 主要设置显示位置
        Args:
            paragraph:
            lst_style:
        """

        # paragraph.paragraph_format.first_line_indent = shared.Inches(1) # 段落首行空格
        # paragraph.paragraph_format.tab_stops.add_tab_stop(shared.Inches(1.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader = WD_TAB_LEADER.LINES)
        top, bottom, left, right = StyleManageApi.parse_margin(lst_style, dict_label_style)
        if top not in [0.0, None]:
            paragraph.paragraph_format.space_before = shared.Pt(top)  # 上
        else:
            paragraph.paragraph_format.space_before = shared.Pt(0)

        if bottom not in [0.0, None]:
            paragraph.paragraph_format.space_after = shared.Pt(bottom)  # 下
        else:
            paragraph.paragraph_format.space_after = shared.Pt(0)

        if left not in [0.0, None]:
            paragraph.paragraph_format.left_indent = shared.Pt(left)  # 左
        # else:
        #     paragraph.paragraph_format.left_indent = shared.Pt(0)  # 左,设置0时，列表会左移一部分

        if right is not None:
            paragraph.paragraph_format.right_indent = shared.Pt(right)  # 右
        else:
            paragraph.paragraph_format.right_indent = shared.Pt(0)  # 右

        # 设置行距
        style = StyleManageApi.get_style("line-height", lst_style, dict_label_style)
        if style is not None:
            paragraph.paragraph_format.line_spacing = StyleManageApi.line_spacing_change(style, lst_style,
                                                                                         dict_label_style)  # 行距(没有单位）
        else:
            paragraph.paragraph_format.line_spacing = 1.0

        # 设置文字对齐方式
        style = StyleManageApi.get_style("text-align", lst_style, dict_label_style)
        if style is not None:
            DocApi.set_text_alingment(paragraph, style)
        else:
            DocApi.set_text_alingment(paragraph, "left")

        # 缩进
        style = StyleManageApi.get_style("text-indent", lst_style, dict_label_style)
        if style is not None:
            paragraph.paragraph_format.left_indent = StyleManageApi.parse_size(style, lst_style, dict_label_style)
            pass

        # 设置背景颜色
        style = StyleManageApi.get_style("background-color", lst_style, dict_label_style)
        if style is not None:
            DocApi.set_paragraph_background_color(paragraph, style)

    @staticmethod
    def set_text_style(run, lst_style, dict_label_style, is_en=True):
        """
        设置文字样式
        Args:
            lst_style:
            dict_label_style:
        """
        # 获取字体
        font_name = StyleManageApi.get_style("font-family", lst_style, dict_label_style)
        if font_name is not None:
            font_name = FontApi.change(font_name, is_en)
        else:
            if is_en:
                font_name = CON_DEFAULT_FONT_EN
            else:
                font_name = CON_DEFAULT_FONT

        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)  # 中文加这句设置才会生效

        # 设置文字大小
        size = StyleManageApi.get_style("font-size", lst_style, dict_label_style, CON_DEFAULT_FONT_SIZE)
        run.font.size = StyleManageApi.parse_size(size, lst_style, dict_label_style, "font-size")

        # 设置文字颜色
        color = StyleManageApi.get_style("color", lst_style, dict_label_style)
        if color is not None:
            r, g, b = StyleManageApi.style_color_2_rgb(color)
            run.font.color.rgb = shared.RGBColor(r, g, b)

        # 设置文字背景颜色(只找当前标签和页面样式）
        color = StyleManageApi.get_style("background-color", lst_style[-1:], dict_label_style)
        if color is not None:
            if color == "inherit":
                color = StyleManageApi.get_style("background-color", lst_style[-1:0], dict_label_style)
                if color not in [None, "inherit"]:
                    run.font.highlight_color = StyleManageApi.background_color_change(color)
            else:
                run.font.highlight_color = StyleManageApi.background_color_change(color)

        color = StyleManageApi.get_style("background", lst_style[-1:], dict_label_style)
        if color is not None:
            if color == "inherit":
                color = StyleManageApi.get_style("background", lst_style[-1:0], dict_label_style)
                if color not in [None, "inherit"]:
                    run.font.highlight_color = StyleManageApi.background_color_change(color)
            else:
                run.font.highlight_color = StyleManageApi.background_color_change(color)

        # 只找当前标签和页面样式
        decoration = StyleManageApi.get_style("text-decoration", lst_style[-1:], dict_label_style)
        if decoration is not None:
            # 设置下划线
            if decoration == "underline":
                run.font.underline = True

            # 删除线
            elif decoration == "line-through":
                run.font.strike = True

            elif decoration == "inherit":
                decoration = StyleManageApi.get_style("text-decoration", lst_style[-1:0], dict_label_style)
                # 设置下划线
                if decoration == "underline":
                    run.font.underline = True

                # 删除线
                elif decoration == "line-through":
                    run.font.strike = True

        # 加粗
        strong = StyleManageApi.get_style("font-weight", lst_style, dict_label_style)
        if (strong in ['bold', 'bolder']) or ((strong is not None) and strong.isdigit() and int(strong) >= 700):
            run.font.bold = True

        # 斜体
        italic = StyleManageApi.get_style("italic", lst_style, dict_label_style)
        if italic is not None:
            run.font.italic = True


class Html2Doc(object):
    """
    1、新建页面默认8.5英寸 7772400 / 914400.0 doc.sections[0].page_width
    section = self.doc.sections[0]
    section.page_height
    10058400
    section.page_width
    7772400
    section.left_margin
    1143000
    section.right_margin
    1143000
    2、
    """

    def __init__(self, content):
        self.content = content
        self.doc = Document()
        # self.dict_page_style = {}  # 页面样式，使用html关键字
        # self.dict_paragraph_style = {}  # 段落样式，使用html关键字
        # self.dict_element_style = {}  # 节点样式，使用html关键字
        # self.begin_paragraph = False
        self.paragraph = None
        self.table = None
        self.cell = None
        self.lst_style = []
        self.lst_label = []

    def _write_text(self, text, lst_write, label):
        """
        写文字
        Args:
            text:
        """
        if text is None:
            return

        text = text.replace("\r\n", "\n")
        style = copy.copy(get_page_style(label))
        if style:
            style.update(self.lst_style[-1])

            if (label == "pre") and ("background-color" in style):  # 有背景色，不能再给文字使用
                del style["background-color"]

            self.lst_style[-1] = style

        for run_text, is_en in FontApi.split(text):
            self.run = self.paragraph.add_run(run_text)
            DocApi.set_text_style(self.run, self.lst_style, style, is_en)

        lst_write.append("text")

    def _write_table(self, element, lst_write, dict_para):
        """
        写表格
        """
        # 从table下级tbody开始找tr
        lst_tr = element[0].findall("tr")
        if not lst_tr:
            return

        if element[0].tag == "thead" and len(element) > 1:
            lst_tr.extend(element[1].findall("tr"))

        row_num = len(lst_tr)
        col_num = 0
        tr = lst_tr[0]
        lst_td = tr.findall("td")
        if not lst_td:
            lst_td = tr.findall("th")

        for td in lst_td:
            dict_td_para = dict(td.items())
            if "colspan" in dict_td_para:
                colspan = int(dict_td_para["colspan"])
                col_num += colspan
            else:
                col_num += 1

        table = self.doc.add_table(row_num, col_num, style='TableGrid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        dict_set_width = {}
        lst_set_height = []
        dict_merge = {}
        dict_table_style = self.lst_style[-1]
        css_table_style = get_page_style("table")
        # css_td_style = get_page_style("td")

        for row_idx, tr in enumerate(lst_tr):
            dict_element_para, dict_tr_para = self._parse_element_para(tr)
            self.lst_label.append(tr.tag)
            self.lst_style.append(dict_tr_para)

            # 设置行高(在td上的样式)
            if ("height" in dict_tr_para) and (row_idx not in lst_set_height):
                height = StyleManageApi.size_2_pt(dict_tr_para["height"], self.lst_style)
                doc_row = table.rows[row_idx]
                DocApi.set_row_height(doc_row, height, "atLeast")
                lst_set_height.append(row_idx)

            lst_td = tr.findall("td")
            if not lst_td:
                lst_td = tr.findall("th")

            td_idx = 0
            for col_idx in range(col_num):
                if [row_idx, col_idx] in dict_merge.get(row_idx, []):
                    continue

                td = lst_td[td_idx]
                css_td_style = get_page_style(td.tag)
                dict_element_para, dict_td_para = self._parse_element_para(td)
                if td.tag == "th":
                    dict_td_para["font-weight"] = 'bold'

                self.lst_label.append(td.tag)
                self.lst_style.append(dict_td_para)

                td_idx += 1
                cell = table.cell(row_idx, col_idx)
                self.cell = cell
                top, bottom, left, right = StyleManageApi.parse_padding(self.lst_style, css_td_style)
                DocApi.set_cell_padding(cell, top, bottom, left, right)

                # 设置合并单元格
                if ("colspan" in dict_td_para) or ("rowspan" in dict_td_para):
                    rowspan = int(dict_td_para.get("rowspan", 1))
                    colspan = int(dict_td_para.get("colspan", 1))
                    # if not (rowspan == 1 and colspan == 1):
                    b_cell = table.cell(row_idx + rowspan - 1, col_idx + colspan - 1)
                    cell.merge(b_cell)

                    for i in range(row_idx, row_idx + rowspan):
                        for j in range(col_idx, col_idx + colspan):
                            # 记录每行合并的列
                            dict_merge.setdefault(i, []).append([i, j])

                # 设置列宽
                if "width" in dict_td_para:
                    width_pt = StyleManageApi.size_2_pt(dict_td_para["width"], self.lst_style, css_td_style)
                    if width_pt > dict_set_width.get(col_idx, 0):
                        dict_set_width[col_idx] = width_pt

                # 设置行高(在tr上的样式)
                if ("height" in dict_td_para) and (row_idx not in lst_set_height):
                    height = StyleManageApi.size_2_pt(dict_td_para["height"], self.lst_style)
                    doc_row = table.rows[row_idx]
                    DocApi.set_row_height(doc_row, height, "atLeast")
                    lst_set_height.append(row_idx)

                # 垂直对齐
                if "valign" in dict_td_para:
                    DocApi.set_cell_vertical_alingment(cell, dict_td_para["valign"])

                if "style" in dict_td_para:
                    # 水平对齐
                    dict_style = self._parse_style(dict_td_para["style"])
                    if "text-align" in dict_style:
                        DocApi.set_text_alingment(cell.paragraphs[0], dict_style["text-align"])

                color = StyleManageApi.get_style("background-color", self.lst_style, css_td_style)
                if color is not None:
                    hex_color = StyleManageApi.style_color_2_hex(color)
                    DocApi.set_cell_background_color(cell, "#%s" % hex_color)

                color = StyleManageApi.get_style("background", self.lst_style, css_td_style)
                if color is not None:
                    hex_color = StyleManageApi.style_color_2_hex(color)
                    DocApi.set_cell_background_color(cell, "#%s" % hex_color)

                dict_para["cell"] = True
                dict_para["first_add_paragraph"] = False
                self.paragraph = cell.paragraphs[-1]
                if td.text:  # 内容在td中
                    self._write_text(td.text, lst_write, "td")
                lst_check = []
                self._parse_element(td, lst_check, dict_para)
                if lst_check == ['br', ]:
                    delete_paragraph(self.paragraph)

                dict_para["cell"] = False
                dict_para["first_add_paragraph"] = True
                del self.lst_label[-1]
                del self.lst_style[-1]

            del self.lst_label[-1]
            del self.lst_style[-1]

        # 设置列宽
        # 计算最大显示列宽
        if col_num > 0:
            max_pt = UnitChangeApi.in_2_pt(MAX_CONTENT_WIDTH_IN)
            table_width = max_pt
            if "width" in dict_table_style:
                width_pt = StyleManageApi.size_2_pt(dict_table_style["width"], self.lst_style, css_table_style)
                if math.isnan(width_pt):
                    table_width = max_pt
                elif width_pt > max_pt:
                    table_width = max_pt
                else:
                    table_width = width_pt

            default_width = table_width / float(col_num)
            ratio = table_width / ((col_num - len(dict_set_width)) * default_width + sum(dict_set_width.values()))

            for col_idx in range(col_num):
                width_pt = dict_set_width.get(col_idx, default_width)
                DocApi.set_col_width(table, col_idx, width_pt * ratio)

    def _write_image(self, dict_image, dict_para, dict_style):
        try:
            if 'src' in dict_image:
                src = dict_image['src'].strip()
                self.run = self.paragraph.add_run()
                width = None
                height = None
                dict_label_style = get_page_style("img")
                if "width" in dict_style:
                    width = shared.Pt(StyleManageApi.size_2_pt(dict_style["width"], [], dict_label_style))

                elif dict_para.get("cell", False):
                    width = self.cell.width - shared.Pt(5)  # 表格内的图片，右边要空一点，经验值

                if "height" in dict_style:
                    height = shared.Pt(StyleManageApi.size_2_pt(dict_style["height"], [], dict_label_style))

                if src.find("http") == 0:
                    src = download_image(src)
                else:
                    src = "%s%s" % (CON_BASE_ADDRESS, dict_image['src'])
                    src = os.path.abspath(src)

                # 获取图片的实际大小
                if (width is None) or (height is None):
                    try:
                        im = Image.open(src)
                    except (BaseException,):
                        # 文件不存在，写上名字
                        if "alt" in dict_style:
                            lst_image = []
                            self._write_text(dict_style["alt"], lst_image, "img")
                        return

                    if width is None:
                        # 图片显示宽度自适应
                        max_px = UnitChangeApi.in_2_px(MAX_CONTENT_WIDTH_IN)
                        if im.size[0] > max_px:
                            width = shared.Inches(MAX_CONTENT_WIDTH_IN)
                        else:
                            width = shared.Pt(UnitChangeApi.px_2_pt(im.size[0]))

                    # if height is None:
                    #     height = StyleChangeApi.size_change(im.size[1])

                self.run.add_picture(src, width=width, height=height)

        except (BaseException,):
            logger.error(traceback.format_exc())
            pass

    def _write_pre(self, element, dict_para, lst_style):
        """
        由于word没有html"padding: 16px;"参数，要用table模拟，所以先插入一行一列的表格作为底色
        """

        # DocApi.set_table_background_color(table, "#ff0000")
        dict_label_style = get_page_style("pre")
        margin_top, margin_bottom, margin_left, margin_right = StyleManageApi.parse_margin(lst_style, dict_label_style)
        if margin_top not in [0, None]:
            if margin_top > 8.0:  # 大于8.0
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.line_spacing = \
                    StyleManageApi.line_spacing_change("%spt" % margin_top,
                                                       lst_style, dict_label_style)
                paragraph.paragraph_format.space_after = shared.Pt(0)  # 后段距离

        table = self.doc.add_table(1, 1)

        if margin_bottom not in [0, None]:
            if margin_bottom > 8.0:
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.line_spacing = \
                    StyleManageApi.line_spacing_change("%spt" % margin_bottom,
                                                       lst_style, dict_label_style)
                paragraph.paragraph_format.space_after = shared.Pt(0)  # 后段距离

        top, bottom, left, right = StyleManageApi.parse_padding(lst_style, dict_label_style)
        DocApi.set_table_padding(table, top, bottom, left, right)
        DocApi.set_table_left_indent(table, left + margin_left)  # padding是外扩，所以必须右移, 再加上margin left
        DocApi.set_col_width(table, 0, UnitChangeApi.in_2_pt(MAX_CONTENT_WIDTH_IN) - margin_left - margin_right)

        color = StyleManageApi.get_style("background-color", lst_style, dict_label_style)
        if color is not None:
            hex_color = StyleManageApi.style_color_2_hex(color)
            #     for row in table.rows:
            cell = table.cell(0, 0)
            DocApi.set_cell_background_color(cell, "#%s" % hex_color)

        dict_para["cell"] = True
        self.cell = table.cell(0, 0)
        self.paragraph = self.cell.paragraphs[-1]
        self.paragraph.paragraph_format.left_indent = shared.Pt(0)

        dict_para["cell"] = False

        if "background-color" in lst_style[-1]:
            # 段落背景色和文字背景色关键字一样，不能重复设置
            del lst_style[-1]["background-color"]

        lst_check = []
        self._write_text(element.text, lst_check, element.tag)
        self._parse_element(element, lst_check, dict_para)

    def _del_paragraph(self, dict_para, add_break=None):
        if add_break is None:
            return

        if add_break:
            delete_paragraph(self.run)
        elif dict_para.get('cell', False):
            delete_paragraph(self.cell.paragraphs[-1])
        else:
            delete_paragraph(self.doc.paragraphs[-1])

    def _create_paragraph(self, element, dict_para):
        def compare_upper_level(label):
            if len(self.lst_label) > 1 and (label == self.lst_label[-2]):
                dict_style = self.lst_style[-2]
                if "label_first" not in dict_style:
                    dict_style["label_first"] = True
                else:
                    # 第二次及以后
                    dict_style["label_first"] = False

                return True

            return False

        def check_previou_label(lst_label):
            for label in lst_label:
                if label in self.lst_label:
                    return True

            return False

        def create_paragraph(dict_para, add_break, style=None):
            if add_break:
                self.run = self.paragraph.add_run()
                self.run.add_break()
            else:
                if dict_para.get('cell', False):
                    self.paragraph = self.cell.add_paragraph(style=style)
                else:
                    self.paragraph = self.doc.add_paragraph(style=style)
                    self.paragraph.paragraph_format.space_after = shared.Pt(0)  # 后段距离

        add_break = False
        if element.tag in ["p"]:
            # 如果上级是li, 不处理
            if compare_upper_level("li") or compare_upper_level("td"):
                dict_style = self.lst_style[-2]
                if dict_style["label_first"]:
                    pass
                else:
                    add_break = True
                    create_paragraph(dict_para, add_break)

            elif check_previou_label(["ol", "ul"]):
                add_break = True
                create_paragraph(dict_para, add_break)
            else:
                create_paragraph(dict_para, False)

        elif element.tag in ["pre"]:
            # dict_label_para = {}
            if compare_upper_level("li"):
                dict_style = self.lst_style[-2]
                if dict_style["label_first"]:
                    if "ol" in self.lst_label:
                        create_paragraph(dict_para, add_break=False, style="List Number")
                    elif "ul" in self.lst_label:
                        create_paragraph(dict_para, add_break=False, style="List Bullet")
                else:
                    if "text-indent" not in self.lst_style[-1]:
                        self.lst_style[-1]["text-indent"] = "0.70cm"
                    create_paragraph(dict_para, add_break=False)
            else:
                create_paragraph(dict_para, add_break=False)

        elif element.tag == "ol":
            create_paragraph(dict_para, add_break=False, style="List Number")
            if dict_para["ol_num"] != 1:
                DocApi.set_list_number_restat(self.doc, self.paragraph)

        elif element.tag == "li":
            if compare_upper_level("ol"):
                dict_style = self.lst_style[-2]
                if not dict_style["label_first"]:
                    create_paragraph(dict_para, add_break=False, style="List Number")

            elif "ul" in self.lst_label:
                create_paragraph(dict_para, add_break=False, style="List Bullet")

        elif element.tag in ["span", "strong"]:
            if check_previou_label(["ol", "ul"]) and not check_previou_label(["p", "pre"]):
                add_break = True
                create_paragraph(dict_para, add_break)
            else:
                add_break = None

        elif element.tag in ["br"]:
            if check_previou_label("p"):
                add_break = True
                create_paragraph(dict_para, add_break)
            elif check_previou_label(["ol", "ul", "pre"]):
                add_break = True
                create_paragraph(dict_para, add_break)
            else:
                create_paragraph(dict_para, add_break)

        elif element.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            create_paragraph(dict_para, add_break)

        return add_break

    def _parse_style(self, style):
        """
        Args:

        """
        dict_style = {}

        if style is None:
            return dict_style

        lst_style = style.strip(";").split(";")
        for s_style in lst_style:
            s_style = s_style.strip()
            if s_style == "":
                continue

            #  'inherit !important'
            try:
                key, val = s_style.split(":")
            except (BaseException,):
                continue

            val = val.replace("!important", "").strip()
            if val == "":
                continue

            dict_style[key.strip().lower()] = val

        return dict_style

    def _parse_element_para(self, element):
        dict_style = self._parse_style(element.get("style"))
        dict_element_para = {}
        for key, val in element.items():
            key = key.lower()
            dict_element_para[key] = val

        if ("width" in dict_element_para) and ("width" not in dict_style):
            dict_style["width"] = dict_element_para["width"]

        if ("height" in dict_element_para) and ("height" not in dict_style):
            dict_style["height"] = dict_element_para["height"]

        if "colspan" in dict_element_para:
            dict_style["colspan"] = dict_element_para["colspan"]

        if "rowspan" in dict_element_para:
            dict_style["rowspan"] = dict_element_para["rowspan"]

        if element.tag in ["strong", "b"]:  # 粗体
            dict_style["font-weight"] = 'bold'

        elif element.tag in ["em", "i"]:  # 斜体
            dict_style["italic"] = True

        elif (element.tag == "img") and ("alt" in dict_element_para):
            dict_style["alt"] = dict_element_para["alt"]

        return dict_element_para, dict_style

    def _parse_element(self, root, lst_write, dict_para=None):
        for (i, element) in enumerate(root):
            # print(ET.tostring(element, encoding='utf-8'))
            # if element.text:
            #     print(element.text)

            run_sub = True  # 是否执行子图例
            element.tag = element.tag.lower()
            dict_element_para, dict_style = self._parse_element_para(element)
            self.lst_label.append(element.tag)
            self.lst_style.append(dict_style)

            if element.tag == "pre":  # 增加段落
                ######################## old 不支持padding样式 ######################
                # self._create_paragraph(element, dict_para)
                # # dict_paragraph_style = self._parse_paragraph_style(element.get("style"))
                # # if "background-color" not in dict_style:
                # #     dict_style["background-color"] = "#f8f8f8"  # 富文本css中自带的样式
                #
                # DocApi.set_paragraph_style(self.paragraph, self.lst_style, get_page_style(element.tag))
                # if "background-color" in dict_style:
                #     # 段落背景色和文字背景色关键字一样，不能重复设置
                #     del dict_style["background-color"]
                #
                # lst_check = []
                # self._write_text(element.text, lst_check, element.tag)
                # self._parse_element(element, lst_check, dict_para)
                #
                # if lst_check[-1] == "br":  # 如果为空要删除
                #     delete_paragraph(self.run)
                ######################### old end ##################################
                self._write_pre(element, dict_para, self.lst_style)
                run_sub = False

            elif element.tag == "br":
                self._create_paragraph(element, dict_para)
                lst_write.append(element.tag)

            elif element.tag == "p":
                add_break = self._create_paragraph(element, dict_para)

                # 如果有图片不能固定行高
                if (element.find("img") is not None) and ("line-height" in dict_style):
                    del dict_style["line-height"]

                if add_break:
                    # DocApi.set_paragraph_style(self.run, self.lst_style, get_page_style(element.tag))
                    pass
                else:
                    DocApi.set_paragraph_style(self.paragraph, self.lst_style, get_page_style(element.tag))

                lst_check = []  # 用来判断P标签中是否有内容，如果有内容则要换行
                self._write_text(element.text, lst_check, element.tag)
                self._parse_element(element, lst_check, dict_para)

                if not lst_check:  # 如果为空要删除
                    if add_break:
                        delete_paragraph(self.run)
                    elif dict_para.get('cell', False):
                        delete_paragraph(self.cell.paragraphs[-1])
                    else:
                        delete_paragraph(self.doc.paragraphs[-1])

                elif lst_check[-1] == "br":
                    delete_paragraph(self.run)

                lst_write.extend(lst_check)

                run_sub = False

            elif element.tag in ["span", "strong"]:
                add_break = self._create_paragraph(element, dict_para)
                if element.text is not None:
                    self._write_text(element.text, lst_write, element.tag)
                else:
                    self._del_paragraph(dict_para, add_break)

            elif element.tag in ["em", "i"]:
                self._write_text(element.text, lst_write, element.tag)

            elif element.tag == "img":
                self._write_image(dict_element_para, dict_para, dict_style)
                lst_write.append(element.tag)

            elif element.tag == "table":
                self._write_table(element, lst_write, dict_para)
                lst_write.append(element.tag)
                run_sub = False

            elif element.tag == "ol":
                dict_para["ol_num"] = dict_para.get("ol_num", 0) + 1
                self._create_paragraph(element, dict_para)

            elif element.tag == "ul":
                pass

            elif element.tag == "li":
                self._create_paragraph(element, dict_para)
                self._parse_element(element, lst_write, dict_para)
                lst_write.append(element.tag)
                run_sub = False

            elif element.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                self._create_paragraph(element, dict_para)
                if "line-height" in self.lst_style[-1]:
                    del self.lst_style[-1]["line-height"]  # 不使用行距，word行距没有第一行的上行距，显示效果为不上下居中对齐

                DocApi.set_paragraph_style(self.paragraph, self.lst_style, get_page_style(element.tag))
                # self.lst_style[-1]["font-weight"] = 'bold'  # 默认字体加粗
                if "background-color" in self.lst_style[-1]:  # 有背景色，不能再给文字使用
                    del self.lst_style[-1]["background-color"]

                self._write_text(element.text, lst_write, element.tag)
                self._parse_element(element, lst_write, dict_para)
                run_sub = False

            elif element.tag == "a":
                text = ""
                if element.text is None:
                    # < a
                    # href = "http://www.python.org/dev/peps/pep-0008/" > < span
                    # style = "font-size:13px;font-family:'Segoe UI','sans-serif';color:#0366D6" > pep - 000
                    # 8 < / span > < / a >
                    # <a href="http://zh.whhuiyu.com/share/doc/4130/change/">
                    # <span style="color: rgb(0, 56, 132); text-decoration: underline;"><br /></span></a>
                    child_element = element.getchildren()
                    if child_element:
                        element = child_element[0]
                        dict_tmp, dict_style = self._parse_element_para(element)
                        if element.text is not None:
                            text = element.text
                        elif element.tail is not None:
                            text = element.text
                        dict_style.update(self.lst_style[-1])

                elif element.text is not None:
                    text = element.text

                DocApi.add_hyperlink(self.paragraph, text, dict_element_para["href"], dict_style)
                lst_write.append(element.tag)
            else:
                print(element.tag)

            # lst_write.append(element.tag)
            if run_sub:
                self._parse_element(element, lst_write, dict_para)

            del self.lst_style[-1]
            del self.lst_label[-1]
            if element.tail:
                # <p>2、<a href="http://www.cnblogs.com/aguncn/p/4947092.html"
                # style="margin: 0px; padding: 0px; color: rgb(7, 93, 179);">http://www.cnblogs.com/aguncn/p/4947092.html</a>，
                # django celery redis简单测试</p>
                # 处理如上"，django celery redis简单测试"数据，使用上级样式
                # print(element.tail)
                self._write_text(element.tail, lst_write, element.tag)

    def to_doc(self, path):
        html = etree.HTML(self.content)
        tree = ET.fromstring(etree.tostring(html))  # 使用etrr格式化html
        root = tree.find("body")  # 去掉html, 从body开始
        lst_write = []
        dict_para = {}
        self._parse_element(root, lst_write, dict_para)
        self.doc.save(path)


if __name__ == "__main__":
    # print(FontApi.split(u"在这种情况下，使用123456就是一个很好的选择。"))

    import MySQLdb

    conn = MySQLdb.connect(host="127.0.0.1", user="root", passwd="123456", use_unicode=1, charset='utf8')
    conn.select_db("check_v2")
    cursor = conn.cursor()
    # for id in range(4123, 4143):
    # id = 4135
    # id = 4122
    id = 4142
    print("id:%s" % id)
    sql = "select `text` from t_doc where id=%s;" % id
    cursor.execute(sql)
    clients = cursor.fetchall()
    if len(clients) > 0:
        content = clients[0][0]
        # content = "<b>b1<c>2</c>b2</b>"
        path = "%s.docx" % id
        Html2Doc(content).to_doc(path)
        # from win32com import client as wc

        # word = wc.Dispatch('Word.Application')
        # doc = word.Documents.Open(path)
        # doc.SaveAs('%s.xml' % id, 13)  # 17对应于下表中的pdf文件
        # doc.Close()
        # word.Quit()
        print(path)
    # Doc2Html("aa.docx", "").to_html()
    # import sqlite3
    #
    # conn = sqlite3.connect(r"C:\wangbin\code\hbjw_oa_new\hbjw_oa.db")
    # cursor = conn.cursor()
    # id = 83
    # sql = "select `content` from t_edoc_content where id=%s;" % id
    # cursor.execute(sql)
    # clients = cursor.fetchall()
    # if len(clients) > 0:
    #     html_content = clients[0][0]
    #     Html2Doc(html_content, "%s.docx" % id).to_doc()
