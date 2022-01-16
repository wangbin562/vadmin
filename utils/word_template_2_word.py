# !/usr/bin/python
# -*- coding=utf-8 -*-

import collections
import re

import docx
from PIL import Image
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from utils import conver
from utils import word_api
from vadmin import common
from vadmin import widgets


class Compiler(object):
    def __init__(self, path, param):
        self.path = path
        self.doc = docx.Document(path)
        self.param = param
        self.paragraph = None

    def parse(self):
        """
        解释文件
        """
        section = self.doc.sections[0]
        self.width = conver.twips_2_px(section.page_width)
        self.content_width = conver.twips_2_px(section.page_width - section.left_margin -
                                               section.right_margin)
        for child in self.doc._body._body:
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, self.doc)
                self.parse_paragraph(paragraph)

            elif isinstance(child, CT_Tbl):
                table = Table(child, self.doc)
                o_table = self.parse_table(table)
            else:
                pass

    def parse_paragraph(self, paragraph, width=None):
        self.paragraph = paragraph
        runs = paragraph.runs
        # print(paragraph.text)

        lst_script = re.findall(r"( *{{.*?}} *)", paragraph.text)

        if not lst_script:
            return

        dict_run = collections.OrderedDict()
        idx = 0
        for run in runs:
            dict_run[idx] = run
            idx += len(run.text)

        is_script = False
        is_script_end = False
        script = ""
        run = runs[0]
        begin_idx = 0
        text_len = len(paragraph.text)
        paragraph_text = paragraph.text
        for i in range(text_len + 1):
            if i == text_len:
                if is_script_end:
                    opera_run = self.del_run_text(dict_run, begin_idx, i - 1)
                    script = re.search(r"({{.*}}?)", script).groups()[0]
                    # name = script.strip("{}")
                    # para = self.param.get(name, "")
                    self.write(opera_run, script)
                break

            char = paragraph_text[i]
            if i in dict_run:
                run = dict_run[i]

            if char == "{" and (text_len > i + 1) and paragraph_text[i + 1] == "{":
                is_script = True
                if not script:
                    begin_idx = i
                script += char

            elif char == "}" and (text_len > i + 1) and paragraph_text[i + 1] == "}":
                script += char

            elif char == "}":
                script += char
                is_script_end = True

            elif char == " " and run.underline:
                if not script:
                    begin_idx = i
                script += char

            elif is_script:
                if is_script_end:
                    opera_run = self.del_run_text(dict_run, begin_idx, i - 1)
                    script = re.search(r"({{.*?}})", script).groups()[0]
                    # name = script.strip("{}")
                    # para = self.param.get(name, "")
                    self.write(opera_run, script)
                    is_script = False
                    is_script_end = False
                    script = ""
                else:
                    script += char

            else:
                script = ""

    def parse_table(self, table):
        merged_cells = word_api.TableApi.get_merged_cells(table)
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                b_merged, b_range, row_span, col_span = \
                    word_api.TableApi.get_merged_info(merged_cells, row_idx, col_idx)

                if b_range:
                    continue

                for paragraph in cell.paragraphs:
                    o_table_sub = self.get_table(paragraph, cell)
                    if o_table_sub:
                        pass

                    else:
                        self.parse_paragraph(paragraph)

    def del_run_text(self, dict_run, begin_idx, end_idx):
        run_count = len(dict_run)
        lst_idx = list(dict_run.keys())
        i = 0
        prefix = ""
        while i < run_count:
            idx1 = lst_idx[i]
            run = dict_run[idx1]

            if (i + 1) < run_count:
                idx2 = lst_idx[i + 1]
                if idx2 < begin_idx:
                    i += 1
                    continue

                if idx1 >= begin_idx and end_idx < idx2:
                    if prefix:
                        run.text = prefix + run.text
                    return run

                elif idx2 > end_idx:
                    run.text = prefix + run.text
                    return run

                elif idx1 >= begin_idx:
                    idx = run.text.find("{")
                    if idx > -1:
                        if run.underline and run.text[0:idx].strip() == "":
                            prefix += run.text
                            run.text = ""
                        else:
                            prefix += run.text[idx:]
                            run.text = run.text[0:idx]
                        # elif prefix.strip() == "":
                        #     prefix = run.text
                        #     run.text = ""
                        # else:

                    elif prefix:
                        prefix += run.text
                        run.text = ""

                    elif run.underline:
                        s = run.text.rstrip(" ")
                        prefix = run.text[len(s):]

                elif prefix:
                    prefix += run.text
                    run.text = ""

                i += 1
            else:
                if prefix:
                    run.text = prefix + run.text

                return run

    def _write(self, run, dict_widget, script):
        run.text = run.text.replace(script, "")
        widget_type = dict_widget.get("type", None)
        if widget_type == "image":
            # width = height = None
            # if "width" in dict_widget:
            #     width = conver.px_2_twips(dict_widget["width"])
            #
            # if "height" in dict_widget:
            #     height = conver.px_2_twips(dict_widget["height"])
            #
            # paragraph = self.paragraph.insert_paragraph_before()
            # run = paragraph.add_run()
            # picture = run.add_picture(dict_widget.get("href", ""), width, height)
            self.write_image(run, dict_widget)
        elif widget_type == "text":
            self.write_text(run, dict_widget)
        elif widget_type == "lite_table":
            self.write_table(run, dict_widget)

    def write(self, run, script):
        widget_script = script.strip("{}")
        para = self.param.get(widget_script, None)
        if isinstance(para, str) and para.find("{") == 0:
            try:
                dict_widget = eval(para)
                self._write(run, dict_widget, script)
            except (BaseException,):
                if para is None:
                    run.text = run.text.replace(script, " " * len(script))
                else:
                    run.text = run.text.replace(script, str(para))

        elif isinstance(para, widgets.Text):
            self.write_text(run, para.render())
        elif isinstance(para, widgets.Image):
            self.write_image(run, para.render())
        elif isinstance(para, widgets.LiteTable):
            self.write_table(run, para.render())
        elif script.find("{") == 0:
            if isinstance(para, int):
                run.text = run.text.replace(script, str(para))
            elif widget_script.lower().find("widget(") == 0:
                widget_script = widget_script.replace('”', '"').replace("，", ",")
                widget_script = "widgets.W%s" % widget_script[1:]
                o_widget = eval(widget_script)
                # if o_widget.get_attr_value("type", None):
                #     o_widget.type = o_widget.type.strip()
                name = o_widget.get_attr_value("name", None)
                text = self.param.get(name, "") or ""
                run.text = run.text.replace(script, text)
            else:
                try:
                    dict_widget = eval(script)
                    self._write(run, dict_widget, script)
                except (BaseException,):
                    if para is None:
                        run.text = run.text.replace(script, " " * len(script))
                    else:
                        run.text = run.text.replace(script, str(para))

    def write_text(self, run, dict_widget):
        """
        写文字
        """
        text = dict_widget.get("text", self.param.get(dict_widget.get("name", None), None)) or ""

        run.text = text.replace("\r\n", "\n")
        word_api.TextApi.set_style(run, dict_widget)
        new_run = run
        # for run_text, is_en in word_api.TextApi.split(text):
        #     # 插入text
        #     new_run = word_api.TextApi.insert(self.paragraph, run, run_text)
        #     # run._element.addprevious(new_run_element)
        #     # self.run = self.paragraph.add_run(run_text)
        #     word_api.TextApi.set_style(new_run, dict_widget, is_en)

        if "bg" in dict_widget and "image" in dict_widget["bg"]:
            self.write_image(new_run, {"type": "image", "href": dict_widget["bg"]["image"],
                                       "float": {}, "width": dict_widget.get("width", None),
                                       "height": dict_widget.get("height", None)})

    def write_image(self, run, dict_image):
        # run.text = ""
        href = dict_image.get("href", None) or dict_image.get("active_href", None)
        if href:
            path = common.get_export_path()
            try:
                word_api.ImageApi.download(href, path)
            except (BaseException,):
                path = href

            width = dict_image.get("width", None)
            height = dict_image.get("height", None)

            if width is not None:
                width = conver.px_2_twips(width)

            if height is not None:
                height = conver.px_2_twips(height)

            if (width is None) or (height is None):
                im = Image.open(path)
                if width is None:
                    if im.size[0] > self.content_width:
                        width = conver.px_2_twips(self.content_width)
                    else:
                        width = conver.px_2_twips(im.size[0])

            word_api.ImageApi.add(run=run, path=path, width=width, height=height,
                                  float=dict_image.get("float", None))
            # run.add_picture(path, width=width, height=height)
            # run.add_inline_picture(path, width=width, height=height)

    def write_table(self, run, dict_table):
        """
        写入table对象
        :param dict_table:vadmin lite_table对象数据
        :return:
        """
        # run.text = ""
        data = dict_table.get("data", [])
        if data:
            row_num = len(data)
            col_num = len(data[0])
        else:
            row_num = 0
            col_num = 0

        table = self.doc.add_table(row_num, col_num)
        self.paragraph._p.addnext(table._element)
        style_name = word_api.TableApi.add_style(self.doc)
        table.style = style_name

        # 设置行高
        word_api.TableApi.set_row_height(table, dict_table)

        # 设置列宽
        word_api.TableApi.set_col_width(table, dict_table)

        # word_api.TableApi.set_left_indent(table, dict_table)
        word_api.TableApi.set_horizontal(table, dict_table)
        word_api.TableApi.set_vertical(table, dict_table)
        # 合并单元格
        word_api.TableApi.merge(table, dict_table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        word_api.TableApi.set_padding(table, dict_table)

        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                if not isinstance(cell_data, list):
                    cell_data = [cell_data]

                paragraph = cell.paragraphs[-1]
                for item in cell_data:
                    if isinstance(item, (float, int)):
                        paragraph.add_run(str(item))

                    elif isinstance(item, str):
                        for run_text, is_en in word_api.TextApi.split(item):
                            run = paragraph.add_run(run_text)
                            # word_api.TextApi.set_style(run, o_text, is_en)

                    elif isinstance(item, widgets.Text):
                        for run_text, is_en in word_api.TextApi.split(item.text):
                            run = paragraph.add_run(run_text)
                            word_api.TextApi.set_style(run, item, is_en)

                    elif isinstance(item, widgets.Image):
                        run = paragraph.add_run()
                        self.write_image(run, item)
                    else:  # 不支持嵌套表格
                        print(str(item))

    def get_table(self, paragraph, cell):
        for table in cell.tables:
            if paragraph._p.getprevious() == table._element:
                return table

    def to_doc(self, path):
        self.doc.save(path)


def generate(template_path, file_path, param):
    """
    template_file：模板文件(只支持*.docx)
    file_path：保存文件路径(只支持*.docx)
    param:参数
    """
    obj = Compiler(template_path, param)
    obj.parse()
    obj.to_doc(file_path)
