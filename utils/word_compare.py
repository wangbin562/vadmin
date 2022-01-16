# !/usr/bin/python
# -*- coding=utf-8 -*-
import codecs
import difflib

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def parse(path):
    doc = docx.Document(path)
    lst_text = []
    number_idx = 0
    for child in doc._body._body:
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)

            text = ""
            lst_num = paragraph._p.xpath('./w:pPr/w:numPr/w:numId')
            if lst_num:
                num_id = lst_num[0].val
                ilvl = paragraph._p.xpath('./w:pPr/w:numPr/w:ilvl')[0].val
                number_idx += 1

                if num_id in [1, 2]:
                    dict_number = {'0': '零', '1': '一', '2': '二', '3': '三', '4': '四',
                                   '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
                    s_number = str(number_idx)
                    lst_number = []
                    for num in s_number:
                        lst_number.append(dict_number[num])
                    text = "%s、" % "".join(lst_number)

                elif num_id == 3:
                    text = "%s. " % number_idx

            else:
                number_idx = 0

            text += paragraph.text.replace("\xa0", "").replace("\u3000", "").replace(" ", "")
            lst_text.append(text)

        elif isinstance(child, CT_Tbl):
            text = ""
            table = Table(child, doc)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for i, paragraph in enumerate(cell.paragraphs):
                        o_table_sub = None
                        for table in cell.tables:
                            if paragraph._p.getprevious() == table._element:
                                o_table_sub = table
                                break

                        if o_table_sub:
                            pass  # 暂时不支持子表

                        else:
                            text += paragraph.text
            lst_text.append(text.replace("\xa0", "").replace("\u3000", "").replace(" ", ""))
        else:
            # print(child)
            pass
    return lst_text


def compare(old_path, new_path):
    """
    各个差异符号表示含义
    '-':包含在第一个序列行中，不包含在第二个序列行中
    '+':包含在第二个序列行中，不包含在第一个序列行中
    '':两个序列行一致
    '?':标志两个序列行存在增量差异
    '^':标志出两个序列存在的差异字符
    """
    lst_text1 = parse(old_path)
    lst_text2 = parse(new_path)
    # print(lst_text1)
    # print(lst_text2)
    d = difflib.ndiff(lst_text1, lst_text2)
    # return list(difflib.restore(list(d), 1)), list(difflib.restore(list(d), 2))
    lst_result = []
    for c in list(d):
        if c.find("+") == 0:
            lst_result.append(c[2:])

    return lst_result


_legend = """
    <table class="diff" summary="Legends">
        <tr> <th colspan="2"> 说明 </th> </tr>
        <tr> <td> <table border="" summary="Colors">
                      <tr><th> 颜色说明 </th> </tr>
                      <tr><td class="diff_add"> 增加 </td></tr>
                      <tr><td class="diff_chg"> 修改 </td> </tr>
                      <tr><td class="diff_sub"> 删除 </td> </tr>
                  </table></td>
             <td> <table border="" summary="Links">
                      <tr><th colspan="2"> 链接说明 </th> </tr>
                      <tr><td>(f)上一修改</td> </tr>
                      <tr><td>(n)下一修改</td> </tr>
                      <tr><td>(t)返回顶部</td> </tr>
                  </table></td> </tr>
    </table>"""


def compare_output_html(old_path, new_path, output_path):
    d = difflib.HtmlDiff()
    lst_text1 = parse(old_path)
    lst_text2 = parse(new_path)
    d._legend = _legend
    result = d.make_file(lst_text1, lst_text2)
    # print(diff)
    result = result.replace(' nowrap="nowrap"', "")
    fh = codecs.open(output_path, "w", encoding='utf-8')
    fh.write(result)
    fh.close()


if __name__ == "__main__":
    result = compare_output_html(r"C:\wangbin\code_sync\code_sync\vadmin_v2\vadmin\templates\刘某某2.docx",
                                 r"C:\wangbin\code_sync\code_sync\vadmin_v2\vadmin\templates\1.docx",
                                 r"C:\wangbin\code_sync\code_sync\vadmin_v2\vadmin\templates\1.html", )
    print(result)
    # import os
    # from win32com import client

    # word = client.Dispatch('Word.Application')

    # def doc_to_docx(path, save_path):
    #     if os.path.splitext(path)[1] == ".doc":
    #         doc = word.Documents.Open(path)  # 目标路径下的文件
    #         doc.SaveAs(os.path.splitext(path)[0] + ".docx", 16)  # 转化后路径下的文件
    #         doc.Close()
    #
    #
    # dir = r"C:\wangbin\code_sync\code_sync\vadmin_v2\vadmin\templates\doc"
    # save_dir = r"C:\wangbin\code_sync\code_sync\vadmin_v2\vadmin\templates\docx"
    # for sub in os.listdir(dir):
    #     file_path = os.path.join(dir, sub)
    #     if os.path.isfile(file_path):
    #         doc_to_docx(file_path, save_dir)
    #
    # word.Quit()
