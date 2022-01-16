# !/usr/bin/python
# -*- coding=utf-8 -*-
"""
"""
import os
import re
import time
import traceback
import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import RGBColor
from docx.shared import Twips
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
from utils import conver
from utils import word_api
from vadmin import step
from vadmin import widgets
from vadmin import common
from vadmin import admin_fields

MIN_HEIGHT = 16


class Compiler(object):
    def __init__(self, doc_path, param=None, show_width=None, is_pdf=False,
                 is_word_widget=False, style_param=None):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        if param is None:
            self.param = {}
        else:
            self.param = param
        self.path, name = os.path.split(doc_path)
        file_name = os.path.splitext(name)[0]
        self.file_name = common.format_file_name(file_name)

        self.lst_paragraph = []
        self.show_width = show_width  # 显示的宽度
        self.ratio = 1
        self.col_width_ratio = 1

        self.page_width = None
        self.page_height = None
        self.content_width = None
        self.content_height = None
        self.top_margin = 0
        self.right_margin = 0
        self.bottom_margin = 0
        self.left_margin = 0
        self.header_distance = 0
        self.footer_distance = 0

        self.margin = []
        self.lst_widget = []
        self.prev_text_run = None
        self.dict_abstract = {}
        self.paragraph = None
        self.prev_num_id = None
        self.prev_level = None
        self.lst_hyperlink = []
        self.lst_anchor_name = []
        self.is_pdf = is_pdf
        self.is_word_widget = is_word_widget  # 是否用word组件显示
        self.style_param = style_param or {}  # 固定样式参数 （font_size, font_color)
        self.image_idx = 0

        self.dict_style = {}
        for style in self.doc.styles.element.style_lst:
            self.dict_style[style.styleId] = style

    def parse(self):
        """
        解释文件
        """
        section = self.doc.sections[0]
        self.page_width = conver.twips_2_px(section.page_width)
        self.page_height = conver.twips_2_px(section.page_height)
        self.header_distance = conver.twips_2_px(section.header_distance)
        self.footer_distance = conver.twips_2_px(section.footer_distance)

        if (self.show_width is not None) and (self.show_width < self.page_width):
            self.ratio = self.show_width / self.page_width  # 显示宽度小于实际宽度，所有内容要同比缩小

        # 文档网络类型（只支持无网络和只指定行网络）
        self.line_type = section._sectPr.find(qn("w:docGrid")).get(qn("w:type"))
        line_pitch = Twips(int(section._sectPr.find(qn("w:docGrid")).get(qn("w:linePitch"))))
        self.line_pitch = conver.twips_2_px(line_pitch)

        self.content_width = self.show_width or conver.twips_2_px(section.page_width - section.left_margin -
                                                                  section.right_margin)
        self.content_height = conver.twips_2_px(section.page_height - section.top_margin -
                                                section.bottom_margin)
        try:
            self.parse_paragraph_number()
        except (BaseException,):
            print(traceback.format_exc())
            pass

        self.top_margin = conver.twips_2_px(section.top_margin)
        self.right_margin = conver.twips_2_px(section.right_margin)
        self.bottom_margin = conver.twips_2_px(section.bottom_margin)
        self.left_margin = conver.twips_2_px(section.left_margin)

        for child in self.doc._body._body:
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, self.doc)
                o_paragraph = self.parse_paragraph(paragraph, '100%')
                self.parse_tabs(paragraph, o_paragraph)
                self.lst_paragraph.append(o_paragraph)

            elif isinstance(child, CT_Tbl):
                table = Table(child, self.doc)
                o_table = self.parse_table(table)
                self.lst_paragraph.append(o_table)
                # self.lst_paragraph.append(widgets.Paragraph())
            else:
                print(type(child))
                pass

        return self.lst_paragraph

    def parse_paragraph_number(self):
        """解析段落符号"""
        numbering = self.doc.part.numbering_part.numbering_definitions._numbering
        dict_abstract_id = {}
        lst_num = numbering.xpath("./w:num")
        for num in lst_num:
            num_id = num.get(qn("w:numId"))
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            if abstract_id in dict_abstract_id:
                continue
            dict_abstract_id[abstract_id] = num_id

        self.dict_abstract = {}
        lst_abstract = numbering.xpath("./w:abstractNum")
        for i, abstract in enumerate(lst_abstract):
            abstract_id = abstract.get(qn("w:abstractNumId"))
            lst_symbol = []
            lst_lvl = abstract.findall(qn("w:lvl"))
            for lvl in lst_lvl:
                type = lvl.find(qn("w:numFmt")).get(qn("w:val"))
                level = int(lvl.get(qn("w:ilvl")))
                left_indent = 0
                try:
                    align = lvl.find(qn("w:lvlJc")).get(qn("w:val"))
                    if align == "left":
                        ind = lvl.find(qn("w:pPr")).find(qn("w:ind"))
                        left = ind.get(qn("w:left"))
                        # firstLine = ind.get(qn("w:firstLine"))
                        # if firstLine is None:
                        hanging = ind.get(qn("w:hanging")) or 0
                        left_indent = conver.twips_2_px(Twips(int(left) - int(hanging)))
                        if left_indent < 0:
                            left_indent = 0
                except (BaseException,):
                    left_indent = 0

                try:
                    format = lvl.find(qn("w:lvlText")).get(qn("w:val"))
                except (BaseException,):
                    format = "%s"

                symbol = {"id": "%s-%s" % (type, format), "type": type, "abstract_id": abstract_id,
                          "format": format, "level": level, "left_indent": left_indent}

                try:
                    start = lvl.find(qn("w:start")).get(qn("w:val"))
                    symbol["start"] = int(start)
                except (BaseException,):
                    pass

                symbol = conver.paragraph_bullet_word_2_html(symbol)
                lst_symbol.append(symbol)

            num_id = dict_abstract_id[abstract_id]
            self.dict_abstract[num_id] = lst_symbol

    def parse_paragraph(self, paragraph, width):
        self.paragraph = paragraph
        o_paragraph = widgets.Paragraph(width=width)
        runs = paragraph.runs
        print(paragraph.text)

        # if "以树型结构直观展现元数据上下级关系" in paragraph.text:
        #     print('a')
        #     pass

        line_spacing_rule, line_spacing = self.get_line_spacing(paragraph)
        o_paragraph.line_spacing_rule, o_paragraph.line_spacing = line_spacing_rule, line_spacing
        space_before = self.get_space_before(paragraph)
        space_after = self.get_space_after(paragraph)
        o_paragraph.space_before = round(space_before * self.ratio, 2)
        o_paragraph.space_after = round(space_after * self.ratio, 2)

        # 对齐
        if paragraph.alignment == WD_TAB_ALIGNMENT.CENTER:
            horizontal = "center"
        elif paragraph.alignment == WD_TAB_ALIGNMENT.RIGHT:
            horizontal = "right"
        elif paragraph.style.paragraph_format.alignment == WD_TAB_ALIGNMENT.CENTER:
            horizontal = "center"
        elif paragraph.style.paragraph_format.alignment == WD_TAB_ALIGNMENT.RIGHT:
            horizontal = "right"
        else:
            horizontal = "left"

        o_paragraph.horizontal = horizontal

        left_indent = self.get_left_indent(paragraph)
        right_indent = self.get_right_indent(paragraph)
        first_line_indent = self.get_first_line_indent(paragraph)

        if first_line_indent < 0:  # first_line_indent等于负数
            first_line_indent = 0

        # if first_line_indent:  # first_line_indent和left_indent同时有效
        o_paragraph.first_line_indent = round(first_line_indent * self.ratio, 2)

        # if left_indent:
        o_paragraph.left_indent = round(left_indent * self.ratio, 2)

        # if right_indent:
        o_paragraph.right_indent = round(right_indent * self.ratio, 2)

        font_size = self.get_font_size(paragraph=paragraph) or 14
        font = {"size": font_size}
        family = self.get_font_family(paragraph=paragraph)
        font["family"] = family
        font.update(self.get_font_style(paragraph=paragraph))
        bg = self.get_font_bg(paragraph=paragraph)

        # try:
        #     snap_to_grid = paragraph._p.find(qn("w:pPr")).find(qn("w:snapToGrid")).val
        # except (BaseException,):
        #     snap_to_grid = True

        # o_paragraph.snap_to_grid = snap_to_grid

        num_id, level = self.get_number(paragraph)
        if level is not None:
            paragraph_symbol = None
            if (num_id is not None) and (str(num_id) in self.dict_abstract):
                paragraph_symbol = self.dict_abstract[str(num_id)]
                if (self.prev_num_id == num_id) and (level > self.prev_level):
                    paragraph_symbol[level]["number"] = 0  # 重现开始计数

                self.prev_num_id = num_id
                self.prev_level = level

            elif (num_id is None) and (self.prev_num_id is not None):
                num_id = self.prev_num_id
                paragraph_symbol = self.dict_abstract[str(num_id)]
            # else:
            #     paragraph_symbol = list(self.dict_abstract.values())[0]

            if self.is_word_widget and paragraph_symbol:
                o_paragraph.paragraph_abstract_id = paragraph_symbol[0]["abstract_id"]
                o_paragraph.paragraph_level = level

            elif paragraph_symbol:
                text = conver.paragraph_number_2_text(paragraph_symbol, level)
                width = paragraph_symbol[level]["left_indent"]
                # if not width:
                #     width = round(word_api.get_text_width(text, font_size) * self.ratio, 2)
                # o_paragraph.left_indent += paragraph_symbol[level]["left_indent"]

                if o_paragraph.first_line_indent != 0 and o_paragraph.first_line_indent < width:
                    o_paragraph.first_line_indent = width
                elif o_paragraph.left_indent != 0 and o_paragraph.left_indent < width:
                    o_paragraph.left_indent = width
                # elif (o_paragraph.first_line_indent + o_paragraph.left_indent) > width:
                #     v = width - o_paragraph.first_line_indent
                #     o_paragraph.first_line_indent = 0
                #     o_paragraph.left_indent = o_paragraph.left_indent - v

                o_widget = widgets.Text(text=text, font=font, bg=bg)
                o_paragraph.append(o_widget)

        o_paragraph.font = font
        o_paragraph.step = self.get_hyperlink(paragraph)

        # 锚点名称
        lst_node = paragraph._element.xpath("w:bookmarkStart")
        for node in lst_node:
            name = node.get(qn("w:name"))
            if name in self.lst_hyperlink:
                o_paragraph.name = name
                self.lst_anchor_name.append(name)
                break

        if not runs:
            return o_paragraph

        lst_script = re.findall(r"( *{{.*?}} *)", paragraph.text)
        if not lst_script:
            for run in runs:
                if run.text:
                    o_widget = self.parse_run(run, run.text, font_size, family)
                else:
                    o_widget = self.parse_image(run)
                    if o_widget and o_widget.get("float"):
                        if not o_paragraph.get("name"):
                            o_paragraph.name = admin_fields.create_uuid_16()

                        o_widget.float["related"] = o_paragraph.name

                if not o_widget:
                    continue

                o_paragraph.append(o_widget)
            return o_paragraph

        dict_run = {}
        idx = 0
        for run in runs:
            dict_run[idx] = run
            idx += len(run.text)

        is_script = False
        is_script_end = False
        script = ""
        text = ""
        run = runs[0]
        text_len = len(paragraph.text)
        for i in range(text_len + 1):
            if (not is_script) and (i in dict_run) and text:
                # text = text.replace("\t", "    ")
                o_widget = self.parse_run(run, text, font_size, family)
                o_paragraph.append(o_widget)
                text = ""

            if i == text_len:
                if text:
                    o_widget = self.parse_run(run, text, font_size, family)
                    o_paragraph.append(o_widget)

                if is_script_end:
                    if run.font.size:
                        widget_width = int(len(script) / 2.0) * conver.twips_2_px(run.font.size)
                    else:
                        widget_width = int(len(script) / 2.0) * font_size

                    name = re.search(r"{{(.*?)}}", script).groups()[0]
                    o_widget = self.make_widget(name, widget_width)
                    o_paragraph.append(o_widget)
                break

            char = paragraph.text[i]
            if i in dict_run:
                run = dict_run[i]
                # if is_script:

            if char == "{" and (text_len > i + 1) and paragraph.text[i + 1] == "{":
                is_script = True
                script += char

            elif char == "}" and (text_len > i + 1) and paragraph.text[i + 1] == "}":
                script += char
            elif char == "}":
                script += char
                is_script_end = True

            elif char == " " and run.underline:
                script += char

            elif is_script:
                if is_script_end:
                    if text:
                        o_widget = self.parse_run(run, text, font_size, family)
                        o_paragraph.append(o_widget)
                        text = ""

                    if run.font.size:
                        widget_width = int(len(script) / 2.0) * conver.twips_2_px(run.font.size)
                    else:
                        widget_width = int(len(script) / 2.0) * font_size

                    name = re.search(r"{{(.*?)}}", script).groups()[0]
                    o_widget = self.make_widget(name, widget_width)
                    o_paragraph.append(o_widget)
                    is_script = False
                    is_script_end = False
                    script = ""
                    text += char
                else:
                    script += char

            else:
                script = ""
                text += char

        return o_paragraph

    def parse_run(self, run, text=None, font_size=None, family=None):
        text = text or run.text
        text = text.replace("\n", "")
        # text = text.replace("\t", "    ")
        if not text:
            return

        font = {}
        o_text = widgets.Text(text=text, inline=True)

        rId, width, height = self.get_bg_image(run)
        if rId:
            image_part = self.doc.part.related_parts[rId]
            name = image_part.partname

            try:
                from django.conf import settings
                path = os.path.abspath(os.path.join(settings.MEDIA_ROOT, "images", name.strip("/\\")))
                url = os.path.join(settings.MEDIA_URL, "images", name.strip("/\\"))
            except (BaseException,):
                path = os.path.join(os.getcwd(), "images", name.strip("/\\"))
                url = path

            tmp_path = os.path.split(path)[0]
            if not os.path.exists(tmp_path):
                os.makedirs(tmp_path)

            with open(path, 'wb') as fp:
                fp.write(image_part._blob)

            o_text.set("bg", {"image": url})
            o_text.set("height", height)

        if run.font.hidden:
            o_text.hide = 1

        # 字体
        font["size"] = self.get_font_size(None, run) or font_size
        font["family"] = self.get_font_family(None, run) or family
        font.update(self.get_font_style(None, run))
        bg = self.get_font_bg(None, run)

        # if o_text.text:
        #     o_text.text = "&thinsp;".join(re.findall("emsp;+|[-=（）()0-9a-zA-Z().]+|[^-=（）()0-9a-zA-Z().]+",
        #                                              o_text.text, re.S))
        #     RE_LOWERCASE = "emsp;+|[0-9a-zA-Z().,`~!@#$%^&*()_+-='/?<>\\|\{}\[\]\" ]+"
        #     if re.findall(RE_LOWERCASE, o_text.text, re.S):
        #         o_text.text = o_text.text.replace("\xa0", "&ensp;").replace(" ", "&ensp;") \
        #             .replace("\u3000", "&ensp;")
        #     else:
        #         o_text.text = o_text.text.replace("\xa0", "&ensp;&ensp;").replace(" ", "&ensp;") \
        #             .replace("\u3000", "&ensp;&ensp;")

        if font:
            o_text.font = font
        if bg:
            o_text.bg = bg
        return o_text

    def parse_image(self, run):
        rId = self.get_image_id(run)
        if not rId:
            return

        image_part = self.doc.part.related_parts[rId]
        if not os.path.exists(self.path):
            os.makedirs(self.path)

        suffix = os.path.splitext(image_part.partname)[1].lower()

        self.image_idx += 1
        image_path = os.path.join(self.path, "%s-%s%s" % (self.file_name, self.image_idx, suffix))
        image_data = image_part._blob
        with open(image_path, 'wb') as fp:
            fp.write(image_data)

        image = Image.open(image_path)
        if suffix not in [".png", ".jpg", ".jpeg"]:
            try:
                image_path = os.path.join(self.path, "%s-%s.png" % (self.file_name, self.image_idx))
                image.save(image_path)
            except (BaseException,):
                image_path = os.path.join(self.path, "%s-%s%s" % (self.file_name, self.image_idx, suffix))

        float = {}
        lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:positionH/wp:posOffset')
        if lst_node:
            node = lst_node[0]
            float["left"] = conver.twips_2_px(int(node.text))

        lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:positionV/wp:posOffset')
        if lst_node:
            node = lst_node[0]
            float["top"] = conver.twips_2_px(int(node.text))

        margin_top = 0
        margin_left = 0
        lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:wrapNone')
        if lst_node:  # 悬浮
            lst_node = run._element.xpath('./w:drawing/wp:anchor')
            value = lst_node[0].get("behindDoc")  # 1:在文字之上 0:在文字之下
        elif float:
            # 四周型环绕 <wp:wrapSquare wrapText="bothSides"/>
            # 精密型环绕 <wp:wrapTight wrapText="bothSides">
            # 穿越型环绕 <wp:wrapThrough wrapText="bothSides">
            # 上下型环绕 <wp:wrapTopAndBottom/>
            margin_top = float["top"]
            margin_left = float["left"]
            float = {}

        width, height = self.get_image_size(run, rId)
        if float:
            o_image = widgets.Image(href=common.get_download_url(image_path), float=float)
        else:
            o_image = widgets.Image(href=common.get_download_url(image_path),
                                    margin_top=margin_top, margin_left=margin_left)

        if width:
            o_image.width = round(width * self.ratio, 2)
            o_image.height = round(height * self.ratio, 2)

        return o_image

    def parse_table(self, table):
        data = list()
        min_row_height = dict()
        cell_horizontal = dict()
        cell_vertical = dict()
        cell_style = dict()
        merged_cells = self.get_table_merged_cells(table)
        horizontal = "left"
        lst_node = table._element.xpath("w:tblPr/w:jc")
        if lst_node:
            node = lst_node[0]
            horizontal = node.val
            if horizontal == WD_TABLE_ALIGNMENT.LEFT:
                horizontal = "left"
            elif horizontal == WD_TABLE_ALIGNMENT.CENTER:
                horizontal = "center"
            elif horizontal == WD_TABLE_ALIGNMENT.RIGHT:
                horizontal = "right"

        for row_idx, row in enumerate(table.rows):
            row_height = self.get_table_row_height(row)
            if row_height:
                min_row_height[row_idx] = row_height

            row_data = []
            for col_idx, cell in enumerate(row.cells):
                b_merged, b_range, row_span, col_span = \
                    self.get_table_merged_info(merged_cells, row_idx, col_idx)

                # 获取表格文字对齐（水平、垂直）
                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                    vertical2 = "center"
                elif cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
                    vertical2 = "bottom"
                else:
                    vertical2 = "top"

                if vertical2 != "top":
                    cell_vertical["%s-%s" % (row_idx, col_idx)] = vertical2

                horizontal2 = self.get_table_cell_horizontal(cell)
                if horizontal2 != "left":
                    cell_horizontal["%s-%s" % (row_idx, col_idx)] = horizontal2

                background_color = self.get_table_cell_background_color(cell)
                if background_color:
                    cell_style["%s-%s" % (row_idx, col_idx)] = {"bg": {"color": background_color}}

                if b_range:
                    row_data.append(None)
                    continue

                # cell_data = widgets.Panel(width="100%", height="100%", vertical=vertical, horizontal=horizontal)
                cell_data = []
                for paragraph in cell.paragraphs:
                    o_table_sub = self.get_table(paragraph, cell)
                    if o_table_sub:
                        pass

                    else:
                        o_panel = self.parse_paragraph(paragraph, "100%")
                        cell_data.append(o_panel)

                row_data.append(cell_data)

            data.append(row_data)

        col_width = dict()
        width = 0
        for col_idx, col in enumerate(table.columns):
            width += col.width

        width = conver.twips_2_px(width)
        if width > self.content_width:  # 表格不能大于内容宽度
            ratio = self.content_width / width
            for col_idx, col in enumerate(table.columns):
                col_width[col_idx] = round(conver.twips_2_px(col.width) * ratio, 2) - 1
        else:
            for col_idx, col in enumerate(table.columns):
                col_width[col_idx] = round(conver.twips_2_px(col.width) * self.col_width_ratio, 2)

        space_left = round(6 * self.ratio, 2)
        space_right = round(6 * self.ratio, 2)

        # 边框样式暂时没有处理，默认全都有边框
        left_indent = round(self.get_table_left_indent(table) * self.ratio, 2)
        border_color = self.style_param.get("border_color", "#000000")
        o_table = widgets.LiteTable(data=data, min_row_height=min_row_height,
                                    # width="auto", max_width=self.width - self.padding[1] - self.padding[3],
                                    col_width=col_width,
                                    width=sum(col_width.values()),
                                    space_left=space_left, space_right=space_right,
                                    merged_cells=merged_cells,
                                    left_indent=left_indent,
                                    horizontal=horizontal,
                                    cell_horizontal=cell_horizontal, cell_vertical=cell_vertical,
                                    cell_style=cell_style,
                                    row_border={"color": border_color}, col_border={"color": border_color},
                                    border={"color": border_color}, auto_wrap=True,
                                    scroll={"x": "hidden"})
        return o_table

    # def get_paragraph_row_height(self, line_spacing_rule, line_spacing, font_size):
    #     # // ONE_POINT_FIVE = 1，1.5倍行距
    #     # // AT_LEAST = 3，最小行距
    #     # // DOUBLE = 2，双倍行距
    #     # // EXACTLY = 4，固定值
    #     # // MULTIPLE = 5，多倍行距
    #     # // SINGL = 0，单倍行距
    #     height = conver.font_size_2_height(font_size)
    #     if line_spacing_rule == 1:
    #         height = height * 1.5
    #     elif line_spacing_rule == 2:
    #         height = height * 2
    #     elif line_spacing_rule == 3:
    #         pass
    #     elif line_spacing_rule == 4:
    #         height = line_spacing
    #     elif line_spacing_rule == 5:
    #         height = height * line_spacing
    #     return height

    def parse_tabs(self, paragraph, o_paragraph):
        lst_node = paragraph._element.xpath('w:pPr/w:tabs/w:tab')
        if lst_node:
            node = lst_node[0]
            leader = node.get(qn("w:leader"))  # 暂时只处理dot
            if leader:
                width = 0
                font_size = o_paragraph.font.get("size")
                o_widget_t = None
                for o_widget in o_paragraph.children:
                    if o_widget.text == "\t":
                        o_widget_t = o_widget
                    else:
                        # text = word_api.replace_chinese_char(o_widget.text)
                        # w = (len(o_widget.text) - len(text)) * (o_widget.font.get("size") or font_size)
                        # width += w
                        # w = len(text) * (o_widget.font.get("size") or font_size) / 2.0
                        # width += w
                        w = word_api.get_text_width(o_widget.text, o_widget.font.get("size") or font_size,
                                                    is_pdf=self.is_pdf)
                        width += w

                if o_widget_t:
                    if width < self.content_width:
                        num = word_api.get_half_char(self.content_width - width,
                                                     o_widget_t.font.get("size") or font_size, self.is_pdf)
                        o_widget_t.text = "." * num
                    else:
                        o_widget_t.text = ""

    def make_widget(self, name, widget_width):
        o_widget = None
        name = name.strip()
        if name.lower().find("widget(") == 0:
            name = name.replace('”', '"').replace("，", ",")
            script = "widgets.W%s" % name[1:]
            o_widget = eval(script)

        elif name.lower().find("widgets.") == 0:
            name = name.replace('”', '"').replace("，", ",")
            o_widget = eval(name)

        if o_widget:
            if o_widget.get_attr_value("type", None):
                o_widget.type = o_widget.type.strip()
                if o_widget.type == "input" and o_widget.get_attr_value("input_type", None) == "textarea":
                    o_widget.border = {"color": "#FF0000"}
                    o_widget.clearable = False
                else:
                    o_widget.border = {"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0}
            else:
                o_widget.type = "input"
                o_widget.clearable = False
                o_widget.border = {"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0}

            if o_widget.get_attr_value("name", None):
                o_widget.name = o_widget.name.strip()
                o_widget.value = self.param.get(o_widget.name, None)

            if o_widget.get_attr_value("width", None) and isinstance(o_widget.width, str):
                if "%" in o_widget.width:
                    pass
                else:
                    o_widget.width = round(int(o_widget.width.strip()) * self.ratio, 2)

            if o_widget.get_attr_value("height", None) and isinstance(o_widget.height, str):
                o_widget.height = round(int(o_widget.height.strip()) * self.ratio, 2)

            if o_widget.get_attr_value("width", None) is None:
                if o_widget.type == "date_picker":
                    o_widget.width = round(140 * self.ratio, 2)
                elif o_widget.type == "time_picker":
                    o_widget.width = round(120 * self.ratio, 2)
                elif o_widget.type == "datetime_picker":
                    o_widget.width = round(200 * self.ratio, 2)

        elif self.param.get(name, None):
            value = self.param[name]
            if isinstance(value, widgets.Widget):
                o_widget = value
                if not o_widget.get_attr_value("name", None):
                    o_widget.name = name

            elif isinstance(value, dict):
                o_widget = value
                if not o_widget.get("name", None):
                    o_widget["name"] = name
            else:
                o_widget = widgets.Input(name=name, value=value, width=round(widget_width * self.ratio, 2),
                                         clearable=False,
                                         border={"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0},
                                         active={"border_color": "#000000"}, focus={"border_color": "#000000"})
        else:
            o_widget = widgets.Input(name=name, width=round(widget_width * self.ratio, 2), clearable=False,
                                     border={"color": "#FF0000", "width": [0, 0, 1, 0], "radius": 0},
                                     active={"border_color": "#000000"}, focus={"border_color": "#000000"})

        return o_widget

    def get_font_family(self, paragraph=None, run=None):
        if run:
            r = run._r
            lst_font = r.xpath('./w:rPr/w:rFonts')
        else:
            r = paragraph._p
            lst_font = r.xpath('./w:pPr/w:rPr/w:rFonts')

        family = None
        if lst_font:
            font = lst_font[0]
            family = font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or font.get(qn("w:hAnsi"))

        if family is None:
            if paragraph:
                # <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"
                # w:eastAsia="宋体" w:cs="Times New Roman"/>
                lst_font = paragraph.style.font._element.xpath('./w:rPr/w:rFonts')
                if lst_font:
                    font = lst_font[0]
                    family = font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or font.get(qn("w:hAnsi"))

                if family is None:
                    lst_node = paragraph.style.element.xpath("w:basedOn")
                    while lst_node:
                        style_id = lst_node[0].get(qn("w:val"))
                        style = self.dict_style[style_id]
                        lst_font = style.xpath('./w:rPr/w:rFonts')
                        if lst_font:
                            font = lst_font[0]
                            family = font.get(qn("w:eastAsia")) or font.get(qn("w:ascii")) or \
                                     font.get(qn("w:hAnsi"))
                            if family:
                                break
                        lst_node = style.xpath('w:basedOn')

        return family

    def get_font_size(self, paragraph=None, run=None):
        size = None
        if "font_size" in self.style_param:
            size = self.style_param["font_size"]
        elif paragraph:
            lst_sz = paragraph._p.xpath('./w:pPr/w:rPr/w:sz')
            lst_szCs = paragraph._p.xpath('./w:rPr/w:szCs')
            if lst_sz:
                size = round(conver.twips_2_px(lst_sz[0].val) * self.ratio, 2)
            elif lst_szCs:
                size = round(conver.twips_2_px(int(lst_szCs[0].get(qn("w:val"))) * 6350) * self.ratio, 2)
            elif paragraph.style.font.size:
                size = round(conver.twips_2_px(paragraph.style.font.size) * self.ratio, 2)

        elif run:
            lst_sz = run._element.xpath('./w:rPr/w:sz')
            lst_szCs = run._element.xpath('./w:rPr/w:szCs')
            if lst_sz:
                size = round(conver.twips_2_px(lst_sz[0].val) * self.ratio, 2)
            elif lst_szCs:
                size = round(conver.twips_2_px(int(lst_szCs[0].get(qn("w:val"))) * 6350) * self.ratio, 2)
            elif run.font.size:
                size = round(conver.twips_2_px(run.font.size) * self.ratio, 2)

        return size

    def get_font_style(self, paragraph=None, run=None):
        font = {}
        if "font_color" in self.style_param:
            font["color"] = self.style_param["font_color"]

        if paragraph:
            lst_node = paragraph._element.xpath("w:r/w:rPr/w:color")
            if lst_node:
                node = lst_node[0]
                font["color"] = "#%s" % str(node.val)

            lst_node = paragraph.style.element.xpath("w:rPr/w:b")
            if lst_node:
                font["weight"] = "bold"

        elif run:
            if run.font.color.rgb:
                if run.font.color.rgb != RGBColor(0, 0, 0):  # docx bug，实际word没有颜色
                    font["color"] = "#%s" % str(run.font.color.rgb)

            if run.font.underline:
                font["decoration"] = "underline"
                # o_text.under_line = font.underline

            if run.font.strike:
                font["decoration"] = "line-through"

            if run.bold or run.style.font.bold:
                font["weight"] = "bold"
            # else:
            #     lst_node = run.element.xpath("w:rPr/w:rFonts")
            #     if lst_node:
            #         if lst_node[0].get(qn("w:hint")) == "eastAsia": # 不确定是否正确，根据xml类容分析
            #             font["weight"] = "bold"

            if run.font.italic:
                font["style"] = "italic"

        return font

    def get_font_bg(self, paragraph=None, run=None):
        bg = {}
        if paragraph:
            lst_node = paragraph._element.xpath("w:r/w:rPr/w:highlight")
            if lst_node:
                node = lst_node[0]
                bg = {"color": conver.word_color_idx_2_rgb(node.val)}

        elif run:
            try:
                if run.font and run.font.highlight_color:
                    bg["color"] = conver.word_color_idx_2_rgb(run.font.highlight_color)
            except (BaseException,):
                pass
        return bg

    def get_table(self, paragraph, cell):
        for table in cell.tables:
            if paragraph._p.getprevious() == table._element:
                return table

    def get_table_merged_cells(self, table):
        """获取合并单元格"""
        row_num = len(table.rows)
        col_num = len(table.columns)

        merged_cells = []
        for row_idx in range(row_num):
            for col_idx in range(col_num):
                b_merged, b_range, row_span, col_span = \
                    self.get_table_merged_info(merged_cells, row_idx, col_idx)

                if b_merged or b_range:
                    continue

                cell = table.cell(row_idx, col_idx)
                end_row, end_col = self.get_table_merged_cell(table, cell, row_idx, col_idx)
                if (row_idx != end_row) or (col_idx != end_col):
                    # merged_cells.append("%s%s:%s%s" % (get_column_letter(col_idx + 1), row_idx + 1,
                    #                                    get_column_letter(c + 1), r + 1))
                    merged_cells.append("%s-%s:%s-%s" % (row_idx, col_idx, end_row, end_col))

        return merged_cells

    def get_table_merged_cell(self, table, cell, row_idx, col_idx):
        row_num = len(table.rows)
        col_num = len(table.columns)

        end_row = row_idx
        end_col = col_idx

        for r in range(row_idx + 1, row_num):
            if table.cell(r, col_idx)._tc == cell._tc:
                end_row = r
            else:
                break

        for c in range(col_idx + 1, col_num):
            if table.cell(row_idx, c)._tc == cell._tc:
                end_col = c
            else:
                break

        return end_row, end_col

    def get_table_merged_info(self, merged_cells, row_idx, col_idx):
        """获取合并的数据"""
        b_merged = False
        b_range = False
        row_span = 0
        col_span = 0
        row_col = '%s-%s' % (row_idx, col_idx)

        for s_range in merged_cells:
            s_range = str(s_range)
            begin_row_col, end_row_col = s_range.split(":")
            # begin_col, begin_row = re.findall("[A-Z]+|[0-9]+", begin_col_row, re.I)
            # end_col, end_row = re.findall("[A-Z]+|[0-9]+", end_col_row, re.I)
            begin_row, begin_col = begin_row_col.split("-")
            begin_row = int(begin_row)
            begin_col = int(begin_col)
            end_row, end_col = end_row_col.split("-")
            end_row = int(end_row)
            end_col = int(end_col)

            if begin_row_col == row_col:  # 开始合并
                b_merged = True
                row_span = end_row - begin_row + 1
                col_span = end_col - begin_col + 1
                break

            elif (row_idx >= begin_row) and (row_idx <= end_row) and \
                    (col_idx >= begin_col) and (col_idx <= end_col):
                b_range = True
                break

        return b_merged, b_range, row_span, col_span

    def get_table_row_height(self, row):
        if row.height:
            return conver.twips_2_px(row.height)

        cell = row.cells[0]
        tc = cell._tc
        # size = tc.xpath('./w:trPr/w:trHeight')
        lst_sz = tc.xpath('./w:p/w:pPr/w:rPr/w:sz')
        if lst_sz:
            sz = lst_sz[0]
            return conver.twips_2_px(sz.val)

    def get_table_cell_horizontal(self, cell):
        horizontal = "left"
        lst_horizontal = cell._tc.xpath("./w:p/w:pPr/w:jc")
        if lst_horizontal:
            horizontal = lst_horizontal[0].val
            if horizontal == WD_TABLE_ALIGNMENT.LEFT:
                horizontal = "left"
            elif horizontal == WD_TABLE_ALIGNMENT.CENTER:
                horizontal = "center"
            elif horizontal == WD_TABLE_ALIGNMENT.RIGHT:
                horizontal = "right"
            else:
                horizontal = "left"
        return horizontal

    def get_table_cell_background_color(self, cell):
        lst_shd = cell._tc.xpath('./w:tcPr/w:shd')
        if lst_shd:
            element = lst_shd[0]
            fill = element.get(qn("w:fill"))
            return "#%s" % fill

    def get_table_left_indent(self, table):
        left_indent = 0
        lst_path = table._tbl.xpath('./w:tblPr/w:tblInd')
        if lst_path:
            left_indent = conver.in_2_px(
                int(lst_path[0].get(qn("w:w"))) * 1 / 1440)  # dxa-指定该值的单位为点的二十分之一（1/1440英寸）
        else:
            lst_path = table._tbl.xpath('./w:tblPr/w:tblpPr')
            if lst_path:
                left_indent = conver.in_2_px(int(lst_path[0].get(qn("w:tblpX"))) / 1400) * -1
        return left_indent

    def get_image_id(self, run):
        namespace = {
            'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
            'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

        if hasattr(run.element, "drawing_lst"):
            for drawing in run.element.drawing_lst:
                for node in drawing.findall('.//a:blip', namespace):
                    for k, v in node.items():
                        if k.find("}embed") >= 0:
                            return v

    def get_image_size(self, run, rId):
        width = None
        height = None
        for shape in self.doc.inline_shapes:
            if rId == shape._inline.graphic.graphicData.pic.blipFill.blip.embed:
                if shape.width:
                    width = conver.twips_2_px(shape.width)
                    height = conver.twips_2_px(shape.height)
                break

        if width is None and height is None:
            lst_node = run._element.xpath('./w:drawing/wp:anchor/wp:extent')
            if lst_node:
                node = lst_node[0]
                width = conver.twips_2_px(int(node.get("cx")))
                height = conver.twips_2_px(int(node.get("cy")))

        return width, height

    def get_number(self, paragraph):
        """解析编号"""
        num_id, ilvl = None, None
        lst_num = paragraph._p.xpath('./w:pPr/w:numPr/w:numId')
        if lst_num:
            num_id = lst_num[0].val

        lst_lvl = paragraph._p.xpath('./w:pPr/w:numPr/w:ilvl')
        if lst_lvl:
            ilvl = lst_lvl[0].val

        if ilvl is None:
            lst_num = paragraph.style._element.xpath('./w:pPr/w:numPr/w:numId')
            if lst_num:
                num_id = lst_num[0].val

            lst_lvl = paragraph.style._element.xpath('./w:pPr/w:numPr/w:ilvl')
            if lst_lvl:
                ilvl = lst_lvl[0].val

        if (ilvl is not None) and (num_id is None):
            num_id = 1

        return num_id, ilvl

    def get_hyperlink(self, paragraph):
        """
        链接有两种格式
        1、在paragraph的run中存在，但没有链接内容，通过style_id中的name_val判断是否为链接文本
        2、在paragraph中不存在，在paragraph中查找是否有链接结点，判断是否在run之前
        """
        o_step = None
        lst_node = paragraph._element.xpath('w:r/w:instrText')
        for node in lst_node:
            hyperlink = node.text
            hyperlink = hyperlink.strip()
            if not hyperlink:
                continue

            if "HYPERLINK" in hyperlink:
                name = re.search('"(.*)"', hyperlink).groups()[0]
                self.lst_hyperlink.append(name)
                o_step = step.AnchorPoint(name=name)
                break

            elif hyperlink.find("http") == 0:
                href = re.search('(.*?)"', hyperlink).groups()[0]
                o_step = step.Get(href=href, new_window=True, unique=True)
                break

        # if hyperlink:
        #     try:
        #         import shutil
        #         from django.conf import settings
        #         path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, "hyperlink"))
        #         if not os.path.exists(path):
        #             os.makedirs(path)
        #         file_path = os.path.join(path, hyperlink)
        #         shutil.copy(os.path.join(os.path.split(self.path)[0], hyperlink), file_path)
        #         hyperlink = "%shyperlink/%s" % (settings.MEDIA_URL, hyperlink.strip("/\\"))
        #     except (BaseException,):
        #         pass

        return o_step

        # if not run.style:
        #     return None, None
        #
        # for style in self.doc.styles.element.style_lst:
        #     if run.style.style_id == style.styleId:
        #         if style.name_val == "Hyperlink":
        #             p = paragraph._p
        #             lst_instr = p.xpath('./w:r/w:instrText')
        #             if lst_instr:
        #                 instr = lst_instr[0]
        #                 o_re = re.search('HYPERLINK "(.*)" ', instr.text, re.I)
        #                 if o_re:
        #                     hyperlink = o_re.groups()[0]
        #                 else:
        #                     hyperlink = instr.text
        #
        #                 return hyperlink, None
        #         break

        # lst_hyperlink = paragraph._p.xpath("./w:hyperlink")
        # if lst_hyperlink:
        #     for hyperlink in lst_hyperlink:
        #         if hyperlink.getnext() == run.element:
        #             rId = hyperlink.get(qn("r:id"))
        #             hyperlink_run = None
        #             for node in lst_hyperlink[0].getchildren():
        #                 # hyperlink_text += node.text
        #                 hyperlink_run = Run(node, paragraph)
        #                 break
        #
        #             for rel in self.doc.part.rels:
        #                 if rel == rId and \
        #                         self.doc.part.rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
        #                     return self.doc.part.rels[rel]._target, hyperlink_run

        # return None, None

    def get_bg_image(self, run):
        rId, width, height = None, None, None
        lst_element = run._r.xpath('./w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip')
        if lst_element:
            element = lst_element[0]
            rId = element.get(qn("r:embed"))

        lst_element = run._r.xpath('./w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic/pic:spPr/a:xfrm/a:ext')
        if lst_element:
            element = lst_element[0]
            width = element.get("cx")
            height = element.get("cy")
            if width is not None:
                width = conver.twips_2_px(int(width))

            if height is not None:
                height = conver.twips_2_px(int(height))

        return rId, width, height

    def get_left_indent(self, paragraph):
        left_indent = 0
        if paragraph.paragraph_format.left_indent:  # 块全部左缩进
            left_indent = conver.twips_2_px(paragraph.paragraph_format.left_indent)
        elif paragraph.style.paragraph_format.left_indent:
            left_indent = conver.twips_2_px(paragraph.style.paragraph_format.left_indent)
        return left_indent

    def get_right_indent(self, paragraph):
        right_indent = 0
        if paragraph.paragraph_format.right_indent:
            right_indent = conver.twips_2_px(paragraph.paragraph_format.right_indent)
        elif paragraph.style.paragraph_format.right_indent:
            right_indent = conver.twips_2_px(paragraph.style.paragraph_format.right_indent)
        return right_indent

    def get_first_line_indent(self, paragraph):
        first_line_indent = 0
        if paragraph.paragraph_format.first_line_indent is not None:
            first_line_indent = conver.twips_2_px(paragraph.paragraph_format.first_line_indent)
        elif paragraph.style.paragraph_format.first_line_indent:
            first_line_indent = conver.twips_2_px(paragraph.style.paragraph_format.first_line_indent)
        return first_line_indent

    def get_space_before(self, paragraph):
        space_before = 0
        if paragraph.paragraph_format.space_before:
            space_before = conver.twips_2_px(paragraph.paragraph_format.space_before)
        elif paragraph.style.paragraph_format.space_before:
            space_before = conver.twips_2_px(paragraph.style.paragraph_format.space_before)
        else:
            p = paragraph._p
            lst_xpath = p.xpath('./w:pPr/w:spacing')
            if lst_xpath:
                before = lst_xpath[0].get(qn("w:beforeLines"))
                if before:
                    space_before = float(before)

        return space_before

    def get_space_after(self, paragraph):
        space_after = 0
        if paragraph.paragraph_format.space_after:
            space_after = int(conver.twips_2_px(paragraph.paragraph_format.space_after))
        elif paragraph.style.paragraph_format.space_after:
            space_after = conver.twips_2_px(paragraph.style.paragraph_format.space_after)
        else:
            p = paragraph._p
            lst_xpath = p.xpath('./w:pPr/w:spacing')
            if lst_xpath:
                after = lst_xpath[0].get(qn("w:afterLines"))
                if after:
                    space_after = float(after)

        return space_after

    def get_line_spacing(self, paragraph):
        line_spacing = 0
        line_spacing_rule = 0
        if paragraph.paragraph_format.line_spacing:
            line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
            if line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                line_spacing = conver.twips_2_px(paragraph.paragraph_format.line_spacing)
            elif line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                # < w: pPr >
                # < w: spacing
                # w: line = "600"
                # w: lineRule = "atLeast" / >
                line_spacing = conver.twips_2_px(paragraph.paragraph_format.line_spacing)
            else:
                line_spacing = paragraph.paragraph_format.line_spacing
            line_spacing_rule = int(line_spacing_rule)

        elif paragraph.style.paragraph_format.line_spacing:
            line_spacing_rule = paragraph.style.paragraph_format.line_spacing_rule
            if line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                line_spacing = conver.twips_2_px(paragraph.style.paragraph_format.line_spacing)
            elif line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                line_spacing = conver.twips_2_px(paragraph.paragraph_format.line_spacing)
            else:
                line_spacing = paragraph.style.paragraph_format.line_spacing
            line_spacing_rule = int(line_spacing_rule)
        else:
            lst_node = paragraph.style.element.xpath("w:basedOn")
            if lst_node:
                style_id = lst_node[0].get(qn("w:val"))
                style = self.dict_style[style_id]
                # {'atLeast': 3, 'exact': 4, 'auto': 5}
                # Twips(int(str_value))
                lst_node = style.xpath("w:pPr/w:spacing")
                if lst_node:
                    line = lst_node[0].get(qn("w:line"))
                    lineRule = lst_node[0].get(qn("w:lineRule"))
                    if line or lineRule:
                        line = Twips(int(line))
                        lineRule = {'atLeast': 3, 'exact': 4, 'auto': 5}.get(lineRule, None)
                        if lineRule == WD_LINE_SPACING.MULTIPLE:
                            if line == Twips(240):
                                line_spacing_rule = WD_LINE_SPACING.SINGLE
                            elif line == Twips(360):
                                line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                            elif line == Twips(480):
                                line_spacing_rule = WD_LINE_SPACING.DOUBLE

                        if line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                            line_spacing = conver.twips_2_px(line)
                        elif line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                            line_spacing = conver.twips_2_px(line)
                        else:
                            line_spacing = line
                        line_spacing_rule = int(line_spacing_rule)

        return line_spacing_rule, line_spacing


def generate(word_path, param=None, show_width=None, padding=None, style_param=None, cache=False):
    """
    template_file：模板文件(只支持*.docx)
    param:参数 {"模板关键字":"显示值", "模板关键字":"widget对象"}
    show_width:
    padding:页边距
    """
    import json
    from vadmin import common
    from vadmin.json_encoder import Encoder
    path, name = os.path.split(word_path)
    file_name = os.path.splitext(name)[0]
    file_name = common.format_file_name(file_name)
    s_padding = str(padding).replace("[", "").replace(",", "").replace("]", "").replace(" ", "")
    widget_path = os.path.join(path, "%s.%s.vadmin" % (file_name, s_padding))
    if cache:
        if os.path.exists(widget_path):
            buf = open(widget_path).read()
            try:
                o_panel = json.loads(buf)
                return o_panel
            except (BaseException,):
                pass

    if show_width is not None:
        show_width = int(show_width)

    if padding:
        obj = Compiler(word_path, param, show_width - padding[1] - padding[3], style_param=style_param)
    else:
        obj = Compiler(word_path, param, show_width, style_param=style_param)

    # begin = time.time()
    lst_widget = obj.parse()
    # print("time:%s" % (time.time() - begin))
    if (show_width is None) or (show_width > obj.page_width):
        width = obj.page_width
    else:
        width = show_width

    if padding:
        o_panel = widgets.Panel(vertical="top", horizontal="left", width=width,
                                padding=padding,
                                bg_color="#FFFFFF", height="100%", scroll={"y": "auto"}, min_height=obj.page_height)
    else:
        o_panel = widgets.Panel(vertical="top", horizontal="left", width=width,
                                padding=[obj.top_margin, obj.right_margin, obj.bottom_margin, obj.left_margin],
                                bg_color="#FFFFFF", height="100%", scroll={"y": "auto"}, min_height=obj.page_height)

    o_panel.append(lst_widget)
    if cache:
        open(widget_path, "w").write(json.dumps(o_panel, ensure_ascii=False, cls=Encoder))
    return o_panel


def generate_word(template_path, param=None, show_width=None, is_margin=True):
    obj = Compiler(template_path, param, show_width, is_word_widget=True)
    begin = time.time()
    lst_widget = obj.parse()
    print("time:%s" % (time.time() - begin))
    # if (show_width is None) or (show_width > obj.page_width):
    #     width = obj.page_width
    # else:
    #     width = show_width

    if is_margin:
        o_word = widgets.Word(top_margin=obj.top_margin,
                              right_margin=obj.right_margin,
                              bottom_margin=obj.bottom_margin,
                              left_margin=obj.left_margin,
                              bg_color="#FFFFFF")
    else:
        o_word = widgets.Word(bg_color="#FFFFFF")

    o_word.page_width = obj.page_width
    o_word.page_height = obj.page_height
    o_word.header_distance = obj.header_distance
    o_word.footer_distance = obj.footer_distance
    o_word.paragraph_symbol = {}

    for (num_id, lst_val) in obj.dict_abstract.items():
        if len(lst_val) > 0:
            o_word.paragraph_symbol[lst_val[0]["abstract_id"]] = lst_val

    o_word.append(lst_widget)
    return o_word


if __name__ == "__main__":
    generate(r"C:\wangbin\code_sync\code_sync\project_management\templates\winning.docx", is_compile=True)
